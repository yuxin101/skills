# template_engine/executor.py - 修复版本
import os
import re
import logging
import pandas as pd
import openpyxl
from openpyxl.cell import MergedCell
from copy import deepcopy
from typing import Dict, Any, List
from docx import Document
from docxtpl import DocxTemplate
import shutil

from .core import *
from .utils import TemplateUtils

logger = logging.getLogger(__name__)

class BaseGenerator:
    """基础生成器 - 统一生成接口"""
    
    def fill_placeholders(self, file_path: str, data: pd.DataFrame, 
                         config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充占位符 - 由具体生成器实现"""
        raise NotImplementedError
    
    def fill_table_data(self, file_path: str, data: pd.DataFrame,
                       config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充表格数据 - 由具体生成器实现"""
        raise NotImplementedError

class ExcelGenerator(BaseGenerator):
    """Excel 生成器实现"""
    
    def fill_placeholders(self, file_path: str, data: pd.DataFrame, 
                         config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充 Excel 占位符"""
        try:
            wb = openpyxl.load_workbook(file_path)
            placeholders = config.get('placeholders', [])
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                # 查找并替换占位符
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for mapping in placeholders:
                                if mapping.placeholder in cell.value:
                                    # 使用第一行数据填充占位符
                                    value = data.iloc[0].get(mapping.data_field, mapping.default_value)
                                    formatted_value = TemplateUtils.format_value(value, mapping.data_type)
                                    cell.value = cell.value.replace(mapping.placeholder, str(formatted_value))
            
            wb.save(file_path)
            logger.info(f"✅ Excel占位符填充完成: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ Excel占位符填充失败: {e}")
            raise
    
    def fill_table_data(self, file_path: str, data: pd.DataFrame,
                       config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充 Excel 表格数据"""
        try:
            logger.info(f"📊 开始填充表格数据，文件: {file_path}, 数据量: {len(data)}")
            logger.info(f"📋 表格配置: {config}")
            logger.info(f"📝 数据列: {list(data.columns)}")
            logger.info(f"📄 数据预览:\n{data.head()}")
            
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            
            insert_row = config.get('insert_row', 1)
            insert_col = config.get('insert_col', 1)
            reserved_rows = config.get('reserved_rows', 1)
            
            logger.info(f"📍 插入位置: 行{insert_row}, 列{insert_col}, 保留行: {reserved_rows}")
            
            # 计算需要插入的行数
            data_count = len(data)
            needed_rows = data_count - reserved_rows
            
            logger.info(f"📏 数据行数: {data_count}, 需要插入: {needed_rows} 行")
            
            if needed_rows > 0:
                insert_pos = insert_row + reserved_rows
                logger.info(f"🔧 在位置 {insert_pos} 插入 {needed_rows} 行")
                ws.insert_rows(insert_pos, amount=needed_rows)
                TemplateUtils.copy_merged_cells(ws, insert_row, insert_pos, reserved_rows)
            
            # 插入数据
            logger.info("🖊️ 开始写入数据...")
            for row_offset, (_, row_data) in enumerate(data.iterrows()):
                row_idx = insert_row + row_offset
                logger.info(f"📝 写入第 {row_idx} 行数据: {dict(row_data)}")
                for col_offset, value in enumerate(row_data):
                    cell = ws.cell(row=row_idx, column=insert_col + col_offset)
                    if not isinstance(cell, MergedCell):
                        sanitized_value = TemplateUtils.sanitize_value(value)
                        cell.value = sanitized_value
                        logger.debug(f"  └─ 单元格({row_idx}, {insert_col + col_offset}) = {sanitized_value}")
            
            wb.save(file_path)
            wb.close()
            logger.info(f"✅ Excel表格数据填充完成: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ Excel表格数据填充失败: {e}")
            raise
class WordGenerator(BaseGenerator):
    """Word 生成器实现 - 混合模式版本"""

    def fill_placeholders(self, file_path: str, data: pd.DataFrame, 
                         config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充 Word 占位符 - 混合模式第一步"""
        try:
            logger.info(f"🔄 DOCX混合模式步骤1: 占位符填充，文件: {file_path}, 数据键: {data_key_value}")
            
            # 构建智能上下文字典
            context = self._build_smart_context(data, config.get('data_key', ''))
            
            logger.info(f"🎯 占位符上下文: {list(context.keys())}")
            
            # 使用 docxtpl 渲染占位符
            tpl = DocxTemplate(file_path)
            tpl.render(context)
            tpl.save(file_path)  # 必须保存以生效
            
            logger.info(f"✅ DOCX占位符填充完成: {file_path}")
            return file_path
                
        except Exception as e:
            logger.error(f"❌ DOCX占位符填充失败: {e}")
            raise

    def fill_table_data(self, file_path: str, data: pd.DataFrame,
                       config: Dict[str, Any], data_key_value: Any = None) -> str:
        """填充 Word 表格数据 - 混合模式第二步"""
        try:
            logger.info(f"📊 DOCX混合模式步骤2: 表格填充，文件: {file_path}, 数据量: {len(data)}")
            
            # 复用XLSX参数体系
            table_index = config.get('reserved_rows', 0)  # reserved_rows → 表格索引
            start_row = config.get('insert_row', 1)       # insert_row → 表格内起始行
            start_col = config.get('insert_col', 1)       # insert_col → 表格内起始列
            
            logger.info(f"📍 表格操作: 第{table_index}个表格, 起始行:{start_row}, 起始列:{start_col}")
            
            # 加载已进行占位符填充的文档
            doc = Document(file_path)
            
            # 检查表格是否存在
            if table_index >= len(doc.tables):
                raise ValueError(f"文档中只有{len(doc.tables)}个表格，无法访问第{table_index}个表格")
            
            table = doc.tables[table_index]
            logger.info(f"📋 目标表格: {len(table.rows)}行 x {len(table.columns)}列")
            
            # 计算需要添加的行数
            data_count = len(data)
            existing_data_rows = len(table.rows) - start_row  # 现有数据行数
            needed_rows = data_count - existing_data_rows
            
            logger.info(f"📏 数据行数: {data_count}, 需要添加: {needed_rows}行")
            
            # 动态添加行
            if needed_rows > 0:
                for i in range(needed_rows):
                    table.add_row()
                logger.info(f"✅ 成功添加{needed_rows}行")
            
            # 填充数据
            logger.info("🖊️ 开始填充表格数据...")
            for row_offset, (_, row_data) in enumerate(data.iterrows()):
                row_idx = start_row + row_offset
                
                # 检查行索引是否有效
                if row_idx >= len(table.rows):
                    logger.warning(f"⚠️ 行索引{row_idx}超出表格范围，跳过")
                    continue
                    
                for col_offset, value in enumerate(row_data):
                    # 检查列索引是否有效
                    if col_offset >= len(table.rows[row_idx].cells):
                        logger.warning(f"⚠️ 列索引{col_offset}超出表格范围，跳过")
                        continue
                        
                    cell = table.rows[row_idx].cells[col_offset]
                    sanitized_value = TemplateUtils.sanitize_value(value)
                    
                    # 清空单元格原有内容
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                    
                    # 添加新内容
                    cell.text = str(sanitized_value)
                    
                    logger.debug(f"  └─ 单元格({row_idx},{col_offset}) = {sanitized_value}")
            
            # 保存文档 - 混合模式最终保存
            doc.save(file_path)
            logger.info(f"✅ DOCX表格填充完成: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ DOCX表格填充失败: {e}")
            raise

    def _build_smart_context(self, group_data: pd.DataFrame, data_key: str) -> Dict[str, Any]:
        """构建智能字段分类的上下文字典"""
        context = {}
        
        # 如果只有一行数据，使用简单模式
        if len(group_data) == 1:
            logger.info("📝 单行数据模式，使用简单字典转换")
            context = group_data.iloc[0].to_dict()
        else:
            logger.info(f"🔍 多行数据模式，智能字段分类，数据行数: {len(group_data)}")
            
            # 分析每个字段的特征
            single_value_fields = []
            multi_value_fields = []
            
            for col in group_data.columns:
                if col == data_key:
                    continue
                    
                unique_count = group_data[col].nunique()
                total_count = len(group_data)
                
                logger.debug(f"  ├─ 字段 '{col}': 唯一值数={unique_count}, 总行数={total_count}")
                
                if unique_count == 1:
                    # 所有行值相同 → 单值字段
                    single_value_fields.append(col)
                else:
                    # 值不同 → 多值字段
                    multi_value_fields.append(col)
            
            # 处理单值字段：取第一行的值
            for col in single_value_fields:
                context[col] = group_data[col].iloc[0]
                logger.debug(f"  ├─ 单值字段 '{col}': {context[col]}")
            
            # 处理多值字段：构建记录列表
            if multi_value_fields:
                multi_value_data = group_data[multi_value_fields]
                context['records'] = multi_value_data.to_dict('records')
                context['record_count'] = len(multi_value_data)
                
                logger.info(f"  ├─ 多值字段: {multi_value_fields}")
                logger.info(f"  ├─ 生成记录数: {context['record_count']}")
        
        # 确保所有键都是字符串
        context = {str(key): value for key, value in context.items()}
        
        return context


class TemplateExecutor:
    """模板执行器 - 修复版本"""
    
    # 生成器实例映射
    _GENERATORS = {
        '.xlsx': ExcelGenerator(),
        '.xls': ExcelGenerator(),
        '.docx': WordGenerator(),
        '.doc': WordGenerator()
    }
    
    @staticmethod
    def execute_single_file_generation(task_config: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行单文件生成 - 供并行任务调用
        
        Args:
            task_config: 任务配置
            
        Returns:
            Dict[str, Any]: 执行结果
        """
        data_key_value = task_config.get('_data_key_value', 'unknown')
        
        try:
            data_group = task_config['_data_group']
            
            logger.info(f"执行并行任务，数据键: {data_key_value}, 数据量: {len(data_group)}")
            
            # 创建执行计划 - 使用 TemplateConfigEngine
            plan = TemplateExecutor._create_direct_plan(task_config, data_group, data_key_value)
            
            # 执行生成计划
            result = TemplateExecutor.execute_generation_plan(plan)
            
            return {
                'status': result.status,
                'message': result.message,
                'generated_files': result.generated_files,
                'data_key': data_key_value,
                'data_count': len(data_group),
                'output_dir': result.output_dir,
                'execution_log': result.execution_log
            }
            
        except Exception as e:
            logger.error(f"并行任务执行失败 {data_key_value}: {e}")
            return {
                'status': 'error',
                'message': str(e),
                'data_key': data_key_value,
                'generated_files': []
            }
    
    @staticmethod
    def execute_serial_generation(frontend_config: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行串行文件生成 - 保持兼容性（理论上不再使用）
        """
        try:
            # 导入在方法内部避免循环导入
            from .config_engine import TemplateConfigEngine
            plan = TemplateConfigEngine.create_execution_plan(frontend_config)
            result = TemplateExecutor.execute_generation_plan(plan)
            
            return {
                'status': result.status,
                'message': result.message,
                'generated_files': result.generated_files,
                'output_dir': result.output_dir,
                'execution_log': result.execution_log,
                'download_info': result.download_info,
                'execution_mode': 'serial'
            }
            
        except Exception as e:
            logger.error(f"串行文件生成执行失败: {e}")
            return {
                'status': 'error',
                'message': f'串行文件生成执行失败: {e}'
            }
    
    @staticmethod
    def execute_generation_plan(plan: ExecutionPlan) -> GenerationResult:
        """
        执行生成计划 - 统一处理逻辑
        
        Args:
            plan: 执行计划
            
        Returns:
            GenerationResult: 生成结果
        """
        logger.info(f"🚀 开始执行生成计划，模式: {plan.mode.value}")
        
        try:
            # 加载数据源（支持内存数据）
            data_context = TemplateExecutor._load_data_sources(plan.data_sources, plan)
            main_data = data_context.get('main_data')
            
            if main_data is None:
                raise ValueError("缺少主数据")
            
            data_key_value = getattr(plan, '_data_key_value', 'unknown')
            
            # 准备输出目录
            os.makedirs(plan.output_dir, exist_ok=True)
            
            # 生成文件名
            file_extension = os.path.splitext(plan.template_path)[1]
            safe_name = re.sub(r'[\\/*?:"<>|]', "", str(data_key_value))
            output_filename = f"{safe_name}{file_extension}"
            output_path = os.path.join(plan.output_dir, output_filename)
            
            # 复制模板文件
            shutil.copy2(plan.template_path, output_path)
            
            # 执行步骤并应用到文件
            execution_log = TemplateExecutor._execute_steps_and_apply_to_file(
                output_path, plan.steps, data_context, data_key_value
            )
            
            generated_files = [output_filename]
            
            return GenerationResult(
                status='success',
                message=f'文件生成完成: {output_filename}',
                generated_files=generated_files,
                output_dir=plan.output_dir,
                execution_log=execution_log,
                download_info={
                    'type': 'single',
                    'files': generated_files,
                    'output_dir': plan.output_dir
                }
            )
            
        except Exception as e:
            logger.error(f"❌ 执行生成计划失败: {e}")
            return GenerationResult(
                status='error',
                message=str(e),
                generated_files=[],
                output_dir=plan.output_dir
            )
    
    @staticmethod
    def _create_direct_plan(config: Dict[str, Any], data_group: pd.DataFrame, data_key_value: Any) -> ExecutionPlan:
        """为并行任务创建执行计划 - 使用 TemplateConfigEngine 统一处理"""
        try:
            # 导入 TemplateConfigEngine
            from .config_engine import TemplateConfigEngine
            
            # 创建基础配置（排除并行任务内部字段）
            base_config = {}
            for k, v in config.items():
                if not k.startswith('_'):
                    base_config[k] = v
            
            logger.info(f"🔍 使用 TemplateConfigEngine 创建计划，模式: {base_config.get('mode')}")
            
            # 使用 TemplateConfigEngine 创建执行计划
            plan = TemplateConfigEngine.create_execution_plan(base_config)
            
            # 存储内存数据和键值
            plan._memory_data = data_group
            plan._data_key_value = data_key_value
            
            # 修改数据源为内存数据
            for data_source in plan.data_sources:
                if data_source.name == "main_data":
                    data_source.file_path = "memory"
                    logger.info(f"📁 修改数据源为内存数据: {data_source.name}")
            
            logger.info(f"✅ 执行计划创建完成，模式: {plan.mode.value}, 步骤数: {len(plan.steps)}")
            return plan
            
        except Exception as e:
            logger.error(f"❌ 创建执行计划失败: {e}")
            raise
    
    @staticmethod
    def _load_data_sources(data_sources: List[DataSource], plan: ExecutionPlan = None) -> Dict[str, pd.DataFrame]:
        """加载数据源 - 支持内存数据"""
        data_context = {}
        
        for data_source in data_sources:
            try:
                # 如果是内存数据，使用 plan 中存储的数据
                if data_source.file_path == "memory" and plan and hasattr(plan, '_memory_data'):
                    df = plan._memory_data
                    logger.info(f"📁 使用内存数据: {data_source.name}, 数据量: {len(df)}")
                else:
                    # 否则从文件加载
                    df = pd.read_excel(data_source.file_path, sheet_name=data_source.sheet_name, dtype=str)
                    logger.info(f"📁 加载数据源成功: {data_source.name}, 数据量: {len(df)}")
                
                data_context[data_source.name] = df
                
            except Exception as e:
                logger.error(f"❌ 加载数据源失败 {data_source.name}: {e}")
                raise
        
        return data_context
    
    @staticmethod
    def _execute_steps_and_apply_to_file(file_path: str, steps: List[ExecutionStep], 
                                       data_context: Dict[str, pd.DataFrame], data_key_value: Any) -> List[Dict]:
        """执行步骤并应用到文件 - 返回执行日志"""
        execution_log = []
        
        for step in steps:
            try:
                logger.info(f"🎯 开始执行步骤: {step.step_id} - {step.description}")
                # 获取步骤数据
                step_data = data_context.get(step.data_source)
                if step_data is None:
                    raise ValueError(f"数据源不存在: {step.data_source}")
                
                # 使用新的方法应用步骤到文件
                TemplateExecutor._apply_step_to_file(file_path, step, step_data, data_key_value)
                
                execution_log.append({
                    'step_id': step.step_id,
                    'description': step.description,
                    'status': 'success',
                    'result': f"步骤 {step.step_id} 执行成功"
                })
                
                logger.info(f"✅ 步骤执行成功: {step.description}")
                
            except Exception as e:
                execution_log.append({
                    'step_id': step.step_id,
                    'description': step.description,
                    'status': 'error',
                    'error': str(e)
                })
                
                logger.error(f"❌ 步骤执行失败: {step.step_id} - {e}")
                if step.error_policy == "stop":
                    raise  # 重新抛出异常以停止执行
        
        return execution_log
    
    @staticmethod
    def _apply_step_to_file(file_path: str, step: ExecutionStep, data: pd.DataFrame, data_key_value: Any):
        """将单个步骤应用到文件"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 获取对应的生成器
        generator = TemplateExecutor._GENERATORS.get(file_ext)
        if not generator:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        # 根据函数名执行对应的生成方法
        function_name = step.function_name
        logger.info(f"🔧 执行步骤: {function_name}, 文件: {file_path}")

        if function_name == "fill_placeholders":
            generator.fill_placeholders(file_path, data, step.config, data_key_value)
        elif function_name == "fill_table_data":
            generator.fill_table_data(file_path, data, step.config, data_key_value)
        else:
            raise ValueError(f"未知的函数名: {function_name}")
    
    @staticmethod
    def _analyze_template_placeholders_direct(template_path: str) -> List[str]:
        """直接分析模板占位符 - 避免循环导入"""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        file_ext = os.path.splitext(template_path)[1].lower()
        placeholders = []
        placeholder_pattern = r'\{\{(.*?)\}\}'
        
        if file_ext in ['.docx', '.doc']:
            try:
                doc = Document(template_path)
                # 检查段落
                for paragraph in doc.paragraphs:
                    matches = re.findall(placeholder_pattern, paragraph.text)
                    placeholders.extend(matches)
                # 检查表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            matches = re.findall(placeholder_pattern, cell.text)
                            placeholders.extend(matches)
            except Exception as e:
                logger.error(f"分析Word模板占位符失败: {e}")
                raise
                
        elif file_ext in ['.xlsx', '.xls']:
            try:
                wb = openpyxl.load_workbook(template_path)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                matches = re.findall(placeholder_pattern, cell.value)
                                placeholders.extend(matches)
            except Exception as e:
                logger.error(f"分析Excel模板占位符失败: {e}")
                raise
        
        return list(set(placeholders))
    
    # ==================== 保持原有公共接口兼容性 ====================
    
    @staticmethod
    def fill_placeholders(template_path: str, config: Dict[str, Any], 
                         data: pd.DataFrame, output_dir: str, data_key_value: Any = None) -> str:
        """填充占位符 - 保持原有接口"""
        file_ext = os.path.splitext(template_path)[1].lower()
        generator = TemplateExecutor._GENERATORS.get(file_ext)
        
        if not generator:
            raise ValueError(f"不支持的模板格式: {file_ext}")
        
        # 使用数据键值生成文件名
        if data_key_value:
            safe_name = re.sub(r'[\\/*?:"<>|]', "", str(data_key_value))
            output_filename = f"{safe_name}{file_ext}"
            output_path = os.path.join(output_dir, output_filename)
            shutil.copy2(template_path, output_path)
            return generator.fill_placeholders(output_path, data, config, data_key_value)
        else:
            # 原有逻辑
            return generator.fill_placeholders(template_path, data, config, data_key_value)
    
    @staticmethod
    def fill_table_data(template_path: str, config: Dict[str, Any], 
                       data: pd.DataFrame, output_dir: str, data_key_value: Any = None) -> str:
        """填充表格数据 - 保持原有接口"""
        file_ext = os.path.splitext(template_path)[1].lower()
        generator = TemplateExecutor._GENERATORS.get(file_ext)
        
        if not generator:
            raise ValueError(f"不支持的模板格式: {file_ext}")
        
        # 使用数据键值生成文件名
        if data_key_value:
            safe_name = re.sub(r'[\\/*?:"<>|]', "", str(data_key_value))
            output_filename = f"{safe_name}{file_ext}"
            output_path = os.path.join(output_dir, output_filename)
            shutil.copy2(template_path, output_path)
            return generator.fill_table_data(output_path, data, config, data_key_value)
        else:
            return generator.fill_table_data(template_path, data, config, data_key_value)
        



# class WordGenerator(BaseGenerator):
#     """Word 生成器实现 - 智能字段分类版本"""
    
#     def fill_placeholders(self, file_path: str, data: pd.DataFrame, 
#                          config: Dict[str, Any], data_key_value: Any = None) -> str:
#         """填充 Word 占位符 - 智能字段分类版本"""
#         try:
#             logger.info(f"🔄 开始填充Word占位符，文件: {file_path}, 数据键: {data_key_value}")
#             logger.info(f"📊 输入数据形状: {data.shape}, 列: {list(data.columns)}")
            
#             # 获取数据键名
#             data_key = config.get('data_key', '')
            
#             # 构建智能上下文字典
#             context = self._build_smart_context(data, data_key)
            
#             logger.info(f"🎯 智能上下文构建完成: {list(context.keys())}")
#             logger.info(f"📋 单值字段: {[k for k in context.keys() if k not in ['records', 'record_count']]}")
#             if 'records' in context:
#                 logger.info(f"📊 多值记录数: {context.get('record_count', 0)}")
            
#             # 使用 docxtpl 渲染
#             tpl = DocxTemplate(file_path)
#             tpl.render(context)
#             tpl.save(file_path)
            
#             logger.info(f"✅ Word占位符填充完成: {file_path}")
#             return file_path
                
#         except Exception as e:
#             logger.error(f"❌ Word占位符填充失败: {e}")
#             raise
    
#     def _build_smart_context(self, group_data: pd.DataFrame, data_key: str) -> Dict[str, Any]:
#         """构建智能字段分类的上下文字典"""
#         context = {}
        
#         # 如果只有一行数据，使用简单模式（向后兼容）
#         if len(group_data) == 1:
#             logger.info("📝 单行数据模式，使用简单字典转换")
#             context = group_data.iloc[0].to_dict()
#             # 确保所有键都是字符串
#             context = {str(key): value for key, value in context.items()}
#             return context
        
#         logger.info(f"🔍 多行数据模式，开始智能字段分类，数据行数: {len(group_data)}")
        
#         # 分析每个字段的特征
#         single_value_fields = []
#         multi_value_fields = []
        
#         for col in group_data.columns:
#             if col == data_key:
#                 continue
                
#             unique_count = group_data[col].nunique()
#             total_count = len(group_data)
            
#             logger.debug(f"  ├─ 字段 '{col}': 唯一值数={unique_count}, 总行数={total_count}")
            
#             if unique_count == 1:
#                 # 所有行值相同 → 单值字段
#                 single_value_fields.append(col)
#             else:
#                 # 值不同 → 多值字段
#                 multi_value_fields.append(col)
        
#         # 处理单值字段：取第一行的值
#         for col in single_value_fields:
#             context[col] = group_data[col].iloc[0]
#             logger.debug(f"  ├─ 单值字段 '{col}': {context[col]}")
        
#         # 处理多值字段：构建记录列表
#         if multi_value_fields:
#             # 只包含多值字段的记录
#             multi_value_data = group_data[multi_value_fields]
#             context['records'] = multi_value_data.to_dict('records')
#             context['record_count'] = len(multi_value_data)
            
#             logger.info(f"  ├─ 多值字段: {multi_value_fields}")
#             logger.info(f"  ├─ 生成记录数: {context['record_count']}")
        
#         # 确保所有键都是字符串
#         context = {str(key): value for key, value in context.items()}
        
#         return context
    
#     def fill_table_data(self, file_path: str, data: pd.DataFrame,
#                        config: Dict[str, Any], data_key_value: Any = None) -> str:
#         """填充 Word 表格数据 - 待实现"""
#         try:
#             logger.info(f"📊 开始Word表格填充，文件: {file_path}, 数据量: {len(data)}")
            
#             # 参数映射（保持与Excel相同的参数名）
#             table_index = config.get('insert_col', 0)  # insert_col -> 表格索引
#             start_row = config.get('insert_row', 1)    # insert_row -> 表格内起始行
#             reserved_rows = config.get('reserved_rows', 0)
            
#             logger.info(f"📍 表格操作: 第{table_index}个表格, 起始行:{start_row}, 保留行:{reserved_rows}")
            
#             # 加载Word文档
#             doc = Document(file_path)
            
#             # 检查表格是否存在
#             if table_index >= len(doc.tables):
#                 raise ValueError(f"文档中只有{len(doc.tables)}个表格，无法访问第{table_index}个表格")
            
#             table = doc.tables[table_index]
#             logger.info(f"📋 目标表格: {len(table.rows)}行 x {len(table.columns)}列")
            
#             # 计算需要添加的行数
#             data_count = len(data)
#             needed_rows = data_count - reserved_rows
            
#             logger.info(f"📏 数据行数: {data_count}, 需要添加: {needed_rows}行")
            
#             # 动态添加行
#             if needed_rows > 0:
#                 for i in range(needed_rows):
#                     table.add_row()
#                 logger.info(f"✅ 成功添加{needed_rows}行")
            
#             # 填充数据
#             logger.info("🖊️ 开始填充表格数据...")
#             for row_offset, (_, row_data) in enumerate(data.iterrows()):
#                 row_idx = start_row + row_offset
                
#                 # 检查行索引是否有效
#                 if row_idx >= len(table.rows):
#                     logger.warning(f"⚠️ 行索引{row_idx}超出表格范围，跳过")
#                     continue
                    
#                 for col_offset, value in enumerate(row_data):
#                     # 检查列索引是否有效
#                     if col_offset >= len(table.rows[row_idx].cells):
#                         logger.warning(f"⚠️ 列索引{col_offset}超出表格范围，跳过")
#                         continue
                        
#                     cell = table.rows[row_idx].cells[col_offset]
#                     sanitized_value = TemplateUtils.sanitize_value(value)
#                     cell.text = str(sanitized_value)
                    
#                     logger.debug(f"  └─ 单元格({row_idx},{col_offset}) = {sanitized_value}")
            
#             # 保存文档
#             doc.save(file_path)
#             logger.info(f"✅ Word表格填充完成: {file_path}")
#             return file_path
            
#         except Exception as e:
#             logger.error(f"❌ Word表格填充失败: {e}")
#             raise