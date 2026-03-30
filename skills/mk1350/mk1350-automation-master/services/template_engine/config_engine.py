# template_engine/config_engine.py
import os
import re
import logging
import pandas as pd
from typing import Dict, Any, List
from docx import Document
import openpyxl

from .core import *

logger = logging.getLogger(__name__)

class TemplateConfigEngine:
    """模板配置引擎 - 负责智能分析和执行计划生成"""
    
    # 步骤函数名映射 - 避免直接函数引用
    _STEP_FUNCTIONS = {
        "placeholder_fill": "fill_placeholders",
        "table_fill": "fill_table_data"
    }
    
    @staticmethod
    def create_execution_plan(frontend_config: Dict[str, Any]) -> ExecutionPlan:
        """
        根据前端配置创建执行计划
        
        Args:
            frontend_config: 前端配置参数
                - mode: 生成模式 (placeholder_only/table_only/mixed)
                - input_template: 模板文件路径
                - input_data: 数据文件路径
                - data_key: 数据主键字段
                - output_dir: 输出目录
                - 其他表格配置参数...
        
        Returns:
            ExecutionPlan: 完整的执行计划
        """
        logger.info(f"🎯 开始创建执行计划，模式: {frontend_config.get('mode')}")
        
        # 解析生成模式
        mode = GenerationMode(frontend_config.get('mode', 'mixed'))
        
        # 构建数据源
        data_sources = TemplateConfigEngine._build_data_sources(frontend_config)
        
        # 构建执行步骤
        steps = TemplateConfigEngine._build_execution_steps(mode, frontend_config, data_sources)
        
        plan = ExecutionPlan(
            template_path=frontend_config['input_template'],
            output_dir=frontend_config['output_dir'],
            mode=mode,
            steps=steps,
            data_sources=data_sources
        )
        
        logger.info(f"✅ 执行计划创建完成，共 {len(steps)} 个步骤")
        return plan
    
    @staticmethod
    def _build_data_sources(config: Dict[str, Any]) -> List[DataSource]:
        """构建数据源配置"""
        data_sources = [
            DataSource(
                name="main_data",
                file_path=config['input_data'],
                sheet_name=config.get('input_data_sheet_name', 'Sheet1'),
                key_field=config.get('data_key')
            )
        ]
        return data_sources
    
    @staticmethod
    def _build_execution_steps(mode: GenerationMode, config: Dict[str, Any], 
                             data_sources: List[DataSource]) -> List[ExecutionStep]:
        """构建执行步骤"""
        steps = []
        
        # 占位符填充步骤 - 先执行
        if mode in [GenerationMode.PLACEHOLDER_ONLY, GenerationMode.MIXED]:
            placeholder_step = TemplateConfigEngine._create_placeholder_step(config)
            if placeholder_step:
                steps.append(placeholder_step)
        
        # 表格填充步骤 - 后执行
        if mode in [GenerationMode.TABLE_ONLY, GenerationMode.MIXED]:
            table_steps = TemplateConfigEngine._create_table_steps(config)
            steps.extend(table_steps)
        
        return steps
    
    @staticmethod
    def _create_placeholder_step(config: Dict[str, Any]) -> Optional[ExecutionStep]:
        """创建占位符填充步骤"""
        template_path = config['input_template']
        
        # 分析模板获取占位符
        placeholders = TemplateConfigEngine.analyze_template_placeholders(template_path)
        if not placeholders:
            logger.warning("⚠️ 未检测到占位符，跳过占位符填充步骤")
            return None
        
        # 自动创建占位符映射
        placeholder_mappings = TemplateConfigEngine._create_placeholder_mappings(placeholders, config)
        
        return ExecutionStep(
            step_id="fill_placeholders",
            action_type="placeholder_fill",
            function_name="fill_placeholders",  # 使用函数名
            description="填充模板占位符",
            config={"placeholders": placeholder_mappings},
            data_source="main_data"
        )
    
    @staticmethod
    def _create_table_steps(config: Dict[str, Any]) -> List[ExecutionStep]:
        """创建表格填充步骤"""
        steps = []
        
        # 构建表格配置
        table_config = {
            'insert_row': int(config.get('insert_row', 1)),
            'insert_col': int(config.get('insert_col', 1)),
            'reserved_rows': int(config.get('reserved_rows', 1)),
            'data_key': config.get('data_key')
        }
        
        logger.info(f"📋 创建表格步骤配置: {table_config}")
        
        steps.append(ExecutionStep(
            step_id="fill_table_data",
            action_type="table_fill",
            function_name="fill_table_data",  # 使用函数名
            description="填充表格数据",
            config=table_config,
            data_source="main_data"
        ))
        
        return steps
    
    @staticmethod
    def analyze_template_placeholders(template_path: str) -> List[str]:
        """
        分析模板中的占位符
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            List[str]: 检测到的占位符列表
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        file_ext = os.path.splitext(template_path)[1].lower()
        
        if file_ext in ['.docx', '.doc']:
            return TemplateConfigEngine._analyze_word_placeholders(template_path)
        elif file_ext in ['.xlsx', '.xls']:
            return TemplateConfigEngine._analyze_excel_placeholders(template_path)
        else:
            raise ValueError(f"不支持的模板格式: {file_ext}")
    
    @staticmethod
    def _analyze_word_placeholders(template_path: str) -> List[str]:
        """分析Word模板中的占位符"""
        placeholders = []
        try:
            doc = Document(template_path)
            placeholder_pattern = r'\{\{(.*?)\}\}'
            
            # 检查段落
            for paragraph in doc.paragraphs:
                matches = re.findall(placeholder_pattern, paragraph.text)
                placeholders.extend(matches)
            
            # 检查表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(placeholder_pattern, cell.text)
                        placeholders.extend(matches)
            
        except Exception as e:
            logger.error(f"分析Word模板占位符失败: {e}")
            raise
        
        return list(set(placeholders))
    
    @staticmethod
    def _analyze_excel_placeholders(template_path: str) -> List[str]:
        """分析Excel模板中的占位符"""
        placeholders = []
        try:
            wb = openpyxl.load_workbook(template_path)
            placeholder_pattern = r'\{\{(.*?)\}\}'
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            matches = re.findall(placeholder_pattern, cell.value)
                            placeholders.extend(matches)
            
        except Exception as e:
            logger.error(f"分析Excel模板占位符失败: {e}")
            raise
        
        return list(set(placeholders))
    
    @staticmethod
    def _create_placeholder_mappings(placeholders: List[str], config: Dict[str, Any]) -> List[PlaceholderMapping]:
        """自动创建占位符映射配置"""
        mappings = []
        
        for placeholder in placeholders:
            # 简单的自动映射规则：占位符名称直接映射到数据字段
            data_field = placeholder.strip()
            
            mappings.append(PlaceholderMapping(
                placeholder=f"{{{{{placeholder}}}}}",
                data_field=data_field,
                data_type="text"
            ))
        
        logger.info(f"📋 自动创建了 {len(mappings)} 个占位符映射")
        return mappings