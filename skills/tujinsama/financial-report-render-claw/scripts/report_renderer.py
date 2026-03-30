#!/usr/bin/env python3
"""报告渲染器 — financial-report-render-claw

支持 HTML / PDF / DOCX 输出格式，共用统一的 JSON 配置。
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

import jinja2


# ── helpers ──────────────────────────────────────────────────────────

def fmt_number(val):
    """智能格式化数字：千分位分隔。"""
    if isinstance(val, (int, float)):
        if abs(val) >= 10000:
            return f"{val:,.0f}"
        return f"{val:,.2f}" if val != int(val) else f"{val:,.0f}"
    return str(val)


def pct_class(positive):
    return "positive" if positive else "negative"


def risk_class(level):
    return {"高": "risk-high", "中": "risk-medium", "低": "risk-low"}.get(level, "")


def status_class(status):
    return {"达标": "status-ok", "未达标": "status-bad", "关注": "status-warn"}.get(status, "")


def load_json(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ── HTML ─────────────────────────────────────────────────────────────

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ config.title }}</title>
<style>
  :root {
    --primary: #1B4F72;
    --secondary: #2E86C1;
    --accent: #27AE60;
    --warning: #F39C12;
    --danger: #E74C3C;
    --neutral: #95A5A6;
    --bg: #FAFBFC;
    --text: #2C3E50;
    --text-light: #7F8C8D;
    --border: #E8ECF0;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    font-family: "PingFang SC", "Microsoft YaHei", "Noto Sans CJK SC", sans-serif;
    color: var(--text);
    background: var(--bg);
    line-height: 1.8;
    font-size: 15px;
  }

  /* 封面 */
  .cover {
    min-height: 100vh;
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
    color: white;
    text-align: center;
    padding: 40px 20px;
    page-break-after: always;
  }
  .cover h1 { font-size: 2.4em; margin-bottom: 12px; font-weight: 700; }
  .cover .subtitle { font-size: 1.1em; opacity: 0.85; margin-bottom: 40px; }
  .cover .meta { font-size: 0.95em; opacity: 0.7; line-height: 2; }

  /* 容器 */
  .container { max-width: 900px; margin: 0 auto; padding: 40px 30px; }

  /* 章节标题 */
  .section { margin-bottom: 48px; page-break-inside: avoid; }
  .section-title {
    font-size: 1.35em;
    font-weight: 700;
    color: var(--primary);
    border-bottom: 3px solid var(--secondary);
    padding-bottom: 8px;
    margin-bottom: 20px;
  }

  /* 高亮卡片 */
  .highlights { display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 16px; margin-top: 16px; }
  .highlight-card {
    background: white;
    border-radius: 10px;
    padding: 18px 20px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
    border-left: 4px solid var(--secondary);
  }
  .highlight-card .label { font-size: 0.85em; color: var(--text-light); }
  .highlight-card .value { font-size: 1.5em; font-weight: 700; margin: 4px 0; }
  .highlight-card .change { font-size: 0.85em; font-weight: 600; }
  .change.positive { color: var(--accent); }
  .change.negative { color: var(--danger); }

  /* KPI 卡片 */
  .kpi-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 16px; margin-top: 16px; }
  .kpi-card {
    background: white;
    border-radius: 10px;
    padding: 18px 20px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
    border-top: 4px solid var(--secondary);
  }
  .kpi-card .label { font-size: 0.85em; color: var(--text-light); }
  .kpi-card .value { font-size: 1.5em; font-weight: 700; margin: 4px 0; }
  .kpi-card .target { font-size: 0.8em; color: var(--text-light); }
  .kpi-card .status { display: inline-block; padding: 2px 10px; border-radius: 12px; font-size: 0.75em; font-weight: 600; margin-top: 6px; }
  .status-ok { background: #E8F8F0; color: var(--accent); }
  .status-bad { background: #FDEDEC; color: var(--danger); }
  .status-warn { background: #FEF5E7; color: var(--warning); }

  /* 图表 */
  .chart-container { text-align: center; margin: 16px 0; }
  .chart-container img { max-width: 100%; height: auto; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }
  .chart-caption { font-size: 0.85em; color: var(--text-light); margin-top: 8px; }
  .chart-analysis { margin-top: 12px; padding: 16px; background: #F0F7FC; border-radius: 8px; border-left: 3px solid var(--secondary); }

  /* 表格 */
  .data-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.9em; }
  .data-table th { background: var(--primary); color: white; padding: 10px 14px; text-align: left; font-weight: 600; }
  .data-table td { padding: 9px 14px; border-bottom: 1px solid var(--border); }
  .data-table tr:last-child td { border-bottom: 2px solid var(--primary); font-weight: 600; }
  .data-table tr:nth-child(even) td { background: #F8FAFB; }

  /* 风险建议 */
  .risk-item { display: flex; align-items: flex-start; gap: 10px; margin-bottom: 10px; padding: 10px 14px; border-radius: 8px; }
  .risk-high { background: #FDEDEC; border-left: 4px solid var(--danger); }
  .risk-medium { background: #FEF5E7; border-left: 4px solid var(--warning); }
  .risk-low { background: #FEF9E7; border-left: 4px solid #F7DC6F; }
  .risk-badge { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 0.75em; font-weight: 700; color: white; flex-shrink: 0; }
  .risk-high .risk-badge { background: var(--danger); }
  .risk-medium .risk-badge { background: var(--warning); }
  .risk-low .risk-badge { background: #F7DC6F; color: var(--text); }
  .rec-list { list-style: none; padding: 0; }
  .rec-list li { padding: 8px 0 8px 24px; position: relative; border-bottom: 1px dashed var(--border); }
  .rec-list li::before { content: "→"; position: absolute; left: 0; color: var(--secondary); font-weight: 700; }

  /* 文本 */
  .text-content p { margin-bottom: 12px; }
  .text-content h2 { font-size: 1.15em; color: var(--primary); margin: 16px 0 8px; }
  .text-content ul { padding-left: 24px; }
  .text-content li { margin-bottom: 6px; }

  /* 页脚 */
  .page-footer { text-align: center; font-size: 0.75em; color: var(--neutral); padding: 20px 0; margin-top: 40px; border-top: 1px solid var(--border); }

  /* 打印优化 */
  @media print {
    body { background: white; }
    .cover { min-height: auto; padding: 60px 20px; }
    .section { page-break-inside: avoid; }
    .highlight-card, .kpi-card { box-shadow: none; border: 1px solid var(--border); }
    .chart-container img { box-shadow: none; }
  }
</style>
</head>
<body>

{% if config.company and config.company.logo %}
{% endif %}

<!-- 封面 -->
<div class="cover">
  <h1>{{ config.title }}</h1>
  {% if config.subtitle %}
  <div class="subtitle">{{ config.subtitle }}</div>
  {% endif %}
  <div class="meta">
    {% if config.company %}{{ config.company.name }}<br>{% endif %}
    {{ config.metadata.period }}<br>
    {% if config.metadata.author %}{{ config.metadata.author }}<br>{% endif %}
    {{ config.metadata.date }}
  </div>
</div>

<!-- 正文 -->
<div class="container">

{% for section in config.sections %}
<div class="section">
  <h2 class="section-title">{{ section.title }}</h2>

  {% if section.type == 'summary' %}
  <div class="text-content"><p>{{ section.content }}</p></div>
  {% if section.highlights %}
  <div class="highlights">
  {% for h in section.highlights %}
    <div class="highlight-card">
      <div class="label">{{ h.label }}</div>
      <div class="value">{{ h.value }}</div>
      {% if h.change %}<div class="change {{ 'positive' if h.positive else 'negative' }}">{{ h.change }}</div>{% endif %}
    </div>
  {% endfor %}
  </div>
  {% endif %}

  {% elif section.type == 'kpi-cards' %}
  <div class="kpi-grid">
  {% for item in section.items %}
    <div class="kpi-card">
      <div class="label">{{ item.label }}</div>
      <div class="value">{{ item.value }}</div>
      <div class="target">目标: {{ item.target }}</div>
      <span class="status {{ item.status | status_class }}">{{ item.status }}</span>
    </div>
  {% endfor %}
  </div>

  {% elif section.type == 'chart' %}
  {% if section.image %}
  <div class="chart-container">
    <img src="{{ section.image }}" alt="{{ section.title }}">
    {% if section.caption %}<div class="chart-caption">{{ section.caption }}</div>{% endif %}
  </div>
  {% endif %}
  {% if section.analysis %}
  <div class="chart-analysis">{{ section.analysis }}</div>
  {% endif %}

  {% elif section.type == 'table' %}
  <table class="data-table">
    <thead><tr>{% for col in section.columns %}<th>{{ col }}</th>{% endfor %}</tr></thead>
    <tbody>
    {% for row in section.rows %}
      <tr>{% for cell in row %}<td>{{ cell }}</td>{% endfor %}</tr>
    {% endfor %}
    </tbody>
  </table>
  {% if section.footnote %}<div style="font-size:0.85em;color:var(--text-light);margin-top:4px;">{{ section.footnote }}</div>{% endif %}

  {% elif section.type == 'text' %}
  <div class="text-content">{{ section.content | safe }}</div>

  {% elif section.type == 'advisory' %}
  {% if section.risks %}
  <h3 style="margin-bottom:12px;">风险提示</h3>
  {% for r in section.risks %}
  <div class="risk-item {{ r.level | risk_class }}">
    <span class="risk-badge">{{ r.level }}</span>
    <span>{{ r.description }}</span>
  </div>
  {% endfor %}
  {% endif %}
  {% if section.recommendations %}
  <h3 style="margin:20px 0 12px;">行动建议</h3>
  <ul class="rec-list">
  {% for rec in section.recommendations %}
    <li>{{ rec }}</li>
  {% endfor %}
  </ul>
  {% endif %}

  {% else %}
  <div class="text-content"><p>{{ section.content | default('') }}</p></div>
  {% endif %}

</div>
{% endfor %}

{% if config.appendix %}
{% for app in config.appendix %}
<div class="section">
  <h2 class="section-title">{{ app.title }}</h2>
  <div class="text-content">{{ app.content | default('') }}</div>
</div>
{% endfor %}
{% endif %}

<div class="page-footer">
  {{ config.title }} · {{ config.metadata.date }}
  {% if config.company %}· {{ config.company.name }}{% endif %}
</div>

</div>
</body>
</html>"""


def render_html(config, output, theme="professional"):
    """渲染 HTML 报告。"""
    env = jinja2.Environment(
        loader=jinja2.DictLoader({"report.html": HTML_TEMPLATE}),
        autoescape=False,
    )
    env.filters["status_class"] = status_class
    env.filters["risk_class"] = risk_class
    env.filters["pct_class"] = pct_class

    template = env.get_template("report.html")
    html = template.render(config=config)

    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(html, encoding="utf-8")
    print(f"[OK] HTML 报告已生成: {output_path}")


def render_pdf(config, output, theme="professional"):
    """通过 weasyprint 渲染 PDF。"""
    try:
        from weasyprint import HTML
    except ImportError:
        print("[ERROR] 需要 weasyprint: pip install weasyprint")
        print("[INFO] 回退方案：先生成 HTML，在浏览器中 Ctrl+P 打印为 PDF")
        html_out = Path(output).with_suffix(".html")
        render_html(config, str(html_out), theme)
        sys.exit(1)

    env = jinja2.Environment(
        loader=jinja2.DictLoader({"report.html": HTML_TEMPLATE}),
        autoescape=False,
    )
    env.filters["status_class"] = status_class
    env.filters["risk_class"] = risk_class
    env.filters["pct_class"] = pct_class

    template = env.get_template("report.html")
    html = template.render(config=config)

    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html).write_pdf(str(output_path))
    print(f"[OK] PDF 报告已生成: {output_path}")


def render_docx(config, output):
    """通过 python-docx 渲染 Word 报告。"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("[ERROR] 需要 python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # 页眉
    section = doc.sections[0]
    if config.get("company", {}).get("name"):
        header = section.header
        header.paragraphs[0].text = config["company"]["name"]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(config["title"])
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    if config.get("subtitle"):
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(config["subtitle"])
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_lines = [config["metadata"]["period"]]
    if config["metadata"].get("author"):
        meta_lines.append(config["metadata"]["author"])
    meta_lines.append(config["metadata"]["date"])
    meta_p.add_run("\n".join(meta_lines)).font.size = Pt(12)

    doc.add_page_break()

    # 正文
    for sec in config.get("sections", []):
        h = doc.add_heading(sec["title"], level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        content = sec.get("content", "")
        if content:
            doc.add_paragraph(content)

        if sec.get("type") == "kpi-cards" and sec.get("items"):
            for item in sec["items"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{item['label']}: ")
                run.bold = True
                p.add_run(f"{item['value']}  (目标: {item.get('target', '-')})")

        if sec.get("type") == "table" and sec.get("columns"):
            table = doc.add_table(rows=1, cols=len(sec["columns"]))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, col in enumerate(sec["columns"]):
                hdr[i].text = col
                for p in hdr[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
            for row_data in sec.get("rows", []):
                row = table.add_row().cells
                for i, cell in enumerate(row_data):
                    row[i].text = str(cell)

        if sec.get("type") == "chart" and sec.get("image"):
            img_path = Path(sec["image"])
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.5))
                if sec.get("caption"):
                    cap_p = doc.add_paragraph(sec["caption"])
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if sec.get("type") == "advisory":
            if sec.get("risks"):
                doc.add_heading("风险提示", level=2)
                for r in sec["risks"]:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{r['level']}] ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C) if r["level"] == "高" else RGBColor(0xF3, 0x9C, 0x12)
                    p.add_run(r["description"])
            if sec.get("recommendations"):
                doc.add_heading("行动建议", level=2)
                for rec in sec["recommendations"]:
                    doc.add_paragraph(rec, style="List Bullet")

    # 附录
    for app in config.get("appendix", []):
        doc.add_heading(app["title"], level=1)
        if app.get("content"):
            doc.add_paragraph(app["content"])

    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] Word 报告已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="报告渲染器")
    parser.add_argument("format", choices=["html", "pdf", "docx"], help="输出格式")
    parser.add_argument("config", help="报告配置 JSON 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出文件路径")
    parser.add_argument("--theme", default="professional", help="主题名称")

    args = parser.parse_args()
    config = load_json(args.config)

    renderers = {"html": render_html, "pdf": render_pdf, "docx": render_docx}
    renderers[args.format](config, args.output, args.theme)


if __name__ == "__main__":
    main()
