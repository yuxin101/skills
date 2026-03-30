"""
模块设计文档生成脚本模板
使用方式：复制到输出目录，按实际模块内容修改后运行
依赖：pip install python-docx
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

# ============================================================
# 配置区域 - 每次使用时修改
# ============================================================
BASE = Path(r"C:\Users\liuchao25\.openclaw\workspace\outputs\<模块>-design-doc-<日期>")
SHOTS = BASE / "screenshots"
OUT = BASE / "<系统名>-<模块名>-设计文档-YYYY-MM-DD.docx"

SYSTEM_NAME = "园区智慧应用管理系统"
MODULE_NAME = "<模块名>"
GEN_TIME = "YYYY-MM-DD HH:MM:SS"
# ============================================================


def set_font(run, size=None, bold=None, font_name="仿宋"):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def para(doc, text="", size=12, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("• " + text)
    set_font(r, size=12)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_font(r, size=sizes.get(level, 12), bold=True)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录（打开 Word 后如未自动刷新，请右键目录选择"更新域"）"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r._element.append(fldChar)
    r._element.append(instrText)
    r._element.append(fldChar2)
    r._element.append(text)
    r._element.append(fldChar3)


def image(doc, name, caption, width_pt=420):
    path = SHOTS / name
    if path.exists():
        doc.add_picture(str(path), width=Pt(width_pt))
        para(doc, caption, size=10, center=True)
    else:
        para(doc, f"【缺图】{name}", size=10)


def table7(doc, title, headers, rows):
    """生成7列表结构表格"""
    heading(doc, title, level=3)
    tbl = doc.add_table(rows=1 + len(rows), cols=7)
    tbl.style = "Table Grid"
    # 表头
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=11, bold=True)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10.5)
    doc.add_paragraph()  # 表后空行


TABLE_HEADERS = ["字段名", "说明", "类型", "是否必填", "默认值", "约束", "备注"]


def build():
    doc = Document()

    # 全局字体
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(12)

    # 封面
    para(doc, SYSTEM_NAME, size=18, bold=True, center=True)
    para(doc, f"{MODULE_NAME} 模块设计文档", size=20, bold=True, center=True)
    para(doc, f"生成时间：{GEN_TIME}", size=11, center=True)
    para(doc, "本文档根据前后端代码及实测页面截图自动整理，包含功能描述、流程图、页面截图、表结构及实现类。", size=11, center=True)
    doc.add_page_break()

    # 目录
    heading(doc, "目录", 1)
    add_toc(doc)
    doc.add_page_break()

    # 1. 模块简介
    heading(doc, "1. 模块简介", 1)
    para(doc, f"{MODULE_NAME} 模块位于企业端，包含以下子功能：<列出子模块>。")
    para(doc, "前端路由前缀：/<路由前缀>/")
    para(doc, "后端接口前缀：/<接口前缀>/")
    para(doc, "前端代码位置：<前端路径>")
    para(doc, "后端代码位置：<后端路径>")
    image(doc, "00-login-page.png", "图1-1 登录页")

    # 2. 功能模块详细设计
    heading(doc, "2. 功能模块详细设计", 1)

    # --- 子模块示例，按实际模块复制此块 ---
    heading(doc, "2.1 <子模块名>", 2)

    heading(doc, "2.1.1 功能描述与业务流程", 3)
    para(doc, "<子模块名>用于<功能说明>，支持<操作列表>。")
    para(doc, "业务流程：")
    bullet(doc, "用户进入列表页 -> 通过<搜索条件>筛选查询")
    bullet(doc, "用户点击"新增" -> 填写<主表字段> -> 填写<子表内容> -> 点击保存 -> 系统校验必填项 -> 保存成功返回列表")
    bullet(doc, "用户点击"提交审核" -> 系统弹出确认框 -> 确认后触发审批流 -> 记录 process_instance_id")
    bullet(doc, "审批人审批通过 -> 状态变更为"已通过" / 驳回 -> 状态变更为"已驳回"")
    bullet(doc, "用户可在列表查看审批状态 -> 点击"查看审批"跳转审批详情")

    heading(doc, "2.1.2 页面截图", 3)
    image(doc, "<模块>-list.png", "图2-1 <子模块>列表页")
    image(doc, "<模块>-form.png", "图2-2 <子模块>新增/编辑表单")
    image(doc, "<模块>-detail.png", "图2-3 <子模块>详情页")

    heading(doc, "2.1.3 数据表结构", 3)
    # 主表
    table7(doc, "主表：<table_name>", TABLE_HEADERS, [
        ("id",          "主键ID",     "bigint",        "是", "AUTO_INCREMENT", "PRIMARY KEY", "主键"),
        ("field_name",  "字段说明",   "varchar(200)",  "是", "NULL",           "",            ""),
        ("status",      "状态",       "tinyint",       "是", "0",              "INDEX",       "0草稿 1审核中 2已通过 3已驳回"),
        ("process_instance_id", "流程实例ID", "varchar(64)", "否", "NULL",    "",            "审批流关联"),
        ("create_time", "创建时间",   "datetime",      "否", "NULL",           "",            ""),
        ("remark",      "备注",       "varchar(500)",  "否", "NULL",           "",            ""),
    ])
    # 子表（如有）
    table7(doc, "子表：<sub_table_name>", TABLE_HEADERS, [
        ("id",          "主键ID",     "bigint",        "是", "AUTO_INCREMENT", "PRIMARY KEY", "主键"),
        ("parent_id",   "主表ID",     "bigint",        "是", "NULL",           "INDEX",       "关联主表id"),
        ("field_name",  "字段说明",   "varchar(200)",  "否", "NULL",           "",            ""),
    ])

    heading(doc, "2.1.4 实现类", 3)
    bullet(doc, "前端列表页：src/views/<模块>/index.vue")
    bullet(doc, "前端表单：src/views/<模块>/<模块>Form.vue")
    bullet(doc, "前端详情：src/views/<模块>/detail.vue")
    bullet(doc, "前端API：src/api/<模块>.ts")
    bullet(doc, "后端Controller：<模块>Controller.java")
    bullet(doc, "后端Service：<模块>Service.java / <模块>ServiceImpl.java")
    bullet(doc, "审批监听：<模块>BpmStatusListener.java")
    bullet(doc, "数据对象：<模块>DO.java")
    # --- 子模块示例结束 ---

    # 3. 总结说明
    heading(doc, "3. 总结说明", 1)
    bullet(doc, "本文档各功能模块均已整理：功能描述/流程、页面截图、表结构、实现类四部分。")
    bullet(doc, "表结构已结合前后端代码及建表SQL整理，主子表完整。")
    bullet(doc, "文档采用 Word 标题样式，已插入自动目录。")

    doc.save(str(OUT))
    print(f"设计文档已生成：{OUT}")


if __name__ == "__main__":
    build()
