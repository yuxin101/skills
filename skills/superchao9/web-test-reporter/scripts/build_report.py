"""
Web 功能测试报告生成脚本模板
使用方式：复制此脚本到报告目录，按实际测试内容修改后运行
依赖：pip install python-docx
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

# ============================================================
# 配置区域 - 每次使用时修改这里
# ============================================================
BASE = Path(r"C:\Users\liuchao25\.openclaw\workspace\reports\<模块>-test-<日期>")
SHOTS = BASE / "screenshots"
OUT = BASE / "<端名>-<模块名>功能测试报告-YYYY-MM-DD.docx"

TITLE = "企业端-<模块名>功能测试报告"
TEST_DATE = "YYYY-MM-DD"
TEST_URL = "http://localhost/"
TEST_SCOPE = "企业端 > <模块名>"
# ============================================================


def set_run_font(run, size=None, bold=None):
    run.font.name = "Microsoft YaHei"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def para(doc, text="", size=10.5, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    sizes = {1: 16, 2: 13.5, 3: 12}
    set_run_font(r, size=sizes.get(level, 11), bold=True)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录（打开 Word 后如未自动刷新，请右键目录选择"更新域"）"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r._element.append(fldChar)
    r._element.append(instrText)
    r._element.append(fldChar2)
    r._element.append(text)
    r._element.append(fldChar3)


def image(doc, name, caption, width=6.0):
    path = SHOTS / name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        para(doc, caption, size=9.5, center=True)
    else:
        para(doc, f"【缺图】{name}", size=10)


def build():
    doc = Document()

    # 全局字体
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    # 封面
    para(doc, TITLE, size=18, bold=True, center=True)
    para(doc, f"测试日期：{TEST_DATE}")
    para(doc, f"测试对象：{TEST_URL}")
    para(doc, f"测试范围：{TEST_SCOPE}")
    para(doc, "执行规范：先新增测试数据，再删除自己新增的数据；确认类操作先验证取消分支，再验证确认分支；关键步骤截图留证。")

    # 目录
    heading(doc, "目录", 1)
    add_toc(doc)

    # 1. 测试概述
    heading(doc, "1. 测试概述", 1)
    para(doc, "本次对<模块名>全部功能页面进行测试，执行完备闭环验证。")
    para(doc, "可闭环模块：<列出模块>")
    para(doc, "受限模块：<列出模块及原因>")
    para(doc, "存在异常模块：<列出模块>")

    # 2. 测试环境
    heading(doc, "2. 测试环境", 1)
    para(doc, f"系统地址：{TEST_URL}")
    para(doc, "账号：租户=园区智慧应用管理系统，用户=admin，密码=admin123")
    para(doc, "测试工具：agent-browser CLI")
    para(doc, f"截图目录：{SHOTS}")

    # 3. 总体结论
    heading(doc, "3. 总体结论", 1)
    para(doc, "3.1 已完成闭环且验证通过的模块：<列出>")
    para(doc, "3.2 受页面能力限制无法闭环的模块：<列出>")
    para(doc, "3.3 新增链路存在异常的模块：<列出>")
    para(doc, "3.4 明确识别的异常：<列出>")

    # 4. 登录过程
    heading(doc, "4. 登录过程", 1)
    image(doc, "00-login-page.png", "图1 登录页初始状态")
    image(doc, "01-login-success.png", "图2 登录成功后首页")

    # 5. 分模块测试记录
    heading(doc, "5. 分模块测试记录", 1)

    # --- 模块示例，按实际模块复制此块 ---
    heading(doc, "5.1 <模块名>", 2)
    para(doc, "测试范围：页面访问、新增、编辑、删除、审核、导出、边界值。")

    heading(doc, "5.1.1 基础验证", 3)
    image(doc, "<模块>-init.png", "图3 <模块>页面初始状态")
    image(doc, "<模块>-search.png", "图4 搜索功能验证")
    para(doc, "结果：<描述>")

    heading(doc, "5.1.2 新增闭环", 3)
    image(doc, "<模块>-add-open.png", "图5 新增弹窗打开")
    image(doc, "<模块>-add-validation.png", "图6 必填项校验")
    image(doc, "<模块>-add-filling.png", "图7 填写中")
    image(doc, "<模块>-add-success.png", "图8 新增成功")
    image(doc, "<模块>-add-list-verify.png", "图9 列表验证新记录")
    para(doc, "结论：<描述>")

    heading(doc, "5.1.3 编辑闭环", 3)
    image(doc, "<模块>-edit-open.png", "图10 编辑弹窗（原数据回填）")
    image(doc, "<模块>-edit-success.png", "图11 编辑保存成功")
    image(doc, "<模块>-edit-list-verify.png", "图12 列表验证修改生效")
    para(doc, "结论：<描述>")

    heading(doc, "5.1.4 删除闭环", 3)
    image(doc, "<模块>-delete-cancel.png", "图13 取消删除后记录仍存在")
    image(doc, "<模块>-delete-success.png", "图14 确认删除后记录消失")
    para(doc, "结论：<描述>")

    heading(doc, "5.1.5 审核双路径", 3)
    image(doc, "<模块>-submit-cancel.png", "图15 取消提交审核状态未变")
    image(doc, "<模块>-submit-success.png", "图16 确认提交审核状态变更")
    image(doc, "<模块>-approval-view.png", "图17 查看审批页")
    para(doc, "结论：<描述>")

    heading(doc, "5.1.6 边界值测试", 3)
    image(doc, "<模块>-boundary-empty.png", "图18 必填项为空报错")
    image(doc, "<模块>-boundary-long.png", "图19 文本超长处理")
    image(doc, "<模块>-boundary-special.png", "图20 特殊字符处理")
    para(doc, "结论：<描述>")
    # --- 模块示例结束 ---

    # 6. 缺陷汇总
    heading(doc, "6. 缺陷汇总", 1)
    heading(doc, "6.1 <缺陷标题>", 2)
    para(doc, "现象：<描述>")
    para(doc, "复现步骤：<步骤>")
    para(doc, "严重程度：高/中/低")

    # 7. 最终结论
    heading(doc, "7. 最终结论", 1)
    para(doc, "7.1 <模块>：<结论>")
    para(doc, "7.2 遗留问题：<描述>")

    doc.save(str(OUT))
    print(f"报告已生成：{OUT}")


if __name__ == "__main__":
    build()
