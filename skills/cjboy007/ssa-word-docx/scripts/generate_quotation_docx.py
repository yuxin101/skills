#!/usr/bin/env python3
"""生成 Word 报价单 - 支持模板填充、样式、表格

🔴 P0-REVISE: 集成数据验证（防止示例数据）
"""

import sys
import json
import argparse
from pathlib import Path
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 🔴 P0: 导入验证模块
sys.path.insert(0, str(Path(__file__).parent.parent / 'quotation-workflow' / 'scripts'))
try:
    from quotation_schema import validate_quotation_data
    VALIDATION_AVAILABLE = True
except ImportError:
    VALIDATION_AVAILABLE = False
    print("⚠️  警告：quotation_schema 模块不可用，将跳过数据验证")

# 颜色定义
COLORS = {
    'header': RGBColor(0x44, 0x72, 0xC4),      # 蓝色
    'title': RGBColor(0x1F, 0x4E, 0x78),       # 深蓝
    'accent': RGBColor(0xFF, 0xC0, 0x00),      # 黄色
    'success': RGBColor(0x70, 0xAD, 0x47),     # 绿色
    'light_gray': RGBColor(0xF2, 0xF2, 0xF2),  # 浅灰
}

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 创建边框元素
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = parse_xml(r'<w:{} {}/>'.format(edge, nsdecls('w')))
                tcBorders.append(element)
            
            for key in edge_data:
                element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_header_paragraph(doc, text, level=1):
    """添加标题段落"""
    if level == 1:
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLORS['title']
    elif level == 2:
        p = doc.add_heading(text, level=2)
        run = p.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLORS['header']
    return p

def create_quotation_docx(output_path, data=None):
    """创建 Word 报价单
    
    Args:
        output_path: 输出文件路径
        data: 报价数据字典
    """
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ============ 标题 ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("报价单\nQUOTATION")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLORS['title']
    
    # ============ 公司信息 ============
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run("Farreach Electronic Co., Ltd.")
    run.font.size = Pt(12)
    run.font.bold = True
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("📧 sale@farreach-electronic.com | 🌐 www.farreach-electronic.com | 📱 +86-756-XXXXXXX")
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()  # 空行
    
    # ============ 客户信息和报价单号 ============
    # 使用表格布局
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    
    # 客户信息标题
    customer_header = info_table.rows[0].cells[0]
    customer_header.merge(info_table.rows[0].cells[1])
    customer_header.text = "客户信息 Customer"
    customer_header.paragraphs[0].runs[0].font.bold = True
    customer_header.paragraphs[0].runs[0].font.size = Pt(11)
    customer_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    customer_header.paragraphs[0].space_after = Pt(6)
    
    # 报价单信息标题
    quotation_header = info_table.rows[0].cells[2]
    quotation_header.merge(info_table.rows[0].cells[3])
    quotation_header.text = "报价单信息 Quotation"
    quotation_header.paragraphs[0].runs[0].font.bold = True
    quotation_header.paragraphs[0].runs[0].font.size = Pt(11)
    quotation_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    quotation_header.paragraphs[0].space_after = Pt(6)
    
    # 客户数据
    customer_data = data.get('customer', {}) if data else {}
    quotation_data = data.get('quotation', {}) if data else {}
    
    # 第 2 行
    info_table.rows[1].cells[0].text = "公司名称 Company:"
    info_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    info_table.rows[1].cells[1].text = customer_data.get('company_name', customer_data.get('name', '_________________'))
    info_table.rows[1].cells[2].text = "报价单号 Quotation No:"
    info_table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
    info_table.rows[1].cells[3].text = quotation_data.get('quotation_no', 'QT-' + datetime.now().strftime('%Y%m%d-001'))
    
    # 第 3 行
    info_table.rows[2].cells[0].text = "联系人 Contact:"
    info_table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    info_table.rows[2].cells[1].text = customer_data.get('contact', '_________________')
    info_table.rows[2].cells[2].text = "报价日期 Date:"
    info_table.rows[2].cells[2].paragraphs[0].runs[0].font.bold = True
    info_table.rows[2].cells[3].text = quotation_data.get('date', datetime.now().strftime('%Y-%m-%d'))
    
    # 第 4 行
    info_table.rows[3].cells[0].text = "邮箱 Email:"
    info_table.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True
    info_table.rows[3].cells[1].text = customer_data.get('email', '_________________')
    info_table.rows[3].cells[2].text = "有效期至 Valid Until:"
    info_table.rows[3].cells[2].paragraphs[0].runs[0].font.bold = True
    info_table.rows[3].cells[3].text = quotation_data.get('valid_until', '')
    
    doc.add_paragraph()  # 空行
    
    # ============ 产品表格 ============
    products = data.get('products', []) if data else []
    
    # 表头
    product_table = doc.add_table(rows=1, cols=6)
    product_table.style = 'Table Grid'
    product_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = [
        "序号\nNo.",
        "产品描述\nProduct Description",
        "规格型号\nSpecification",
        "数量\nQuantity",
        "单价\nUnit Price (USD)",
        "金额\nAmount (USD)"
    ]
    
    header_row = product_table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 产品数据行
    for idx, product in enumerate(products, start=1):
        row = product_table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = product.get('description', '')
        row.cells[2].text = product.get('specification', '')
        row.cells[3].text = str(product.get('quantity', 0))
        row.cells[4].text = f"${product.get('unit_price', product.get('unitPrice', 0)):.2f}"
        
        # 金额计算
        amount = product.get('quantity', 0) * product.get('unit_price', product.get('unitPrice', 0))
        row.cells[5].text = f"${amount:.2f}"
        
        # 设置单元格格式
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()  # 空行
    
    # ============ 汇总区域 ============
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # 币别
    summary_table.rows[0].cells[0].text = "币别 Currency:"
    summary_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[0].cells[1].text = data.get('currency', 'USD') if data else 'USD'
    
    # 付款条款
    summary_table.rows[1].cells[0].text = "付款条款 Payment Terms:"
    summary_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[1].cells[1].text = data.get('payment_terms', 'T/T 30% deposit, 70% before shipment') if data else 'T/T'
    
    # 交货期
    summary_table.rows[2].cells[0].text = "交货期 Lead Time:"
    summary_table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[2].cells[1].text = data.get('lead_time', '15-20 days after deposit') if data else '15-20 days'
    
    # 计算总额
    subtotal = sum(p.get('quantity', 0) * p.get('unit_price', p.get('unitPrice', 0)) for p in products) if products else 0
    freight = data.get('freight', 0) if data else 0
    tax = data.get('tax', 0) if data else 0
    total = subtotal + freight + tax
    
    # 小计
    summary_table.rows[3].cells[0].text = "小计 Subtotal:"
    summary_table.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[3].cells[1].text = f"${subtotal:.2f}"
    
    # 总计
    summary_table.rows[4].cells[0].text = "总计 Total:"
    summary_table.rows[4].cells[0].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[4].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    summary_table.rows[4].cells[1].text = f"${total:.2f}"
    summary_table.rows[4].cells[1].paragraphs[0].runs[0].font.bold = True
    summary_table.rows[4].cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # 空行
    
    # ============ 备注和条款 ============
    notes_para = doc.add_paragraph()
    notes_run = notes_para.add_run("备注 Remarks:\n")
    notes_run.font.bold = True
    notes_para.add_run(data.get('notes', '1. 以上价格基于当前原材料成本，如有变动将另行通知。\n2. 最终价格以确认为准。') if data else '')
    
    terms_para = doc.add_paragraph()
    terms_run = terms_para.add_run("条款 Terms & Conditions:\n")
    terms_run.font.bold = True
    terms_para.add_run("1. 报价有效期：30 天\n")
    terms_para.add_run("2. 付款方式：T/T 或 L/C\n")
    terms_para.add_run("3. 包装：标准出口包装\n")
    terms_para.add_run("4. 运输：FOB Shenzhen 或 CIF")
    
    doc.add_paragraph()  # 空行
    
    # ============ 签名区域 ============
    signature_para = doc.add_paragraph()
    sig_run = signature_para.add_run("授权签名 Authorized Signature:\n\n")
    sig_run.font.bold = True
    signature_para.add_run("_________________________\n")
    signature_para.add_run("Sales Manager")
    signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 保存文件
    doc.save(output_path)
    
    return output_path

def main():
    parser = argparse.ArgumentParser(
        description='生成 Word 报价单',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 创建空白模板
  python3 generate_quotation_docx.py --template output/quotation_template.docx
  
  # 从 JSON 数据生成报价单
  python3 generate_quotation_docx.py --data quotation_data.json --output output/QT-20260314-001.docx
  
  # 从命令行数据生成（快速测试）
  python3 generate_quotation_docx.py --output test.docx --quick-test
        '''
    )
    
    parser.add_argument('--template', '-t', help='创建空白模板到指定路径')
    parser.add_argument('--data', '-d', help='报价数据 JSON 文件路径')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--quick-test', action='store_true', help='使用测试数据快速生成')
    
    args = parser.parse_args()
    
    # 创建空白模板
    if args.template:
        output = create_quotation_docx(args.template)
        print(f"✅ 模板已创建：{output}")
        return
    
    # 从数据生成
    if args.output:
        data = None
        
        if args.data:
            with open(args.data, 'r', encoding='utf-8') as f:
                data = json.load(f)
        elif args.quick_test:
            data = {
                'customer': {
                    'company_name': 'Test Customer Inc.',
                    'contact': 'John Doe',
                    'email': 'john@test.com'
                },
                'quotation': {
                    'quotation_no': 'QT-20260314-001',
                    'date': '2026-03-14',
                    'valid_until': '2026-04-13'
                },
                'products': [
                    {
                        'description': 'HDMI 2.1 Ultra High Speed Cable',
                        'specification': '8K@60Hz, 48Gbps, 2m',
                        'quantity': 500,
                        'unit_price': 8.50
                    },
                    {
                        'description': 'USB-C to USB-C Cable',
                        'specification': 'USB 4.0, 80Gbps, 100W PD, 1m',
                        'quantity': 1000,
                        'unit_price': 12.00
                    }
                ],
                'currency': 'USD',
                'payment_terms': 'T/T 30% deposit, 70% before shipment',
                'lead_time': '15-20 days after deposit',
                'freight': 150.00,
                'tax': 0,
                'notes': '1. 以上价格基于当前原材料成本\n2. 最终价格以确认为准'
            }
        
        if not args.data and not args.quick_test:
            print("❌ 请提供 --data 或 --quick-test", file=sys.stderr)
            sys.exit(1)
        
        # 🔴 P0: 数据验证（强制，无交互确认）
        if VALIDATION_AVAILABLE and not args.quick_test:
            print("🔍 验证报价单数据...")
            valid, errors = validate_quotation_data(data)
            
            if not valid:
                print("❌ 数据验证失败，Word 报价单生成已终止:")
                for i, err in enumerate(errors, start=1):
                    print(f"  {i}. {err}")
                print()
                print("请检查数据文件，确保使用真实客户信息。")
                print("如需要测试，请使用 --quick-test 参数")
                sys.exit(1)
            
            print("✅ 数据验证通过")
            print()
        
        output = create_quotation_docx(args.output, data)
        print(f"✅ Word 报价单已生成：{output}")
        return
    
    # 没有参数时显示帮助
    parser.print_help()

if __name__ == '__main__':
    main()
