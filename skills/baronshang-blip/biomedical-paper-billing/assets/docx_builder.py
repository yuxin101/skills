#!/usr/bin/env python3
"""
docx_builder.py — Reusable Word document builder for biomedical manuscripts.

Usage:
    from assets.docx_builder import ManuscriptBuilder
    mb = ManuscriptBuilder()
    mb.title("My Paper Title")
    mb.running_title("Short Running Title")
    mb.section("Abstract")
    mb.abstract_label("Background", "GBS is ...")
    mb.body("Regular body paragraph text...")
    mb.section("Introduction")
    mb.body("First paragraph...")
    mb.subsection("Data Source")
    mb.body("Methods paragraph...")
    mb.reference_list(["1. Author A. Title. Journal. 2024;..."])
    mb.save("/workspace/output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class ManuscriptBuilder:
    def __init__(self):
        self.doc = Document()
        # Page margins
        for section in self.doc.sections:
            section.top_margin    = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin   = Cm(3.17)
            section.right_margin  = Cm(3.17)
        # Base style
        normal = self.doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(0)

    # ── Inline markup ─────────────────────────────────────────
    def _inline(self, para, text, size=11):
        """Parse **bold** and *italic* and add runs to para."""
        for part in re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text):
            if part.startswith('**') and part.endswith('**'):
                r = para.add_run(part[2:-2]); r.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = para.add_run(part[1:-1]); r.italic = True
            else:
                r = para.add_run(part)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(size)

    # ── Paragraph types ───────────────────────────────────────
    def title(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)

    def running_title(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(12)
        self._inline(p, text)

    def section(self, text):
        """Top-level section heading (Abstract, Introduction, Methods…)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(11)

    def subsection(self, text):
        """Sub-section heading (bold-italic)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text); r.bold = True; r.italic = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(11)

    def abstract_label(self, label, text):
        """e.g., abstract_label('Background', 'GBS is ...')"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing  = Pt(22)
        rl = p.add_run(label + ': '); rl.bold = True
        rl.font.name = 'Times New Roman'; rl.font.size = Pt(11)
        self._inline(p, text)

    def body(self, text):
        """Standard body paragraph (justified, first-line indent, double-spaced)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing  = Pt(22)
        self._inline(p, text)

    def formula(self, text):
        """Centred italic formula line."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text); r.italic = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(11)

    def reference_list(self, refs):
        """refs: list of formatted reference strings."""
        for ref in refs:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(-0.35)
            p.paragraph_format.left_indent       = Inches(0.35)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing  = Pt(20)
            self._inline(p, ref, size=10)

    def save(self, path):
        self.doc.save(path)
        print(f"Saved → {path}")


# ── Convenience: build from plain-text markdown-ish source ───
def build_from_markdown(md_path, out_path):
    """
    Lightweight parser: recognises
      # Title
      Running title: ...
      ## Section
      ### Subsection
      **Label:** text   (abstract labels)
      ΔYLDs = ...       (formula)
      1. ref text       (references, after ## References)
      Everything else → body paragraph
    """
    mb = ManuscriptBuilder()
    in_refs = False
    with open(md_path) as f:
        for raw in f:
            line = raw.strip()
            if not line:
                continue
            if line.startswith('# ') and not line.startswith('## '):
                mb.title(line[2:])
            elif line.lower().startswith('running title:'):
                mb.running_title(line)
            elif line.startswith('## '):
                heading = line[3:]
                if 'reference' in heading.lower():
                    in_refs = True
                mb.section(heading)
            elif line.startswith('### '):
                mb.subsection(line[4:])
            elif in_refs and re.match(r'^\d+\.', line):
                mb.reference_list([line])
            elif line.startswith('ΔYLDs'):
                mb.formula(line)
            elif re.match(r'^\*\*\w[^*]+:\*\*', line):
                m = re.match(r'^\*\*([^*]+):\*\*\s*(.*)', line)
                if m:
                    mb.abstract_label(m.group(1), m.group(2))
                else:
                    mb.body(line)
            else:
                mb.body(line)
    mb.save(out_path)


if __name__ == '__main__':
    import sys
    if len(sys.argv) == 3:
        build_from_markdown(sys.argv[1], sys.argv[2])
    else:
        print("Usage: docx_builder.py <input.md> <output.docx>")
