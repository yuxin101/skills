#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PRD 文档解析脚本

功能：解析 Word/PDF 文档，提取纯文本内容供智能体分析
支持格式：.docx, .pdf, .txt, .md
"""

import sys
import os


def parse_docx(file_path: str) -> str:
    """解析 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        return "[ERROR] 缺少依赖：python-docx，请安装后重试"

    try:
        doc = Document(file_path)
        content = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    
    except Exception as e:
        return f"[ERROR] 解析 Word 文档失败：{str(e)}"


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文档"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "[ERROR] 缺少依赖：PyPDF2，请安装后重试"

    try:
        reader = PdfReader(file_path)
        content = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                content.append(f"--- 第 {page_num} 页 ---")
                content.append(text.strip())
        
        return "\n".join(content)
    
    except Exception as e:
        return f"[ERROR] 解析 PDF 文档失败：{str(e)}"


def parse_txt(file_path: str) -> str:
    """解析纯文本文件"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        return "[ERROR] 无法识别文件编码，请将文件转换为 UTF-8 格式"
    
    except Exception as e:
        return f"[ERROR] 读取文件失败：{str(e)}"


def parse_document(file_path: str) -> str:
    """
    解析文档并返回纯文本内容
    
    参数：
        file_path: 文档路径
    
    返回：
        提取的纯文本内容
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件不存在：{file_path}"
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        return parse_docx(file_path)
    elif file_ext == '.pdf':
        return parse_pdf(file_path)
    elif file_ext in ['.txt', '.md']:
        return parse_txt(file_path)
    else:
        return f"[ERROR] 不支持的文件格式：{file_ext}，仅支持 .docx, .pdf, .txt, .md 格式"


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法：python doc_parser.py <文档路径>")
        print("支持格式：.docx, .pdf, .txt, .md")
        sys.exit(1)
    
    file_path = sys.argv[1]
    content = parse_document(file_path)
    print(content)


if __name__ == "__main__":
    main()
