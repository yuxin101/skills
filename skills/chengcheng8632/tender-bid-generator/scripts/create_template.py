#!/usr/bin/env python3
"""
生成投标文件Word模板
运行此脚本生成标准的投标文件模板
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_bid_template():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('投标文件')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（正本/副本）')
    run.font.size = Pt(16)
    run.font.name = '宋体'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 项目信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    cells = [
        ('项目名称', '[请填写项目名称]'),
        ('项目编号', '[请填写项目编号]'),
        ('投标单位', '[请填写投标单位全称]'),
        ('投标日期', '[请填写投标日期]')
    ]
    
    for i, (label, value) in enumerate(cells):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # 设置第一列加粗
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_page_break()
    
    # 目录
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目    录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    
    doc.add_paragraph()
    
    toc_items = [
        '一、投标函',
        '二、法定代表人授权书',
        '三、开标一览表/投标报价表',
        '四、资格证明文件',
        '    4.1 营业执照',
        '    4.2 资质证书',
        '    4.3 财务状况证明',
        '    4.4 业绩证明材料',
        '    4.5 其他资格证明',
        '五、技术响应文件',
        '    5.1 技术方案',
        '    5.2 技术规格响应表',
        '    5.3 项目实施计划',
        '    5.4 质量保证措施',
        '六、商务响应文件',
        '    6.1 交货期承诺',
        '    6.2 售后服务承诺',
        '    6.3 培训方案',
        '    6.4 其他商务条款响应',
        '七、其他材料'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5 if item.startswith('    ') else 0)
    
    doc.add_page_break()
    
    # 各章节模板
    sections = [
        ('一、投标函', '''
致：[采购单位名称]

我方已仔细研究了[项目名称]（项目编号：[项目编号]）的招标文件，愿意以人民币[金额大写]元（¥[金额小写]）的投标总价参与投标。

我方承诺：
1. 投标有效期为投标截止之日起[XX]天；
2. 严格按照招标文件要求履行义务；
3. 提供的所有资料真实有效；
4. 如中标，将按照招标文件要求签订合同并履行合同义务。

投标人：[公司名称]（盖章）
法定代表人或授权代表：（签字）
日期：XXXX年XX月XX日
        '''),
        ('二、法定代表人授权书', '''
本授权书声明：注册于[注册地址]的[公司名称]的[法定代表人姓名]代表本公司授权[被授权人姓名]（身份证号：[身份证号码]）为本公司的合法代理人，就[项目名称]的投标及合同执行，以本公司名义处理一切与之有关的事务。

本授权书于XXXX年XX月XX日签字生效，有效期至[有效期截止日期]。

被授权人无转委托权。

特此声明。

法定代表人：（签字或盖章）
被授权人：（签字）
单位名称：（盖章）
日期：XXXX年XX月XX日

附：被授权人身份证复印件（正反面）
        '''),
        ('三、投标报价表', '''
项目名称：[项目名称]
项目编号：[项目编号]

序号 | 项目内容 | 数量 | 单价（元） | 总价（元） | 备注
-----|---------|-----|-----------|-----------|-----
1    | [内容]  | [数] | [单价]    | [总价]    | [备注]
2    | [内容]  | [数] | [单价]    | [总价]    | [备注]
...  | ...     | ... | ...       | ...       | ...
     | 合计    |     |           | [总金额]  |

投标总价（大写）：人民币[金额大写]元整

投标人：（盖章）
法定代表人或授权代表：（签字）
日期：XXXX年XX月XX日
        '''),
        ('四、资格证明文件', '''
4.1 营业执照副本复印件
[粘贴营业执照副本复印件，加盖公章]

4.2 资质证书复印件
[粘贴相关资质证书复印件，加盖公章]

4.3 财务状况证明
[粘贴经审计的财务报表或银行资信证明，加盖公章]

4.4 业绩证明材料
[提供类似项目业绩清单及合同复印件，加盖公章]

4.5 无重大违法记录声明
[按照招标文件要求的格式提供声明]

4.6 其他资格证明
[根据招标文件要求提供其他证明材料]
        '''),
        ('五、技术响应文件', '''
5.1 技术方案概述

[在此详细描述技术方案，包括：]
- 总体技术架构
- 核心技术路线
- 关键技术特点
- 技术创新点

5.2 技术规格响应表

| 序号 | 招标要求 | 投标响应 | 偏离情况 | 说明 |
|-----|---------|---------|---------|-----|
| 1   | [复制招标要求] | [实际参数] | 无偏离/正偏离 | [说明] |
| 2   | [复制招标要求] | [实际参数] | 无偏离/正偏离 | [说明] |
| ... | ... | ... | ... | ... |

5.3 项目实施计划

[描述项目实施的时间安排、人员配置、里程碑等]

5.4 质量保证措施

[描述质量管理体系、测试方案、质量控制措施等]
        '''),
        ('六、商务响应文件', '''
6.1 交货期/服务期承诺

我方承诺严格按照招标文件要求的交货期/服务期执行：
- 合同签订后[X]天内完成[具体交付内容]
- [其他时间承诺]

6.2 售后服务承诺

我方承诺提供以下售后服务：
- 质保期：[X]年
- 响应时间：[X]小时内响应，[X]小时内到达现场
- 服务热线：[电话号码]
- 备品备件供应承诺
- 其他服务承诺

6.3 培训方案

[描述培训内容、培训方式、培训时间安排等]

6.4 商务条款响应

[逐条响应招标文件中的商务条款要求]
        '''),
        ('七、其他材料', '''
[根据招标文件要求或需要补充的其他材料，如：]
- 企业简介
- 荣誉证书
- 专利证书
- 产品彩页
- 其他支持材料
        ''')
    ]
    
    for title, content in sections:
        # 章节标题
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        
        doc.add_paragraph()
        
        # 内容
        for line in content.strip().split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Inches(0.3)
        
        doc.add_page_break()
    
    return doc

def main():
    try:
        doc = create_bid_template()
        output_path = 'assets/bid-template.docx'
        doc.save(output_path)
        print(f'模板已生成: {output_path}')
    except ImportError:
        print('错误：未安装python-docx，请运行 pip install python-docx')
        return 1
    return 0

if __name__ == '__main__':
    exit(main())
