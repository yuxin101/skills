"""
md2docx - Markdown to Word Document Converter

Converts Markdown files to Word documents using Pandoc with enhanced options
to ensure all Markdown syntax is properly converted.

This module provides functionality to convert Markdown files to Word documents
with proper font handling for Chinese characters and table borders.
It uses a two-stage process: first using Pandoc for the conversion,
then post-processing with python-docx to enhance formatting.
"""

import subprocess
import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Optional, Union
from docx import Document
from docx.oxml.ns import qn


class ConversionError(Exception):
    """Base exception for conversion errors."""
    pass


class FileNotFoundError(ConversionError):
    """Raised when input file is not found."""
    pass


class DependencyError(ConversionError):
    """Raised when required dependencies are missing."""
    pass


class ValidationError(ConversionError):
    """Raised when validation fails."""
    pass


def _apply_font_settings(paragraphs, font_name: str = 'Microsoft YaHei'):
    """
    Apply font settings to a collection of paragraphs.

    Args:
        paragraphs: Collection of paragraph objects
        font_name: Font name to apply
    """
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_font_and_borders(docx_path: Union[str, Path]) -> bool:
    """
    Set fonts and borders for the Word document using python-docx.

    Args:
        docx_path: Path to the Word document

    Returns:
        True if successful, False otherwise

    Raises:
        FileNotFoundError: If the document file doesn't exist
    """
    try:
        # Verify document exists
        if not Path(docx_path).exists():
            raise FileNotFoundError(f"Word文档不存在: {docx_path}")

        # Open the document
        doc = Document(docx_path)

        # Define font settings
        font_name = 'Microsoft YaHei'  # Preferred font

        # Apply font to all paragraphs
        _apply_font_settings(doc.paragraphs, font_name)

        # Apply font to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Apply font to text in cells
                    _apply_font_settings(cell.paragraphs, font_name)

                    # Set border for each cell
                    # Access the table cell properties directly
                    tc_pr = cell._tc.get_or_add_tcPr()

                    # Add new borders
                    borders = ['top', 'bottom', 'left', 'right']
                    for border in borders:
                        # Create border element
                        border_elem = tc_pr.makeelement(
                            f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{border}'
                        )
                        border_elem.set(qn('w:val'), 'single')
                        border_elem.set(qn('w:sz'), '4')  # Size: 1/8 pt
                        border_elem.set(qn('w:space'), '0')
                        border_elem.set(qn('w:color'), 'auto')
                        tc_pr.append(border_elem)

        # Apply font to headers and footers if they exist
        for section in doc.sections:
            # Header
            if section.header:
                _apply_font_settings(section.header.paragraphs, font_name)

            # Footer
            if section.footer:
                _apply_font_settings(section.footer.paragraphs, font_name)

        # Save the modified document
        doc.save(docx_path)
        print(f"成功设置字体和边框: {docx_path}")
        return True

    except FileNotFoundError as e:
        raise e
    except Exception as e:
        error_msg = (
            f"设置字体和边框失败: {str(e)}\n"
            f"可能的原因及解决方案:\n"
            f"1. 文件被其他程序占用，请关闭Word等程序后重试\n"
            f"2. 文件权限不足，请检查文件访问权限\n"
            f"3. 文件损坏，请使用备份文件重试"
        )
        print(error_msg)
        return False


def convert_md_to_docx(
    input_file: Union[str, Path],
    output_file: Union[str, Path],
    template: Optional[Union[str, Path]] = None,
    reference_docx: Optional[Union[str, Path]] = None
) -> bool:
    """
    Convert a markdown file to a Word document using two-stage process:
    1. Pandoc conversion from Markdown to Word
    2. Post-processing with python-docx to set fonts and borders

    Args:
        input_file: Input markdown file path
        output_file: Output Word document path
        template: Custom template file (optional)
        reference_docx: Reference Word document for styling (optional)

    Returns:
        True if conversion successful, False otherwise

    Raises:
        FileNotFoundError: If input file doesn't exist
        DependencyError: If Pandoc is not installed or not in PATH
        ValidationError: If file formats are incorrect
    """
    input_file = Path(input_file)
    output_file = Path(output_file)

    # Validate input file
    if not input_file.exists():
        raise FileNotFoundError(
            f"输入Markdown文件不存在: {input_file}\n"
            f"请检查文件路径是否正确"
        )

    # Validate file extension
    if input_file.suffix.lower() != '.md':
        raise ValidationError(
            f"输入文件不是有效的Markdown文件: {input_file}\n"
            f"文件扩展名应为 .md"
        )

    # Create output directory if it doesn't exist
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # To add date at the end of document, we first modify the Markdown file content temporarily
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Add date at the end of the document
    date_line = f"\n\n---\n\n{datetime.now().strftime('%Y年%m月%d日')}"
    temp_input_file = input_file.with_suffix(input_file.suffix + '.temp.md')

    try:
        with open(temp_input_file, 'w', encoding='utf-8') as f:
            f.write(content + date_line)

        # Build pandoc command with comprehensive options
        cmd = [
            "pandoc",
            str(temp_input_file),  # Use temporary file
            "-f", "markdown",
            "-t", "docx",
            "-o", str(output_file),
            "--standalone",                    # Produce output with header and footer
            "--toc",                           # Include table of contents
            "--toc-depth=6",                   # Support H1-H6 in TOC
            "--syntax-highlighting=tango",     # Syntax highlighting style (replaces deprecated --highlight-style)
            "--wrap=none",                     # Preserve line breaks
            "--columns=10000",                 # Prevent text wrapping
        ]

        # Additional options for better Markdown support
        cmd.extend([
            "--markdown-headings=atx",         # Use ATX headings (# H1, ## H2, etc.)
            "--preserve-tabs",                 # Preserve tab characters
        ])

        # Add template if provided
        if template:
            template_path = Path(template)
            if not template_path.exists():
                raise FileNotFoundError(f"模板文件不存在: {template_path}")
            cmd.extend(["--reference-doc", str(template_path)])
        elif reference_docx:
            ref_path = Path(reference_docx)
            if not ref_path.exists():
                raise FileNotFoundError(f"参考文档不存在: {ref_path}")
            cmd.extend(["--reference-doc", str(ref_path)])

        try:
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                check=True
            )
            print(f"第一阶段完成: {input_file} -> {output_file}")
        except subprocess.CalledProcessError as e:
            print("错误: Pandoc转换失败")
            print(f"命令: {' '.join(cmd)}")
            print(f"错误输出: {e.stderr}")
            return False
        except FileNotFoundError:
            raise DependencyError(
                "Pandoc 未安装或未加入 PATH。\n"
                "解决方法：\n"
                "1. Windows: winget install pandoc\n"
                "2. macOS: brew install pandoc\n"
                "3. Linux: sudo apt install pandoc\n"
                "详细说明：https://pandoc.org/installing.html"
            )

        # Second stage: Use python-docx post-processing to set fonts and table borders
        print("开始第二阶段：设置字体和表格边框...")
        success = set_font_and_borders(output_file)

        if success:
            print(f"成功: {input_file} -> {output_file} (两阶段转换完成)")
            return True
        else:
            print(f"警告: {input_file} -> {output_file} (Pandoc转换成功，但后处理失败)")
            return False

    finally:
        # Clean up temporary file
        if temp_input_file.exists():
            temp_input_file.unlink()


def batch_convert_md_to_docx(
    input_dir: Union[str, Path],
    output_dir: Union[str, Path],
    recursive: bool = False
) -> dict[str, int]:
    """
    Convert all markdown files in a directory to Word documents.

    Args:
        input_dir: Input directory containing markdown files
        output_dir: Output directory for Word documents
        recursive: Whether to process subdirectories recursively

    Returns:
        Dictionary with success/failure counts
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    if not input_path.exists():
        raise FileNotFoundError(f"输入目录不存在: {input_dir}")

    # Find all markdown files
    if recursive:
        md_files = list(input_path.rglob("*.md"))
    else:
        md_files = list(input_path.glob("*.md"))

    results = {"success": 0, "failed": 0}

    for md_file in md_files:
        # Calculate relative path to maintain directory structure
        rel_path = md_file.relative_to(input_path)
        docx_file = output_path / rel_path.with_suffix('.docx')

        try:
            success = convert_md_to_docx(str(md_file), str(docx_file))
            if success:
                results["success"] += 1
            else:
                results["failed"] += 1
        except Exception as e:
            print(f"处理文件时出错 {md_file}: {str(e)}")
            results["failed"] += 1

    return results


def validate_pandoc_available() -> bool:
    """
    Check if Pandoc is available in the system.

    Returns:
        True if Pandoc is available, False otherwise
    """
    try:
        subprocess.run(
            ['pandoc', '--version'],
            capture_output=True,
            text=True,
            check=True
        )
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python md2docx.py <input.md> <output.docx>")
        print("      python md2docx.py <input_dir> <output_dir> [--recursive]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    # Check if input is a directory
    if os.path.isdir(input_path):
        recursive = "--recursive" in sys.argv
        try:
            results = batch_convert_md_to_docx(input_path, output_path, recursive)
            print(f"批量转换完成: {results['success']} 成功, {results['failed']} 失败")
        except Exception as e:
            print(f"批量转换失败: {str(e)}")
            sys.exit(1)
    else:
        try:
            success = convert_md_to_docx(input_path, output_path)
            if success:
                print("转换完成!")
            else:
                print("转换失败!")
                sys.exit(1)
        except ConversionError as e:
            print(f"转换错误: {str(e)}")
            sys.exit(1)
        except Exception as e:
            print(f"未知错误: {str(e)}")
            sys.exit(1)