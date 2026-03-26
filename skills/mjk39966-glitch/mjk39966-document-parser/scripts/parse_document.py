#!/usr/bin/env python3
"""
Document parser supporting .docx, .pdf, and .txt formats.
Extracts text content and tables from documents.
"""

import sys
import json
from pathlib import Path

def parse_txt(file_path):
    """Parse plain text file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return {
        'text': content,
        'tables': [],
        'metadata': {
            'format': 'txt',
            'size': len(content)
        }
    }

def parse_docx(file_path):
    """Parse DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return {
            'error': 'python-docx not installed. Run: pip install python-docx'
        }
    
    doc = Document(file_path)
    
    # Extract paragraphs
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    return {
        'text': '\n\n'.join(text_parts),
        'tables': tables,
        'metadata': {
            'format': 'docx',
            'paragraphs': len(doc.paragraphs),
            'tables': len(tables)
        }
    }

def parse_pdf(file_path):
    """Parse PDF file using PyPDF2 and tabula."""
    try:
        import PyPDF2
    except ImportError:
        return {
            'error': 'PyPDF2 not installed. Run: pip install PyPDF2'
        }
    
    text_parts = []
    
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"=== Page {page_num + 1} ===\n{text}")
    
    # Try to extract tables with pdfplumber (better table support than PyPDF2)
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
    except ImportError:
        # Fallback: note that table extraction needs pdfplumber
        tables = [["Note: Install pdfplumber for table extraction (pip install pdfplumber)"]]
    
    return {
        'text': '\n\n'.join(text_parts),
        'tables': tables,
        'metadata': {
            'format': 'pdf',
            'pages': num_pages,
            'tables': len(tables)
        }
    }

def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            'error': 'Usage: parse_document.py <file_path> [--format json|text]'
        }))
        sys.exit(1)
    
    file_path = Path(sys.argv[1])
    output_format = 'json'
    
    if '--format' in sys.argv:
        idx = sys.argv.index('--format')
        if idx + 1 < len(sys.argv):
            output_format = sys.argv[idx + 1]
    
    if not file_path.exists():
        print(json.dumps({
            'error': f'File not found: {file_path}'
        }))
        sys.exit(1)
    
    # Determine parser based on extension
    ext = file_path.suffix.lower()
    
    if ext == '.txt':
        result = parse_txt(file_path)
    elif ext == '.docx':
        result = parse_docx(file_path)
    elif ext == '.pdf':
        result = parse_pdf(file_path)
    else:
        result = {
            'error': f'Unsupported format: {ext}. Supported: .txt, .docx, .pdf'
        }
    
    # Output result
    if output_format == 'text':
        if 'error' in result:
            print(f"ERROR: {result['error']}")
        else:
            print("=" * 60)
            print("EXTRACTED TEXT:")
            print("=" * 60)
            print(result['text'])
            
            if result['tables']:
                print("\n" + "=" * 60)
                print(f"TABLES FOUND: {len(result['tables'])}")
                print("=" * 60)
                for i, table in enumerate(result['tables'], 1):
                    print(f"\nTable {i}:")
                    for row in table:
                        print(" | ".join(str(cell) for cell in row))
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == '__main__':
    main()
