#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换工具
使用配置系统驱动格式化，支持自定义YAML配置和预设格式

配置说明详见: assets/presets/*.yaml 和 references/config-reference.md
"""

import os
import argparse
import re
import glob
import tempfile

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from PIL import Image

# 导入配置模块
from config import Config, load_config, get_preset, get_default_preset, list_presets, get_config, set_config

# 导入功能模块
from formatter import (
    convert_quotes_to_chinese,
    parse_text_formatting,
    set_run_format_with_styles,
    set_paragraph_format,
    parse_alignment,
    hex_to_rgb,
)
from table_handler import (
    is_table_row,
    create_word_table,
    create_word_table_from_html,
)
from chart_handler import create_mermaid_chart


# ============================================================================
# 图片处理
# ============================================================================

def get_image_output_path(md_file_path, png_filename):
    """获取图片输出路径，确保目录存在"""
    md_dir = os.path.dirname(os.path.abspath(md_file_path))
    md_filename_base = os.path.splitext(os.path.basename(md_file_path))[0]
    image_dir = os.path.join(md_dir, f"{md_filename_base}_images")

    if not os.path.exists(image_dir):
        try:
            os.makedirs(image_dir)
            print(f"📂 创建图片目录: {os.path.relpath(image_dir)}")
        except OSError as e:
            print(f"⚠️ 创建目录失败: {e}")
            return None

    return os.path.join(image_dir, png_filename)


def _postprocess_image_for_word(image, target_display_cm, target_dpi=260):
    """根据目标显示宽度与DPI对图像进行高质量下采样"""
    try:
        target_inches = float(target_display_cm) / 2.54
        target_px_width = max(1, int(target_inches * target_dpi))
        if image.width > target_px_width:
            new_height = int(image.height * (target_px_width / image.width))
            image = image.resize((target_px_width, new_height), Image.LANCZOS)
    except Exception:
        pass
    return image


def insert_image_to_word(doc, image):
    """将PIL图片对象插入Word文档"""
    config = get_config()
    image_config = config.get('image', {})
    page_config = config.get('page', {})

    display_ratio = image_config.get('display_ratio', 0.92)
    max_width_cm = image_config.get('max_width_cm', 14.2)
    target_dpi = image_config.get('target_dpi', 260)

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        page_width = page_config.get('width', 21.0)
        margin_left = page_config.get('margin_left', 3.18)
        margin_right = page_config.get('margin_right', 3.18)
        available_width_cm = page_width - margin_left - margin_right
        target_display_cm = min(available_width_cm * display_ratio, max_width_cm)
        image = _postprocess_image_for_word(image, target_display_cm, target_dpi=target_dpi)
        try:
            image.save(temp_file.name, format='PNG', optimize=True, compress_level=9)
        except Exception:
            image.save(temp_file.name, format='PNG', optimize=True)
        temp_filename = temp_file.name

    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img_width_cm = target_display_cm
        run = paragraph.add_run()
        run.add_picture(temp_filename, width=Cm(img_width_cm))
    finally:
        try:
            os.unlink(temp_filename)
        except:
            pass


# ============================================================================
# 文档结构元素
# ============================================================================

def add_horizontal_line(doc):
    """添加分割线"""
    config = get_config()
    hr_config = config.get('horizontal_rule', {})
    
    p = doc.add_paragraph()
    p.alignment = parse_alignment(hr_config.get('alignment', 'center'))
    
    character = hr_config.get('character', '─')
    repeat_count = hr_config.get('repeat_count', 55)
    run = p.add_run(character * repeat_count)
    
    font_name = hr_config.get('font', 'Times New Roman')
    font_size = hr_config.get('size', 12)
    color_hex = hr_config.get('color', '#808080')
    
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = hex_to_rgb(color_hex)


def add_task_list(doc, line):
    """添加任务列表"""
    config = get_config()
    task_config = config.get('lists.task', {})
    
    is_checked = line.startswith(('- [x]', '- [X]'))
    text = line[5:].strip()
    p = doc.add_paragraph()
    
    checked_mark = task_config.get('checked', '☑')
    unchecked_mark = task_config.get('unchecked', '☐')
    checkbox_run = p.add_run(f'{checked_mark} ' if is_checked else f'{unchecked_mark} ')
    set_run_format_with_styles(checkbox_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_bullet_list(doc, line):
    """添加无序列表"""
    config = get_config()
    bullet_config = config.get('lists.bullet', {})
    
    text = line[2:].strip()
    p = doc.add_paragraph()
    
    marker = bullet_config.get('marker', '•')
    bullet_run = p.add_run(f'{marker} ')
    set_run_format_with_styles(bullet_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_numbered_list(doc, line):
    """添加有序列表"""
    p = doc.add_paragraph()
    parse_text_formatting(p, line)
    set_paragraph_format(p)


def add_quote(doc, text):
    """添加引用块"""
    config = get_config()
    quote_config = config.get('quote', {})
    
    lines = text.split('\n')
    
    bg_color = quote_config.get('background_color', '#EAEAEA')
    left_indent = quote_config.get('left_indent_inches', 0.2)
    font_size = quote_config.get('font_size', 9)
    line_spacing = quote_config.get('line_spacing', 1.5)
    
    for line_index, line in enumerate(lines):
        if not line.strip():
            p = doc.add_paragraph()
            set_paragraph_format(p, is_quote=True)
            continue
        
        p = doc.add_paragraph()
        
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color.lstrip('#'))
        pPr.append(shd)
        
        p.paragraph_format.left_indent = Inches(left_indent)
        p.paragraph_format.line_spacing = line_spacing
        
        bullet_match = re.match(r'^\s*([-*+])\s+', line)
        number_match = re.match(r'^\s*(\d+\.)\s+', line)
        
        list_marker_run = None
        
        if bullet_match:
            indent_and_bullet = '    •  '
            list_marker_run = p.add_run(indent_and_bullet)
            line = line[bullet_match.end():]
        elif number_match:
            indent_and_number = f'    {number_match.group(1)} '
            list_marker_run = p.add_run(indent_and_number)
            line = line[number_match.end():]
        
        if list_marker_run:
            list_marker_run.font.size = Pt(font_size)
            set_run_format_with_styles(list_marker_run, {}, is_quote=True)
        
        parse_text_formatting(p, line, is_quote=True)
        set_paragraph_format(p, is_quote=True)
        
        for run in p.runs:
            run.font.size = Pt(font_size)


def add_code_block(doc, code_lines, language):
    """添加代码块"""
    config = get_config()
    code_config = config.get('code_block', {})
    
    label_config = code_config.get('label', {})
    if language:
        lang_p = doc.add_paragraph()
        lang_run = lang_p.add_run(f"[{language}]")
        lang_run.font.name = label_config.get('font', 'Times New Roman')
        lang_run.font.size = Pt(label_config.get('size', 10))
        lang_run.font.color.rgb = hex_to_rgb(label_config.get('color', '#808080'))
    
    content_config = code_config.get('content', {})
    left_indent = content_config.get('left_indent', 24)
    line_spacing = content_config.get('line_spacing', 1.2)
    font_name = content_config.get('font', 'Times New Roman')
    font_size = content_config.get('size', 10)
    color_hex = content_config.get('color', '#333333')
    
    for code_line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(code_line or ' ')
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = hex_to_rgb(color_hex)
        p.paragraph_format.left_indent = Pt(left_indent)
        p.paragraph_format.line_spacing = line_spacing


def add_page_number(doc):
    """添加页码"""
    config = get_config()
    page_number_config = config.get('page_number', {})
    
    if not page_number_config.get('enabled', True):
        return
    
    try:
        section = doc.sections[0]
        footer = section.footer
        
        for para in footer.paragraphs:
            para.clear()
        
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        
        position = page_number_config.get('position', 'center')
        if position == 'left':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif position == 'right':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        page_format = page_number_config.get('format', '1/x')
        if '1' in page_format:
            run = footer_para.add_run()
            fld_char_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_begin)
            instr_text = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
            run._r.append(instr_text)
            fld_char_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_end)
        
        if '/' in page_format:
            sep_run = footer_para.add_run("/")
        
        if 'x' in page_format:
            total_run = footer_para.add_run()
            fld_char_begin2 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_begin2)
            instr_text2 = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> NUMPAGES </w:instrText>')
            total_run._r.append(instr_text2)
            fld_char_end2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_end2)
        
        font_name = page_number_config.get('font', 'Times New Roman')
        font_size = page_number_config.get('size', 10.5)
        
        for run in footer_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    
    except Exception as e:
        print(f"⚠️  页码添加失败，将跳过页码设置: {e}")
        pass


# ============================================================================
# 工具函数
# ============================================================================

def find_template_file():
    """查找模板文件"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_dir = os.path.dirname(script_dir)
    templates_dir = os.path.join(skill_dir, 'assets', 'templates')
    docx_files = glob.glob(os.path.join(templates_dir, "*.docx"))
    
    for docx_file in docx_files:
        filename = os.path.basename(docx_file).lower()
        if not any(keyword in filename for keyword in ['完整版', 'test', 'output', '输出']):
            if '模板' in filename or 'template' in filename:
                return docx_file
    
    return docx_files[0] if docx_files else None


def find_md_files():
    """查找脚本所在目录下的所有 .md 文件"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_files = glob.glob(os.path.join(script_dir, "*.md"))
    return md_files


def generate_output_filename(md_file):
    """根据输入文件名生成输出文件名"""
    base_name = os.path.splitext(md_file)[0]
    return f"{base_name}_完整版.docx"


def debug_quotes_in_file(file_path):
    """简化的引号调试"""
    print("🔍 检查文件中的引号...")
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    ascii_double = content.count('"')
    chinese_open = content.count('"')
    chinese_close = content.count('"')
    
    print(f"📊 引号统计: ASCII双引号={ascii_double}, 中文开引号={chinese_open}, 中文闭引号={chinese_close}")
    
    for i, line in enumerate(content.split('\n'), 1):
        if '"' in line:
            print(f"🎯 测试第{i}行: {line.strip()}")
            _ = convert_quotes_to_chinese(line.strip())
            break
    
    print("-" * 30)


# ============================================================================
# 核心转换流程
# ============================================================================

def create_word_document(md_file_path, output_path, template_file=None, config: Config = None):
    """从Markdown文件创建格式化的Word文档"""
    if config is None:
        config = get_config()

    print(f"📄 正在处理: {md_file_path}")
    print(f"📋 使用配置: {config.name}")

    if config.get('quotes.convert_to_chinese', True):
        debug_quotes_in_file(md_file_path)

    # 用于存储模板的header/footer XML元素
    template_header_xmls = []
    template_footer_xmls = []
    use_template_headers = False

    # 创建或加载文档
    if template_file and template_file != "none" and os.path.exists(template_file):
        print(f"📋 使用模板文件: {os.path.basename(template_file)}")

        # 直接使用模板文件，保留所有格式（包括页眉页脚）
        print("📄 直接打开模板文件")
        doc = Document(template_file)

        # 清空模板中的正文内容（保留页眉页脚和sectPr）
        # 获取body元素
        body = doc._element.body

        # 记住sectPr的位置和内容
        sectPr = body.find(qn('w:sectPr'))

        # 移除body中的所有子元素（除了sectPr）
        for child in list(body):
            if child.tag != qn('w:sectPr'):
                body.remove(child)

        use_template_headers = True
        print("✅ 已清空模板内容，保留页眉页脚")
    else:
        print("📄 创建新文档（不使用模板）")
        doc = Document()
        template_header_xmls = []
        template_footer_xmls = []
        template_header_rels = []
        template_footer_rels = []
        template_media_files = []
        template_sectPr_refs = []
        template_doc_rels = {}

    # 设置默认字体
    try:
        normal_style = doc.styles['Normal']
        font_config = config.get('fonts.default', {})
        normal_style.font.name = font_config.get('ascii', 'Times New Roman')
        normal_style.font.size = Pt(font_config.get('size', 10.5))
        normal_style._element.rPr.rFonts.set(qn('w:ascii'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('name', '仿宋_GB2312'))
        normal_style._element.rPr.rFonts.set(qn('w:cs'), font_config.get('ascii', 'Times New Roman'))
    except Exception as _:
        pass

    # 设置页面大小和页边距
    for section in doc.sections:
        page_config = config.get('page', {})
        section.page_width = Cm(page_config.get('width', 21.0))
        section.page_height = Cm(page_config.get('height', 29.7))
        section.top_margin = Cm(page_config.get('margin_top', 2.54))
        section.bottom_margin = Cm(page_config.get('margin_bottom', 2.54))
        section.left_margin = Cm(page_config.get('margin_left', 3.18))
        section.right_margin = Cm(page_config.get('margin_right', 3.18))
    
    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(md_file_path, 'r', encoding='gbk') as f:
            content = f.read()
    
    lines = content.split('\n')
    has_body_before_first_h2 = False
    has_seen_h2 = False
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Mermaid 图表
        if re.match(r'^```\s*mermaid\b', line):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            
            if mermaid_lines:
                mermaid_code = '\n'.join(mermaid_lines)
                create_mermaid_chart(
                    doc,
                    lambda img: insert_image_to_word(doc, img),
                    get_image_output_path,
                    lambda: doc.add_paragraph(),
                    lambda p: set_paragraph_format(p),
                    mermaid_code,
                    md_file_path
                )
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Mermaid图表")
            continue
        
        # 代码块
        if line.startswith('```'):
            code_lines = []
            language = line[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(doc, code_lines, language)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            print("✅ 处理代码块")
            continue
        
        # HTML 表格
        if '<table>' in line.lower():
            html_table_content = []
            while i < len(lines):
                html_table_content.append(lines[i])
                if '</table>' in lines[i].lower():
                    i += 1
                    break
                i += 1
            if html_table_content:
                create_word_table_from_html(doc, '\n'.join(html_table_content))
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
        
        # Markdown 表格
        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                create_word_table(doc, table_lines)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Markdown表格: {len(table_lines)} 行")
            continue
        
        # 分割线
        if line in ['---', '***', '___']:
            add_horizontal_line(doc)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 任务列表
        if line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            add_task_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 无序列表
        if line.startswith(('- ', '* ', '+ ')):
            add_bullet_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+\.\s', line):
            add_numbered_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 引用块
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i][1:].strip())
                i += 1
            if quote_lines:
                add_quote(doc, '\n'.join(quote_lines))
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
        
        # 标题
        if line.startswith('# '):
            title = convert_quotes_to_chinese(line[2:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=1)
            set_paragraph_format(p, title_level=1)
        elif line.startswith('## '):
            if has_seen_h2 or has_body_before_first_h2:
                doc.add_paragraph("")
            title = convert_quotes_to_chinese(line[3:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=2)
            set_paragraph_format(p, title_level=2)
            has_seen_h2 = True
        elif line.startswith('### '):
            title = convert_quotes_to_chinese(line[4:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=3)
            set_paragraph_format(p, title_level=3)
        elif line.startswith('#### '):
            title = convert_quotes_to_chinese(line[5:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=4)
            set_paragraph_format(p, title_level=4)
        else:
            if line:
                p = doc.add_paragraph()
                parse_text_formatting(p, line)
                set_paragraph_format(p)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
        
        i += 1
    
    # 添加页码（仅在没有模板时）
    if not use_template_headers:
        add_page_number(doc)

    # 保存文档
    doc.save(output_path)

    print(f"✅ Word文档已生成: {output_path}")


# ============================================================================
# CLI 入口
# ============================================================================

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='Markdown到Word文档转换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.md
  %(prog)s input.md --preset=academic
  %(prog)s input.md --config=my-config.yaml
  %(prog)s input.md output.docx
  %(prog)s --list-presets
        """
    )
    
    parser.add_argument('input', nargs='?', help='输入的 Markdown 文件')
    parser.add_argument('output', nargs='?', help='输出的 Word 文件')
    parser.add_argument('--preset', '-p', help='使用预设配置', default='legal')
    parser.add_argument('--config', '-c', help='使用自定义配置文件 (YAML格式)')
    parser.add_argument('--list-presets', action='store_true', help='列出所有可用的预设配置')
    parser.add_argument('--template', '-t', help='Word模板文件路径')
    
    args = parser.parse_args()
    
    if args.list_presets:
        print("可用的预设配置:")
        presets = list_presets()
        if presets:
            for preset in presets:
                cfg = get_preset(preset)
                if cfg:
                    print(f"  - {preset}: {cfg.description}")
        else:
            print("  没有可用的预设配置")
        return
    
    config = None
    if args.config:
        config = load_config(args.config)
        if config is None:
            print(f"❌ 无法加载配置文件: {args.config}")
            return
        print(f"📋 使用配置文件: {args.config}")
    elif args.preset:
        config = get_preset(args.preset)
        if config is None:
            print(f"❌ 预设不存在: {args.preset}")
            print(f"可用预设: {', '.join(list_presets())}")
            return
        print(f"📋 使用预设: {args.preset}")
    
    if config is None:
        config = get_default_preset()
    
    set_config(config)
    
    if not args.input:
        auto_mode(config)
        return
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = args.input
    if not os.path.isabs(md_file):
        alt = os.path.join(script_dir, md_file)
        if os.path.exists(alt):
            md_file = alt
    
    if not os.path.exists(md_file):
        print(f"❌ 错误: 找不到文件 {md_file}")
        return
    
    output_file = args.output if args.output else generate_output_filename(md_file)
    template_file = args.template if args.template else find_template_file()
    
    try:
        create_word_document(md_file, output_file, template_file, config)
        print_success_info(output_file, config)
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()


def auto_mode(config: Config):
    """自动模式：处理当前目录下的所有.md文件"""
    md_files = find_md_files()
    
    if not md_files:
        print("❌ 当前目录下没有找到.md文件")
        print("\n💡 使用方法:")
        print("1. 将此脚本放在包含.md文件的文件夹中")
        print("2. 或者运行: python md2word.py 输入文件.md")
        print("3. 或者运行: python md2word.py 输入文件.md --preset=academic")
        print("\n📋 可用预设:")
        presets = list_presets()
        if presets:
            for preset in presets:
                cfg = get_preset(preset)
                if cfg:
                    print(f"  - {preset}: {cfg.description}")
        return
    
    print(f"🔍 找到 {len(md_files)} 个Markdown文件:")
    for i, md_file in enumerate(md_files, 1):
        print(f"  {i}. {md_file}")
    
    print("\n开始转换...")
    
    template_file = find_template_file()
    success_count = 0
    
    for md_file in md_files:
        output_file = generate_output_filename(md_file)
        try:
            create_word_document(md_file, output_file, template_file, config)
            success_count += 1
        except Exception as e:
            print(f"❌ 处理 {md_file} 时出错: {e}")
    
    print(f"\n✅ 转换完成！成功处理 {success_count}/{len(md_files)} 个文件")
    print_success_info(None, config)


def print_success_info(filename=None, config: Config = None):
    """打印成功信息"""
    if config is None:
        config = get_config()
    
    print("\n📋 自动应用的格式:")
    
    page_config = config.get('page', {})
    print(f"📄 页面大小: {page_config.get('width', 21.0)}cm × {page_config.get('height', 29.7)}cm")
    print(f"📐 页边距: 上下{page_config.get('margin_top', 2.54)}cm，左右{page_config.get('margin_left', 3.18)}cm")
    
    font_config = config.get('fonts.default', {})
    print(f"📝 字体: {font_config.get('name', '仿宋_GB2312')}")
    print(f"📏 字号: {font_config.get('size', 12)}pt")
    
    paragraph_config = config.get('paragraph', {})
    print(f"📐 行距: {paragraph_config.get('line_spacing', 1.5)}倍")
    
    title1_config = config.get('titles.level1', {})
    print(f"🎯 一级标题: {title1_config.get('size', 15)}pt，{'加粗' if title1_config.get('bold') else '常规'}")
    
    page_number_config = config.get('page_number', {})
    if page_number_config.get('enabled', True):
        print(f"📄 页码设置: {page_number_config.get('format', '1/x')}格式")
    
    quotes_config = config.get('quotes', {})
    if quotes_config.get('convert_to_chinese', True):
        print("💬 引号转换: 英文引号自动转为中文引号")
    
    print("📊 表格支持: Markdown表格自动转换")
    print("📈 图表支持: Mermaid图表本地渲染")
    print("✨ 格式支持: **加粗**、*斜体*、<u>下划线</u>、~~删除线~~")
    print("\n🎯 完全无需手动调整！直接可用！")
    
    if filename:
        print(f"\n📁 输出文件: {filename}")


if __name__ == "__main__":
    main()
