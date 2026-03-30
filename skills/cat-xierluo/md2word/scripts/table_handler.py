#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格处理模块
处理 Markdown 表格和 HTML 表格的解析与转换
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from bs4 import BeautifulSoup

# 导入配置模块
from config import Config, get_config


def set_cell_background_color(cell, color_hex):
    """设置单元格背景色"""
    if not color_hex:
        return
    color = color_hex.lstrip('#')
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def is_separator_line(line):
    """判断是否是表格分隔行。分隔行必须包含'-'，且只能包含'|', '-', ':', ' '等符号。"""
    line = line.strip()
    if not line or '-' not in line:
        return False
    return all(c in '|-: 	' for c in line)


def is_table_row(line):
    """判断是否是表格行"""
    if not line or not line.strip():
        return False

    line = line.strip()

    # 检查是否是分隔行
    if is_separator_line(line):
        return True

    # 检查是否是数据行（包含 |）
    # 这里的逻辑保持宽松，依赖于主循环中对其他块级元素的优先判断
    if '|' in line:
        return True

    return False


def create_word_table(doc, table_lines):
    """从Markdown表格行创建Word表格"""

    if len(table_lines) < 2:
        return

    # 解析表格数据
    rows_data = []
    header_row = None

    for i, line in enumerate(table_lines):
        # 跳过分隔行（包含横线的行）
        if is_separator_line(line):
            continue

        # 解析单元格
        cells = parse_table_row(line)
        if cells:
            if header_row is None:
                header_row = cells
            else:
                rows_data.append(cells)

    if not header_row:
        return

    # 确定列数
    max_cols = len(header_row)
    for row in rows_data:
        max_cols = max(max_cols, len(row))

    # 创建Word表格
    total_rows = 1 + len(rows_data)  # 标题行 + 数据行
    table = doc.add_table(rows=total_rows, cols=max_cols)

    # 获取表格配置
    config = get_config()
    table_config = config.get('table', {})
    border_enabled = table_config.get('border_enabled', True)
    border_color = table_config.get('border_color', '#000000')
    border_width = table_config.get('border_width', 4)
    row_height_cm = table_config.get('row_height_cm', 0.8)
    alignment_str = table_config.get('alignment', 'center')
    line_spacing = table_config.get('line_spacing', 1.2)
    cell_margin = table_config.get('cell_margin', {})
    vertical_align_str = table_config.get('vertical_align', 'center')

    # 设置表格对齐方式
    alignment_map = {
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'right': WD_TABLE_ALIGNMENT.RIGHT
    }
    table.alignment = alignment_map.get(alignment_str.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # 设置垂直对齐
    vertical_align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'center': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    vertical_align = vertical_align_map.get(vertical_align_str.lower(), WD_ALIGN_VERTICAL.CENTER)

    # 统一设置边框和内边距、行高等
    if border_enabled:
        try:
            tbl = table._tbl
            color = border_color.lstrip('#')
            borders_xml = f'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            tbl.tblPr.append(parse_xml(borders_xml))
        except Exception:
            pass

    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = f'''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass

    # 行高与段落行距统一
    try:
        for row in table.rows:
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                cell.vertical_alignment = vertical_align
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = line_spacing
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass

    # 获取表头背景色配置
    header_bg_color = config.get('table.header', {}).get('background_color')

    # 填充标题行
    header_cells = table.rows[0].cells
    for j, cell_text in enumerate(header_row):
        if j < len(header_cells):
            cell = header_cells[j]
            # 处理表格单元格中的格式
            if contains_markdown_formatting(cell_text.strip()):
                parse_table_cell_formatting(cell, cell_text.strip(), is_header=True)
            else:
                # 导入 convert_quotes_to_chinese 避免循环导入
                from formatter import convert_quotes_to_chinese
                cell.text = convert_quotes_to_chinese(cell_text.strip())
                set_table_cell_format(cell, is_header=True)
            # 应用表头背景色
            if header_bg_color:
                set_cell_background_color(cell, header_bg_color)

    # 获取交替行颜色配置
    row_even_color = config.get('table.row_even', {}).get('background_color')
    row_odd_color = config.get('table.row_odd', {}).get('background_color')

    # 填充数据行
    for i, row_data in enumerate(rows_data):
        if i + 1 < len(table.rows):
            row_cells = table.rows[i + 1].cells
            # 确定当前行颜色（奇偶交替）
            row_bg_color = row_odd_color if i % 2 == 0 else row_even_color
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    # 处理表格单元格中的格式
                    if contains_markdown_formatting(cell_text.strip()):
                        parse_table_cell_formatting(cell, cell_text.strip(), is_header=False)
                    else:
                        # 导入 convert_quotes_to_chinese 避免循环导入
                        from formatter import convert_quotes_to_chinese
                        cell.text = convert_quotes_to_chinese(cell_text.strip())
                        set_table_cell_format(cell, is_header=False)
                    # 应用交替行背景色
                    if row_bg_color:
                        set_cell_background_color(cell, row_bg_color)

    # 调整列宽
    adjust_table_column_width(table)


def parse_table_row(line):
    """解析表格行，提取单元格内容"""
    if not line or not line.strip():
        return []

    line = line.strip()

    # 移除开头和结尾的 |
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]

    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]

    # 过滤掉空单元格（但保留有意义的空单元格）
    return cells


def contains_markdown_formatting(text):
    """检查文本是否包含Markdown格式标记"""
    format_patterns = [
        r'\*\*\*.*?\*\*\*',  # 加粗斜体
        r'\*\*.*?\*\*',      # 加粗
        r'\*.*?\*',          # 斜体
        r'___.*?___',        # 加粗斜体
        r'__.*?__',          # 加粗
        r'_.*?_',            # 斜体
        r'<u>.*?</u>',       # 下划线
        r'~~.*?~~',          # 删除线
        r'`.*?`',            # 行内代码
        r'<br\s*/?>',       # 换行标签
        r'\$.*?\$',         # LaTeX数学公式
    ]

    for pattern in format_patterns:
        if re.search(pattern, text):
            return True
    return False


def parse_table_cell_formatting(cell, text, is_header=False):
    """解析表格单元格中的格式化文本"""
    # 清空单元格
    cell.text = ""

    # 导入 convert_quotes_to_chinese 和 parse_formatted_text 避免循环导入
    from formatter import convert_quotes_to_chinese, parse_formatted_text

    # 转换引号
    text = convert_quotes_to_chinese(text)

    # 支持<br>换行：拆分后逐段处理
    parts_by_br = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)

    # 解析格式
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]

    for idx, segment in enumerate(parts_by_br):
        if idx > 0:
            cell.paragraphs[0].add_run().add_break()
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = cell.paragraphs[0].add_run(part_text)
                set_table_run_format(run, formats, is_header)


def set_table_run_format(run, formats, is_header=False):
    """设置表格单元格run格式"""
    config = get_config()

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.size = Pt(font_size)
    font.color.rgb = hex_to_rgb(color_hex)
    font.bold = bold if is_header else False

    # 设置字体映射：英文和数字用Times New Roman，中文用配置的字体
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

    # 应用Markdown格式
    if formats.get('bold', False):
        font.bold = True
    if formats.get('italic', False):
        font.italic = True
    if formats.get('underline', False):
        font.underline = True
    if formats.get('strikethrough', False):
        font.strike = True
    if formats.get('code', False):
        # 表格中代码使用Times New Roman，稍小字号
        code_config = config.get('inline_code', {})
        font.name = code_config.get('font', 'Times New Roman')
        font.size = Pt(9)
        font.color.rgb = hex_to_rgb(code_config.get('color', '#333333'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return
    if formats.get('math', False):
        # 表格中数学公式使用Times New Roman，斜体，深蓝色
        math_config = config.get('math', {})
        font.name = math_config.get('font', 'Times New Roman')
        font.size = Pt(math_config.get('size', 10))
        font.italic = math_config.get('italic', True)
        font.color.rgb = hex_to_rgb(math_config.get('color', '#00008B'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return


def set_table_cell_format(cell, is_header=False):
    """设置表格单元格格式"""
    config = get_config()
    table_config = config.get('table', {})
    line_spacing = table_config.get('line_spacing', 1.2)

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    # 设置段落格式
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing

        # 设置文字格式
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            font.color.rgb = hex_to_rgb(color_hex)
            font.bold = bold if is_header else False

            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def adjust_table_column_width(table):
    """调整表格列宽"""
    try:
        # 获取表格总宽度（页面宽度减去页边距）
        available_width = Cm(21.0 - 3.18 * 2)  # A4宽度减去左右页边距

        # 平均分配列宽
        col_count = len(table.columns)
        if col_count > 0:
            col_width = int(available_width / col_count)  # 转换为整数
            for column in table.columns:
                column.width = col_width
    except Exception as e:
        print(f"⚠️  表格列宽调整失败: {e}")


def parse_html_table(html_content):
    """解析HTML表格内容，返回表格数据"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        table = soup.find('table')
        if not table:
            return None

        rows_data = []
        for tr in table.find_all('tr'):
            row_cells = []
            for cell in tr.find_all(['td', 'th']):
                # 获取单元格文本内容，保留基本格式
                cell_text = cell.get_text(strip=True)
                row_cells.append(cell_text)
            if row_cells:  # 只添加非空行
                rows_data.append(row_cells)

        return rows_data
    except Exception as e:
        print(f"⚠️  HTML表格解析失败: {e}")
        return None


def create_word_table_from_html(doc, html_content):
    """从HTML表格创建Word表格"""
    rows_data = parse_html_table(html_content)
    if not rows_data or len(rows_data) < 1:
        print("⚠️  HTML表格数据为空或格式不正确")
        return

    # 导入 convert_quotes_to_chinese 避免循环导入
    from formatter import convert_quotes_to_chinese

    # 获取表格配置
    config = get_config()
    table_config = config.get('table', {})
    border_enabled = table_config.get('border_enabled', True)
    border_color = table_config.get('border_color', '#000000')
    border_width = table_config.get('border_width', 4)
    row_height_cm = table_config.get('row_height_cm', 0.8)
    line_spacing = table_config.get('line_spacing', 1.2)
    cell_margin = table_config.get('cell_margin', {})
    vertical_align_str = table_config.get('vertical_align', 'center')

    # 创建Word表格
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))

    # 设置表格对齐方式
    alignment_str = table_config.get('alignment', 'center')
    alignment_map = {
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'right': WD_TABLE_ALIGNMENT.RIGHT
    }
    table.alignment = alignment_map.get(alignment_str.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # 设置垂直对齐
    vertical_align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'center': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    vertical_align = vertical_align_map.get(vertical_align_str.lower(), WD_ALIGN_VERTICAL.CENTER)

    # 设置表格边框和单元格边距
    if border_enabled:
        try:
            tbl = table._tbl
            color = border_color.lstrip('#')
            borders_xml = f'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            tbl.tblPr.append(parse_xml(borders_xml))
        except Exception:
            pass

    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = f'''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass

    # 设置行高和单元格对齐
    try:
        for row in table.rows:
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                cell.vertical_alignment = vertical_align
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = line_spacing
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass

    # 填充表格数据
    for i, row_data in enumerate(rows_data):
        if i < len(table.rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_cells):
                if j < len(row_cells):
                    cell = row_cells[j]
                    cell.text = convert_quotes_to_chinese(cell_text.strip())
                    # 第一行作为标题行处理
                    set_table_cell_format(cell, is_header=(i == 0))

    # 调整列宽
    adjust_table_column_width(table)
    print(f"✅ 处理HTML表格: {len(rows_data)} 行")


def hex_to_rgb(hex_color: str):
    """将十六进制颜色转换为 RGBColor"""
    from docx.shared import RGBColor
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 6:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    return RGBColor(0, 0, 0)  # 默认黑色
