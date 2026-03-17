# 基金日报自动生成脚本（整合新闻版）
# 每天下午 4:00 运行

import requests
import json
import time
import re
import subprocess
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ============ 中国法定节假日配置（2026 年）============
# 格式：'YYYY-MM-DD'，包含所有调休的工作日和法定节假日
CHINESE_HOLIDAYS_2026 = {
    # 元旦：1 月 1 日放假，1 月 2 日（周五）补休
    '2026-01-01': 'holiday',  # 元旦
    '2026-01-02': 'holiday',  # 元旦补休
    
    # 春节：2 月 17 日 -23 日放假调休
    '2026-02-17': 'holiday',  # 春节
    '2026-02-18': 'holiday',
    '2026-02-19': 'holiday',
    '2026-02-20': 'holiday',
    '2026-02-21': 'holiday',
    '2026-02-22': 'holiday',
    '2026-02-23': 'holiday',
    '2026-02-15': 'workday',  # 周日上班
    
    # 清明节：4 月 5 日 -7 日放假
    '2026-04-05': 'holiday',
    '2026-04-06': 'holiday',
    '2026-04-07': 'holiday',
    '2026-04-05': 'holiday',  # 周日
    
    # 劳动节：5 月 1 日 -5 日放假
    '2026-05-01': 'holiday',
    '2026-05-02': 'holiday',
    '2026-05-03': 'holiday',
    '2026-05-04': 'holiday',
    '2026-05-05': 'holiday',
    '2026-05-09': 'workday',  # 周六上班
    
    # 端午节：6 月 19 日 -21 日放假
    '2026-06-19': 'holiday',
    '2026-06-20': 'holiday',
    '2026-06-21': 'holiday',
    
    # 中秋节：9 月 25 日 -27 日放假
    '2026-09-25': 'holiday',
    '2026-09-26': 'holiday',
    '2026-09-27': 'holiday',
    '2026-09-20': 'workday',  # 周日上班
    
    # 国庆节：10 月 1 日 -7 日放假
    '2026-10-01': 'holiday',
    '2026-10-02': 'holiday',
    '2026-10-03': 'holiday',
    '2026-10-04': 'holiday',
    '2026-10-05': 'holiday',
    '2026-10-06': 'holiday',
    '2026-10-07': 'holiday',
    '2026-10-10': 'workday',  # 周六上班
}

# ============ 配置 ============
FUND_CODES = ['001407', '017091', '050025']  # 基金代码
FUND_NAMES = {
    '001407': '景顺长城稳健回报混合 C',
    '017091': '景顺长城纳斯达克科技 ETF 联接 (QDII)A',
    '050025': '博时标普 500ETF 联接 (QDII)A'
}

# Telegram 配置
TELEGRAM_BOT_TOKEN = "8599058765:AAH_4PW048zLJxGQ2BosnG2uh3bRwAeu8XA"
TELEGRAM_CHAT_ID = "7333732220"  # Joshua 的 Chat ID
# =================================

def is_trading_day(date):
    """
    判断指定日期是否为交易日
    返回：(是否交易日，说明)
    """
    # 检查周末
    if date.weekday() >= 5:  # 周六=5，周日=6
        return False, "周末"
    
    # 检查法定节假日配置
    date_str = date.strftime('%Y-%m-%d')
    if date_str in CHINESE_HOLIDAYS_2026:
        holiday_type = CHINESE_HOLIDAYS_2026[date_str]
        if holiday_type == 'holiday':
            return False, "法定节假日"
        elif holiday_type == 'workday':
            return True, "调休工作日"
    
    return True, "交易日"

def get_last_trading_day(from_date=None):
    """
    获取从指定日期往前的最近一个交易日
    """
    if from_date is None:
        from_date = datetime.now()
    
    current = from_date
    days_checked = 0
    
    while days_checked < 10:  # 最多往前查 10 天
        is_trading, reason = is_trading_day(current)
        if is_trading:
            return current, reason
        
        current -= timedelta(days=1)
        days_checked += 1
    
    # 如果找不到，返回原日期
    return from_date, "未知"

def get_fund_report_date():
    """
    获取基金报告应该使用的数据日期
    规则：
    1. 如果是交易日且 >= 20:00，用今天
    2. 如果是交易日且 < 20:00，用上一个交易日
    3. 如果是非交易日，用上一个交易日
    """
    now = datetime.now()
    today = now.date()
    current_hour = now.hour
    
    # 判断今天是否为交易日
    is_today_trading, reason = is_trading_day(today)
    
    if is_today_trading:
        # 今天是交易日
        if current_hour >= 20:
            # 已过 20 点，净值已公布
            return today, "净值已公布", "今日"
        else:
            # 未到 20 点，用上一个交易日的数据
            last_trading, _ = get_last_trading_day(today - timedelta(days=1))
            return last_trading, "净值未公布", "上一交易日"
    else:
        # 今天非交易日，用上一个交易日的数据
        last_trading, _ = get_last_trading_day(today - timedelta(days=1))
        return last_trading, "非交易日", "上一交易日"

def get_fund_data(fund_code, report_date=None):
    """获取基金最新数据（从天天基金网 API）"""
    try:
        # 天天基金网实时 API（获取估值和最新净值）
        url = f"http://fundgz.1234567.com.cn/js/{fund_code}.js"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Referer': 'http://fund.eastmoney.com/'
        }
        response = requests.get(url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            # 解析 JSONP 数据
            content = response.text
            json_str = content[8:-2]  # 去掉 jsonpgz(...) 包装
            data = json.loads(json_str)
            
            # 判断是估值还是净值
            gztime = data.get('gztime', '')
            jjzcrq = data.get('jjzcrq', '')
            
            # 使用传入的报告日期（如果有）
            if report_date:
                nav_date = report_date.strftime('%Y-%m-%d')
            else:
                nav_date = jjzcrq if jjzcrq else gztime[:10]
            
            # 判断当前时间状态
            now = datetime.now()
            is_trading, _ = is_trading_day(now.date())
            
            if is_trading and now.hour >= 15 and now.hour < 20:
                # 交易日 15:00-20:00，交易已结束但净值未公布
                data_status = '交易中'
                change = float(data.get('gszzl', 0))
            elif is_trading and now.hour < 15:
                # 交易日 15:00 前，显示估值
                data_status = '估值'
                change = float(data.get('gszzl', 0))
            else:
                # 非交易日或 20:00 后，显示净值
                data_status = '净值'
                change = 0  # 净值不显示涨跌，只显示净值
            
            return {
                'code': fund_code,
                'name': data.get('name', FUND_NAMES.get(fund_code, '')),
                '估值涨跌': float(data.get('gszzl', 0)),
                '单位净值': float(data.get('dwjz', 0)),
                '累计净值': float(data.get('ljjz', 0)),
                '净值日期': nav_date,
                '数据状态': data_status
            }
    except Exception as e:
        print(f"[ERROR] 获取基金 {fund_code} 数据失败：{e}")
    
    return None

def get_fund_history_returns(fund_code):
    """获取基金历史收益率（通过获取历史净值计算）"""
    try:
        # 获取历史净值数据
        end_date = datetime.now()
        start_date = end_date - timedelta(days=180)
        
        url = f"http://api.fund.eastmoney.com/f10/lsjz"
        params = {
            'fundCode': fund_code,
            'pageIndex': 1,
            'pageSize': 180,
            'startDate': start_date.strftime('%Y%m%d'),
            'endDate': end_date.strftime('%Y%m%d')
        }
        headers = {
            'User-Agent': 'Mozilla/5.0',
            'Referer': f'http://fund.eastmoney.com/{fund_code}.html'
        }
        
        response = requests.get(url, headers=headers, params=params, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('Data') and data['Data'].get('Data'):
                nav_list = data['Data']['Data']
                
                if len(nav_list) > 0:
                    latest_nav = float(nav_list[0]['DWJZ'])
                    returns = {}
                    
                    if len(nav_list) >= 30:
                        month_nav = float(nav_list[30]['DWJZ'])
                        returns['1 个月'] = ((latest_nav - month_nav) / month_nav) * 100
                    else:
                        returns['1 个月'] = 0.0
                    
                    if len(nav_list) >= 90:
                        quarter_nav = float(nav_list[90]['DWJZ'])
                        returns['3 个月'] = ((latest_nav - quarter_nav) / quarter_nav) * 100
                    else:
                        returns['3 个月'] = 0.0
                    
                    if len(nav_list) >= 180:
                        half_year_nav = float(nav_list[180]['DWJZ'])
                        returns['半年'] = ((latest_nav - half_year_nav) / half_year_nav) * 100
                    else:
                        returns['半年'] = 0.0
                    
                    return returns
    except Exception as e:
        print(f"  [WARN] 获取基金 {fund_code} 历史数据失败：{e}")
    
    return {'1 个月': 0.0, '3 个月': 0.0, '半年': 0.0}

def get_finance_news(limit=10):
    """获取财经新闻（使用 news-market 技能）"""
    try:
        script_path = r"C:\Users\ZhuanZ\.openclaw\workspace\skills\news-market\scripts\news_market.py"
        result = subprocess.run(
            ['python', script_path, 'category', '--cat', 'securities', '--limit', str(limit)],
            capture_output=True,
            text=True,
            timeout=30,
            encoding='utf-8'
        )
        
        if result.returncode == 0:
            # 解析输出
            news_list = []
            lines = result.stdout.split('\n')
            
            current_news = {}
            for line in lines:
                line = line.strip()
                if line.startswith('📰'):
                    continue
                elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                    if current_news:
                        news_list.append(current_news)
                    current_news = {'title': line.split('. ', 1)[1] if '. ' in line else line}
                elif line.startswith('📌') and current_news:
                    current_news['source'] = line.replace('📌 ', '')
                elif line.startswith('🔗') and current_news:
                    current_news['url'] = line.replace('🔗 ', '')
                elif line.startswith('📝') and current_news:
                    current_news['summary'] = line.replace('📝 ', '')
            
            if current_news:
                news_list.append(current_news)
            
            return news_list[:limit]
    except Exception as e:
        print(f"[WARN] 获取新闻失败：{e}")
    
    return []

def send_telegram_message(bot_token, chat_id, message):
    """发送 Telegram 消息"""
    try:
        url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
        data = {
            "chat_id": chat_id,
            "text": message,
            "parse_mode": "Markdown"
        }
        response = requests.post(url, json=data, timeout=30)
        return response.status_code == 200
    except Exception as e:
        print(f"[ERROR] Telegram 发送失败：{e}")
        return False

def send_telegram_file(bot_token, chat_id, file_path, caption):
    """发送 Telegram 文件"""
    try:
        url = f"https://api.telegram.org/bot{bot_token}/sendDocument"
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {
                'chat_id': chat_id,
                'caption': caption
            }
            response = requests.post(url, files=files, data=data, timeout=60)
            return response.status_code == 200
    except Exception as e:
        print(f"[ERROR] Telegram 文件发送失败：{e}")
        return False

def create_fund_report(fund_data_list, news_list=None, report_date=None):
    """创建基金日报 Word 文档（包含新闻）"""
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('基金日报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期
    today = datetime.now()
    date_str = today.strftime('%Y 年 %m 月 %d 日 %H:%M')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'生成时间：{date_str}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 数据日期说明
    if report_date:
        nav_date_str = report_date.strftime('%Y-%m-%d')
        data_info = doc.add_paragraph()
        data_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = data_info.add_run(f'📅 数据日期：{nav_date_str}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)
    
    doc.add_paragraph()
    
    # 基金数据表格
    doc.add_heading('📊 基金数据概览', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '基金名称'
    hdr_cells[1].text = '基金代码'
    hdr_cells[2].text = '今日涨跌'
    hdr_cells[3].text = '单位净值'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for fund in fund_data_list:
        row_cells = table.add_row().cells
        row_cells[0].text = fund['name']
        row_cells[1].text = fund['code']
        change = fund['估值涨跌']
        row_cells[2].text = f"{change:+.2f}%"
        row_cells[3].text = f"{fund['单位净值']:.4f}"
    
    doc.add_paragraph()
    
    # 数据说明
    doc.add_paragraph('注：交易日 15:00 前显示估值数据，20:00 后显示实际净值数据。', style='Intense Quote')
    doc.add_paragraph()
    
    # 财经新闻
    doc.add_heading('📰 财经要闻', level=1)
    
    if news_list:
        for i, news in enumerate(news_list[:10], 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {news.get("title", "")}\n').bold = True
            p.add_run(f'   来源：{news.get("source", "")}\n')
            p.add_run(f'   链接：{news.get("url", "")}')
    else:
        doc.add_paragraph('暂无新闻数据')
    
    doc.add_paragraph()
    
    # 总结
    doc.add_heading('📝 总结', level=1)
    
    avg_change = sum(f.get('估值涨跌', 0) for f in fund_data_list) / len(fund_data_list)
    
    summary = doc.add_paragraph()
    summary.add_run('最新表现：').bold = True
    if avg_change > 0:
        summary.add_run(f'整体上涨，平均涨幅 {avg_change:.2f}%')
    elif avg_change < 0:
        summary.add_run(f'整体下跌，平均跌幅 {abs(avg_change):.2f}%')
    else:
        summary.add_run('基本持平')
    
    doc.add_paragraph()
    doc.add_paragraph('注：基金数据来自天天基金网，新闻来自东方财富、雪球等财经媒体。')
    
    # 保存文档
    base_dir = r"D:\System\Desktop\基金日报"
    folder_name = f"{today.strftime('%Y%m%d')}_{today.strftime('%m月%d日')}基金日报"
    output_dir = os.path.join(base_dir, folder_name)
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    filename = f"基金日报_{today.strftime('%Y%m%d')}.docx"
    output_path = os.path.join(output_dir, filename)
    
    doc.save(output_path)
    return output_path

def create_telegram_message(fund_data_list, news_list=None, report_date_info=None):
    """创建 Telegram 消息（包含新闻）"""
    today = datetime.now()
    date_str = today.strftime('%Y 年 %m 月 %d 日')
    
    # 报告日期信息
    if report_date_info:
        nav_date, status_reason, date_type = report_date_info
        nav_date_str = nav_date.strftime('%Y-%m-%d')
    else:
        nav_date_str = "未知"
        status_reason = ""
        date_type = ""
    
    message = f"""
━━━━━━━━━━━━━━━━━━
   📊 最新基金日报
   {date_str}
━━━━━━━━━━━━━━━━━━
📅 数据日期：{nav_date_str} ({status_reason})

"""
    
    for fund in fund_data_list:
        nav_date = fund.get('净值日期', '')[:10]
        nav = fund.get('单位净值', 0)
        change = fund.get('估值涨跌', 0)
        status = fund.get('数据状态', '未知')
        
        message += f"【{fund['name']}】\n"
        message += f"├ 代码：{fund['code']}\n"
        
        if status == '估值' or status == '交易中':
            message += f"├ 估值涨跌：{change:+.2f}%\n"
        else:
            message += f"├ 日涨跌：{change:+.2f}%\n"
        
        message += f"└ 单位净值：{nav:.4f}\n\n"
    
    message += "━━━━━━━━━━━━━━━━━━\n"
    message += "📰 财经要闻\n"
    message += "━━━━━━━━━━━━━━━━━━\n\n"
    
    if news_list:
        for i, news in enumerate(news_list[:5], 1):
            title = news.get('title', '无标题')
            source = news.get('source', '未知')
            message += f"{i}. {title}\n"
            message += f"   📌 {source}\n"
            message += f"   🔗 {news.get('url', '')}\n\n"
    else:
        message += "暂无新闻数据\n\n"
    
    message += "━━━━━━━━━━━━━━━━━━\n"
    message += "注：交易日 15:00 前显示估值，20:00 后显示净值\n"
    message += "━━━━━━━━━━━━━━━━━━\n"
    
    return message

if __name__ == "__main__":
    print("=" * 60)
    print("基金日报自动生成（整合新闻版）")
    print("=" * 60)
    print(f"开始时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()
    
    # 获取报告日期信息
    print("[0/6] 正在判断数据日期...")
    report_date, status_reason, date_type = get_fund_report_date()
    print(f"    报告数据日期：{report_date.strftime('%Y-%m-%d')}")
    print(f"    状态：{status_reason} ({date_type})")
    print()
    
    # 获取基金数据
    print("[1/6] 正在获取基金数据...")
    fund_data_list = []
    
    for code in FUND_CODES:
        print(f"  获取 {code} - {FUND_NAMES[code]}...")
        data = get_fund_data(code, report_date)
        if data:
            fund_data_list.append(data)
            nav_date = data.get('净值日期', '')[:10]
            status = data.get('数据状态', '未知')
            change = data.get('估值涨跌', 0)
            print(f"    [OK] {data['name']}")
            print(f"         日期：{nav_date} ({status}) | 涨跌：{change:+.2f}% | 净值：{data['单位净值']:.4f}")
        time.sleep(1)
    
    if not fund_data_list:
        print("[ERROR] 未能获取任何基金数据")
        exit(1)
    
    print(f"\n[OK] 成功获取 {len(fund_data_list)} 只基金数据")
    print()
    
    # 获取财经新闻
    print("[2/6] 正在获取财经新闻...")
    news_list = get_finance_news(limit=10)
    print(f"[OK] 获取到 {len(news_list)} 条新闻")
    for i, news in enumerate(news_list[:3], 1):
        print(f"  {i}. {news.get('title', '')[:50]}...")
    print()
    
    # 创建 Word 报告
    print("[3/6] 正在生成 Word 报告...")
    output_path = create_fund_report(fund_data_list, news_list, report_date)
    print(f"[OK] 报告已保存：{output_path}")
    print()
    
    # 发送 Telegram 消息
    print("[4/6] 正在发送 Telegram 消息...")
    message = create_telegram_message(fund_data_list, news_list, (report_date, status_reason, date_type))
    send_telegram_message(TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID, message)
    print("[OK] Telegram 消息已发送")
    print()
    
    # 发送 Word 文档
    print("[5/6] 正在发送 Word 文档...")
    time.sleep(2)
    caption = f"📊 基金日报 - {datetime.now().strftime('%Y-%m-%d')}"
    send_telegram_file(TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID, output_path, caption)
    print("[OK] Word 文档已发送")
    print()
    
    # 打印总结
    print("[6/6] 生成总结")
    print(f"    数据日期：{report_date.strftime('%Y-%m-%d')}")
    print(f"    状态说明：{status_reason}")
    print(f"    下次更新：明天 16:00（如果明天是交易日）")
    print()
    
    print("=" * 60)
    print("完成！")
    print(f"结束时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)
