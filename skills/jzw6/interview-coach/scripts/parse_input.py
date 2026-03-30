#!/usr/bin/env python3
"""
解析简历或 JD 文件，输出纯文本到 stdout。

支持格式：PDF、DOCX、DOC、JPG/PNG/JPEG/WEBP、TXT
OCR 使用 PaddleOCR AIStudio 云端 API（仅图片/扫描版 PDF 需要 token）

用法：
  python parse_input.py <文件路径>

环境变量（可在脚本同目录的 .env 文件中配置，也可直接设置系统环境变量）：
  PADDLEOCR_TOKEN       AIStudio Access Token（图片/扫描PDF时必填）
  PADDLEOCR_SERVER_URL  PaddleOCR 服务地址（可选，默认使用官方地址）

获取 Token：https://aistudio.baidu.com/paddleocr → 右上角 → Access Token
"""

import sys
import os
import base64

# ── 配置 ──────────────────────────────────────────────────────────────────────

def _load_env() -> None:
    """
    从脚本同目录的 .env 文件加载配置。
    已存在的环境变量不会被覆盖，安全地合并配置。
    """
    # 同时支持 scripts/.env 和上一级目录 .env（兼容不同部署结构）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(script_dir, ".env"),
        os.path.join(script_dir, "..", ".env"),
    ]
    for env_path in candidates:
        env_path = os.path.normpath(env_path)
        if not os.path.exists(env_path):
            continue
        with open(env_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, _, val = line.partition("=")
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                if key and val and key not in os.environ:
                    os.environ[key] = val
        break  # 找到第一个 .env 就停止

_load_env()

TOKEN      = os.getenv("PADDLEOCR_TOKEN", "").strip()
SERVER_URL = os.getenv(
    "PADDLEOCR_SERVER_URL",
    "https://g4hcg732gdjfofu1.aistudio-app.com"
).rstrip("/")

STRUCT_ENDPOINT = f"{SERVER_URL}/layout-parsing"

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
PDF_EXT    = ".pdf"
DOCX_EXTS  = {".docx", ".doc"}
TXT_EXT    = ".txt"

# ── 依赖安装 ──────────────────────────────────────────────────────────────────

def _install(pkg: str) -> None:
    """尝试用 uv 安装，不存在时回退到 pip。"""
    import subprocess
    print(f"[安装依赖] {pkg} ...", file=sys.stderr)
    try:
        subprocess.check_call(
            ["uv", "pip", "install", pkg, "--system", "-q"],
            stderr=subprocess.DEVNULL,
        )
    except FileNotFoundError:
        print("[安装依赖] uv 未找到，改用 pip ...", file=sys.stderr)
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", pkg, "-q"],
                stderr=subprocess.DEVNULL,
            )
        except subprocess.CalledProcessError:
            raise RuntimeError(
                f"自动安装 {pkg} 失败，请手动执行：pip install {pkg}"
            )

# ── OCR：图片 ─────────────────────────────────────────────────────────────────

def ocr_image(file_path: str) -> str:
    """调用 PaddleOCR AIStudio API 识别图片文字。"""
    _require_token("图片 OCR")

    try:
        import requests
    except ImportError:
        _install("requests")
        import requests

    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    payload = {"file": b64, "fileType": 1}
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"token {TOKEN}",
    }

    resp = _post_with_error_handling(requests, payload, headers, timeout=60)
    return _extract_text_from_response(resp)


# ── OCR：扫描版 PDF ───────────────────────────────────────────────────────────

def _parse_pdf_via_api(file_path: str) -> str:
    """调用 PaddleOCR PP-StructureV3 API 解析扫描版 PDF。"""
    _require_token("扫描版 PDF 解析")

    try:
        import requests
    except ImportError:
        _install("requests")
        import requests

    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    payload = {"file": b64, "fileType": 0}
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"token {TOKEN}",
    }

    resp = _post_with_error_handling(requests, payload, headers, timeout=120)
    return _extract_text_from_response(resp)


def _parse_pdf_local(file_path: str) -> str:
    """用 pdfplumber 提取原生文本 PDF（无需 token，速度快）。"""
    try:
        import pdfplumber
    except ImportError:
        _install("pdfplumber")
        import pdfplumber

    with pdfplumber.open(file_path) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def parse_pdf(file_path: str) -> str:
    """
    PDF 解析策略：
    1. 先用 pdfplumber 提取原生文本（快、无需 token）
    2. 文本过少（可能是扫描件）→ 提示用户需要 token，改用 API OCR
    """
    text = _parse_pdf_local(file_path)
    if len(text.strip()) >= 300:
        return text

    # 文本不足，判断是否为扫描件
    if not TOKEN:
        raise EnvironmentError(
            "此 PDF 可能是扫描件（提取到的文字过少），需要使用 OCR 解析。\n"
            "请配置 PADDLEOCR_TOKEN 后重试：\n"
            "  1. 访问 https://aistudio.baidu.com/paddleocr 获取 Access Token\n"
            "  2. 在 .env 文件中添加：PADDLEOCR_TOKEN=你的token\n"
            "  或者：直接粘贴简历/JD 文字内容，无需 token。"
        )
    print("[解析] PDF 原生文本不足，改用 PaddleOCR API 解析...", file=sys.stderr)
    return _parse_pdf_via_api(file_path)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        _install("python-docx")
        from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 同时提取表格中的文字（简历常用表格排版）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


# ── TXT ───────────────────────────────────────────────────────────────────────

def parse_txt(file_path: str) -> str:
    # 自动检测编码，兼容 GBK/UTF-8
    encodings = ["utf-8", "gbk", "utf-16"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # 全部失败则忽略非法字节
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def _require_token(feature: str) -> None:
    """检查 token 是否已配置，未配置时给出清晰的操作指引。"""
    if not TOKEN:
        raise EnvironmentError(
            f"{feature} 需要 PaddleOCR Token，但当前未配置。\n"
            "配置方法：\n"
            "  1. 访问 https://aistudio.baidu.com/paddleocr 获取 Access Token\n"
            "  2. 在 skill 目录的 .env 文件中添加：PADDLEOCR_TOKEN=你的token\n"
            "     或设置系统环境变量：export PADDLEOCR_TOKEN=你的token\n"
            "\n"
            "如果只有文本 PDF / Word / TXT，无需 token，直接使用即可。"
        )


def _post_with_error_handling(requests_mod, payload, headers, timeout: int):
    """统一处理 API 请求错误，给出友好提示。"""
    try:
        resp = requests_mod.post(
            STRUCT_ENDPOINT, json=payload, headers=headers, timeout=timeout
        )
        resp.raise_for_status()
        return resp
    except requests_mod.exceptions.Timeout:
        raise RuntimeError(
            f"OCR 请求超时（>{timeout}s），请检查网络连接后重试。\n"
            "也可以直接粘贴文字内容，跳过文件解析。"
        )
    except requests_mod.exceptions.HTTPError as e:
        status = e.response.status_code
        if status == 401:
            raise RuntimeError(
                "Token 无效或已过期（HTTP 401）。\n"
                "请到 https://aistudio.baidu.com/paddleocr 重新获取 Token。"
            )
        raise RuntimeError(
            f"OCR API 返回错误（HTTP {status}）。\n"
            "请检查 PADDLEOCR_TOKEN 是否正确，或稍后重试。"
        )
    except requests_mod.exceptions.ConnectionError:
        raise RuntimeError(
            "无法连接到 OCR 服务，请检查网络连接。\n"
            "也可以直接粘贴文字内容，跳过文件解析。"
        )


def _extract_text_from_response(resp) -> str:
    """从 PaddleOCR API 响应中提取文字内容。"""
    data = resp.json()
    parts = []
    for res in data.get("result", {}).get("layoutParsingResults", []):
        text = res.get("markdown", {}).get("text", "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


# ── 入口 ──────────────────────────────────────────────────────────────────────

PARSERS = {
    **{ext: ocr_image  for ext in IMAGE_EXTS},
    PDF_EXT: parse_pdf,
    **{ext: parse_docx for ext in DOCX_EXTS},
    TXT_EXT: parse_txt,
}

SUPPORTED_EXTS = ", ".join(sorted(PARSERS.keys()))


def main() -> None:
    if len(sys.argv) < 2:
        print(f"用法：python3 parse_input.py <文件路径>", file=sys.stderr)
        print(f"支持格式：{SUPPORTED_EXTS}", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"错误：文件不存在 → {file_path}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(file_path):
        print(f"错误：路径不是文件 → {file_path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(file_path)[1].lower()
    parser = PARSERS.get(ext)

    if not parser:
        print(
            f"错误：不支持的文件格式 {ext}\n"
            f"支持的格式：{SUPPORTED_EXTS}",
            file=sys.stderr,
        )
        sys.exit(1)

    try:
        text = parser(file_path)
    except EnvironmentError as e:
        # Token 未配置等配置问题，给出明确指引
        print(f"\n⚠️  配置问题：\n{e}", file=sys.stderr)
        sys.exit(2)  # 用不同退出码区分"配置问题"和"解析失败"
    except RuntimeError as e:
        # API 调用失败等运行时问题
        print(f"\n❌  解析失败：\n{e}", file=sys.stderr)
        sys.exit(3)
    except Exception as e:
        print(f"\n❌  未预期的错误：{e}", file=sys.stderr)
        sys.exit(1)

    if not text.strip():
        print(
            "警告：未提取到任何文字内容。\n"
            "请尝试：直接粘贴文字内容，或换用其他格式的文件。",
            file=sys.stderr,
        )
        sys.exit(1)

    # 强制 UTF-8 输出，避免 Windows GBK 乱码（兼容 Python 3.6+）
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    else:
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    print(text)


if __name__ == "__main__":
    main()
