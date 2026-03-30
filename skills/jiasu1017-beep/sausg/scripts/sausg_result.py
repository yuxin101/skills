#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SAUSG 结果读取模块
用于从SAUSG计算结果中提取数据

功能：
- 读取基本周期和频率 (FRQ文件)
- 读取底部反力 (NSF文件)
- 从Word报告中提取基底剪力、剪重比、位移角
"""

import os
import re
from typing import Optional, List, Dict, Any

# 尝试导入python-docx用于读取Word报告
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def read_frq_file(frq_path: str) -> Dict[str, List[str]]:
    """
    读取FRQ文件，获取基本周期和频率

    Args:
        frq_path: FRQ文件路径

    Returns:
        dict: 包含periods, frequencies, freq_hz的字典
    """
    results = {
        "periods": [],
        "frequencies": [],
        "freq_hz": []
    }

    if not os.path.exists(frq_path):
        return results

    # FRQ文件是GBK编码
    for encoding in ['gbk', 'utf-8', 'cp1252', 'latin1']:
        try:
            with open(frq_path, 'r', encoding=encoding) as f:
                lines = f.readlines()
            break
        except Exception:
            continue
    else:
        return results

    # 解析FRQ文件格式：
    # 第1行: 计算结果: Nmode = 10
    # 第2行: 序号  圆频率  频率  周期
    # 后续行: 序号  圆频率(rad/s)  频率(Hz)  周期(s)
    for line in lines:
        # 跳过标题行
        if '周期' in line or '圆频率' in line or 'mode' in line.lower():
            continue
        parts = line.strip().split()
        if len(parts) >= 4:
            try:
                mode_num = int(parts[0])
                omega = float(parts[1])    # 圆频率 rad/s
                freq = float(parts[2])      # 频率 Hz
                period = float(parts[3])    # 周期 s
                if mode_num <= 6:  # 只取前6阶
                    results["periods"].append(f"T{mode_num}={period:.4f}s")
                    results["frequencies"].append(f"ω{mode_num}={omega:.4f}rad/s")
                    results["freq_hz"].append(f"f{mode_num}={freq:.4f}Hz")
            except (ValueError, IndexError):
                continue

    return results


def read_nsf_file(nsf_path: str) -> Dict[str, Any]:
    """
    读取NSF文件，获取底部反力和楼层总重

    Args:
        nsf_path: NSF文件路径

    Returns:
        dict: 包含reactions和total_weight的字典
    """
    results = {
        "reactions": None,
        "total_weight": None
    }

    if not os.path.exists(nsf_path):
        return results

    # NSF文件是GBK编码
    for encoding in ['gbk', 'utf-8', 'cp1252', 'latin1']:
        try:
            with open(nsf_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except Exception:
            continue
    else:
        return results

    # 提取底部反力
    # 格式: Rx = xxx, Ry = xxx, Rz = xxx
    reaction_match = re.search(
        r'Rx\s*=\s*([-+]?\d+\.?\d*[eE]?[+-]?\d*),?\s*Ry\s*=\s*([-+]?\d+\.?\d*[eE]?[+-]?\d*),?\s*Rz\s*=\s*([-+]?\d+\.?\d*[eE]?[+-]?\d*)',
        content
    )
    if reaction_match:
        rx, ry, rz = reaction_match.groups()
        results["reactions"] = f"Rx={rx}kN, Ry={ry}kN, Rz={rz}kN"

    # 提取楼层总重 (Fz)
    weight_match = re.search(r'Fz\s*=\s*([-+]?\d+\.?\d*[eE]?[+-]?\d*)', content)
    if weight_match:
        results["total_weight"] = f"{abs(float(weight_match.group(1))):.2f}kN"

    return results


def read_report_base_shear(report_path: str) -> Dict[str, Dict[str, float]]:
    """
    从Word报告中提取基底剪力和剪重比

    Args:
        report_path: 报告文件路径

    Returns:
        dict: {"X": {"shear": xxx, "ratio": xxx}, "Y": {...}} 或空字典
    """
    if not HAS_DOCX or not os.path.exists(report_path):
        return {}

    results = {}

    try:
        doc = Document(report_path)

        # 遍历表格查找表6.5-1 (基底剪力表格, 索引8)
        for i, table in enumerate(doc.tables):
            if i == 8 or i == 9:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    # 查找包含RH1的行
                    if any('RH1' in cell for cell in row_data):
                        full_text = ' '.join(row_data)
                        # 移除RH1TG065文本，再提取数值
                        clean_text = re.sub(r'RH1TG\d+', '', full_text)
                        numbers = re.findall(r'\d+\.\d+', clean_text)
                        if len(numbers) >= 4:
                            try:
                                results["X"] = {"shear": float(numbers[0]), "ratio": float(numbers[1])}
                                results["Y"] = {"shear": float(numbers[2]), "ratio": float(numbers[3])}
                            except (ValueError, IndexError):
                                pass

    except Exception:
        pass

    return results


def read_report_drift_ratio(report_path: str) -> Dict[str, Dict[str, Any]]:
    """
    从Word报告中提取最大层间位移角

    Args:
        report_path: 报告文件路径

    Returns:
        dict: {"X": {"displacement": xxx, "drift": "1/xxx", "layer": x}, "Y": {...}} 或空字典
    """
    if not HAS_DOCX or not os.path.exists(report_path):
        return {}

    results = {}

    try:
        doc = Document(report_path)

        # 遍历表格查找表6.6-1 (位移角表格, 索引11)
        for i, table in enumerate(doc.tables):
            if i == 11 or i == 12:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any('RH1' in cell for cell in row_data):
                        full_text = ' '.join(row_data)
                        # 格式: RH1TG065 | 0.214 | 1/111 | 4 | 0.192 | 1/116 | 4
                        match = re.search(
                            r'(\d+\.?\d*)\s+1/(\d+)\s+(\d+)\s+(\d+\.?\d*)\s+1/(\d+)\s+(\d+)',
                            full_text
                        )
                        if match:
                            try:
                                results["X"] = {
                                    "displacement": float(match.group(1)),
                                    "drift": "1/" + match.group(2),
                                    "layer": int(match.group(3))
                                }
                                results["Y"] = {
                                    "displacement": float(match.group(4)),
                                    "drift": "1/" + match.group(5),
                                    "layer": int(match.group(6))
                                }
                            except (ValueError, IndexError):
                                pass

    except Exception:
        pass

    return results


def read_main_results(model_dir: str, model_name: str = None) -> Dict[str, Any]:
    """
    读取主要计算结果

    Args:
        model_dir: 模型目录路径
        model_name: 模型名称（可选，不带扩展名）

    Returns:
        dict: 包含所有主要结果的字典
    """
    results = {
        "periods": [],
        "frequencies": [],
        "freq_hz": [],
        "reactions": None,
        "total_weight": None,
        "base_shear": {},
        "drift_ratio": {},
        "reports": []
    }

    if not model_name:
        model_name = ""

    # 读取FRQ文件（基本周期和频率）
    static_dir = os.path.join(model_dir, "StaticResult")
    if os.path.isdir(static_dir):
        frq_files = [f for f in os.listdir(static_dir) if f.upper().endswith('.FRQ')]
        if frq_files:
            frq_results = read_frq_file(os.path.join(static_dir, frq_files[0]))
            results.update(frq_results)

        # 读取NSF文件（底部反力）
        nsf_files = [f for f in os.listdir(static_dir) if f.upper().endswith('.NSF')]
        if nsf_files:
            nsf_results = read_nsf_file(os.path.join(static_dir, nsf_files[0]))
            results["reactions"] = nsf_results.get("reactions")
            results["total_weight"] = nsf_results.get("total_weight")

    # 查找计算报告
    docx_files = [f for f in os.listdir(model_dir) if f.upper().endswith('.DOCX')]
    # 过滤掉临时文件
    results["reports"] = [r for r in docx_files if not r.startswith('~$') and not r.startswith('tmp')]

    # 从Word报告提取基底剪力和位移角
    if HAS_DOCX and results["reports"]:
        try:
            report_path = os.path.join(model_dir, results["reports"][0])
            results["base_shear"] = read_report_base_shear(report_path)
            results["drift_ratio"] = read_report_drift_ratio(report_path)
        except Exception:
            pass

    return results


def format_results(results: Dict[str, Any]) -> str:
    """
    格式化结果为可读字符串

    Args:
        results: read_main_results返回的结果字典

    Returns:
        str: 格式化的结果字符串
    """
    lines = []

    if results.get("periods"):
        lines.append(f"  基本周期: {', '.join(results['periods'][:3])}")

    if results.get("frequencies"):
        lines.append(f"  圆频率: {', '.join(results['frequencies'][:3])}")

    if results.get("freq_hz"):
        lines.append(f"  频率: {', '.join(results['freq_hz'][:3])}")

    if results.get("total_weight"):
        lines.append(f"  楼层总重: {results['total_weight']}")

    if results.get("reactions"):
        lines.append(f"  底部反力: {results['reactions']}")

    # 基底剪力
    if results.get("base_shear"):
        bs = results["base_shear"]
        if bs.get("X") and bs.get("Y"):
            lines.append(
                f"  基底剪力: X向={bs['X']['shear']:.1f}kN(剪重比{bs['X']['ratio']:.2f}), "
                f"Y向={bs['Y']['shear']:.1f}kN(剪重比{bs['Y']['ratio']:.2f})"
            )

    # 最大层间位移角
    if results.get("drift_ratio"):
        dr = results["drift_ratio"]
        x_drift = dr.get('X', {}).get('drift', 'N/A')
        y_drift = dr.get('Y', {}).get('drift', 'N/A')
        lines.append(f"  最大层间位移角: X向={x_drift}, Y向={y_drift}")

    if results.get("reports"):
        lines.append(f"  计算报告: {', '.join(results['reports'][:1])}")

    return "\n".join(lines) if lines else "  无结果文件"


def get_result_summary(model_dir: str, model_name: str = None) -> Dict[str, Any]:
    """
    获取结果摘要（简洁版本）

    Args:
        model_dir: 模型目录路径
        model_name: 模型名称

    Returns:
        dict: 简洁的结果摘要
    """
    results = read_main_results(model_dir, model_name)

    summary = {
        "periods": results.get("periods", [])[:3],
        "base_shear": results.get("base_shear", {}),
        "drift_ratio": results.get("drift_ratio", {}),
        "total_weight": results.get("total_weight"),
        "reports": results.get("reports", [])
    }

    return summary


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("SAUSG 结果读取脚本")
        print("=" * 50)
        print("用法: python sausg_result.py <模型目录> [模型名称]")
        print()
        print("示例:")
        print("  python .claude/skills/scripts/sausg_result.py MulProject/P1 P1")
        print("  python .claude/skills/scripts/sausg_result.py Project/Example")
        sys.exit(1)

    model_dir = sys.argv[1]
    model_name = sys.argv[2] if len(sys.argv) > 2 else None

    results = read_main_results(model_dir, model_name)
    print("主要结果:")
    print(format_results(results))
