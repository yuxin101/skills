#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档处理技能核心实现
"""
import os
import docx
import openpyxl
import pdfplumber
import docx2txt
from PIL import Image
import pytesseract

class DocumentProcessor:
    """文档处理器类"""
    def __init__(self, temp_dir="/tmp/document_processing"):
        """初始化文档处理器"""
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
    def convert_format(self, input_path, output_path):
        """转换文档格式"""
        input_ext = os.path.splitext(input_path)[1].lower()
        output_ext = os.path.splitext(output_path)[1].lower()
        
        if input_ext == ".pdf" and output_ext == ".docx":
            return self._pdf_to_docx(input_path, output_path)
        elif input_ext == ".docx" and output_ext == ".pdf":
            return self._docx_to_pdf(input_path, output_path)
        elif input_ext == ".xlsx" and output_ext == ".csv":
            return self._xlsx_to_csv(input_path, output_path)
        else:
            raise ValueError(f"不支持的格式转换: {input_ext} → {output_ext}")
        
    def _pdf_to_docx(self, pdf_path, docx_path):
        """PDF转Word"""
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(docx_path)
        return True
        
    def _docx_to_pdf(self, docx_path, pdf_path):
        """Word转PDF"""
        # 需要LibreOffice或其他转换工具支持
        # 这里提供一个占位实现
        raise NotImplementedError("Word转PDF需要LibreOffice支持")
        
    def _xlsx_to_csv(self, xlsx_path, csv_path):
        """Excel转CSV"""
        workbook = openpyxl.load_workbook(xlsx_path)
        sheet = workbook.active
        
        with open(csv_path, 'w', encoding='utf-8') as f:
            for row in sheet.iter_rows(values_only=True):
                f.write(",".join(str(cell) for cell in row) + "\n")
        
        return True
        
    def extract_content(self, file_path):
        """提取文档内容"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return self._extract_pdf_content(file_path)
        elif ext == ".docx":
            return self._extract_docx_content(file_path)
        elif ext == ".xlsx":
            return self._extract_xlsx_content(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return self._extract_image_content(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
        
    def _extract_pdf_content(self, pdf_path):
        """提取PDF内容"""
        content = {}
        with pdfplumber.open(pdf_path) as pdf:
            content["text"] = ""
            content["tables"] = []
            
            for page_num, page in enumerate(pdf.pages):
                content["text"] += page.extract_text() or ""
                tables = page.extract_tables()
                if tables:
                    content["tables"].extend(tables)
        
        return content
        
    def _extract_docx_content(self, docx_path):
        """提取Word内容"""
        return docx2txt.process(docx_path)
        
    def _extract_xlsx_content(self, xlsx_path):
        """提取Excel内容"""
        workbook = openpyxl.load_workbook(xlsx_path)
        content = {}
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_content = []
            
            for row in sheet.iter_rows(values_only=True):
                sheet_content.append(list(row))
            
            content[sheet_name] = sheet_content
        
        return content
        
    def _extract_image_content(self, image_path):
        """提取图片内容（OCR）"""
        try:
            text = pytesseract.image_to_string(Image.open(image_path))
            return text
        except Exception as e:
            raise ValueError(f"OCR识别失败: {str(e)}")
        
    def batch_process(self, input_dir, input_ext, output_ext):
        """批量处理文档"""
        processed_files = []
        
        for filename in os.listdir(input_dir):
            if filename.lower().endswith(input_ext.lower()):
                input_path = os.path.join(input_dir, filename)
                output_filename = os.path.splitext(filename)[0] + output_ext
                output_path = os.path.join(self.temp_dir, output_filename)
                
                try:
                    self.convert_format(input_path, output_path)
                    processed_files.append(output_path)
                except Exception as e:
                    print(f"处理失败 {filename}: {str(e)}")
        
        return processed_files
        
    def sync_to_feishu(self, file_path):
        """同步文档到飞书"""
        # 飞书集成实现
        # 1. 上传文件到飞书云空间
        # 2. 更新飞书多维表格
        # 3. 发送飞书消息通知
        pass
        
    def excel_to_word_by_column_format(self, xlsx_path, docx_path, column_formats, sheet_name=None):
        """
        Excel 转换为 Word，每一列对应不同格式
        
        Args:
            xlsx_path: Excel 文件路径
            docx_path: 输出 Word 文件路径
            column_formats: 列格式配置字典，格式为 {列索引: 格式类型}
                支持的格式类型:
                - 'title': 文档标题（一级）
                - 'heading_1': 一级标题
                - 'heading_2': 二级标题  
                - 'heading_3': 三级标题
                - 'normal': 普通正文
                - 'bold': 加粗正文
                - 'italic': 斜体正文
                - 'bullet': 项目符号列表
            示例: {0: 'heading_1', 1: 'normal', 2: 'italic'}
            → 第一列作为一级标题，第二列作为正文，第三列作为斜体
        sheet_name: 指定工作表名称，不指定则使用第一个工作表
        
        Returns:
            bool: 转换是否成功
        """
        # 加载Excel
        workbook = openpyxl.load_workbook(xlsx_path)
        if sheet_name:
            sheet = workbook[sheet_name]
        else:
            sheet = workbook.active
            
        # 创建Word文档
        doc = docx.Document()
        
        # 逐行处理
        for row in sheet.iter_rows(values_only=True):
            for col_idx, cell_value in enumerate(row):
                if cell_value is None:
                    continue
                    
                cell_value = str(cell_value).strip()
                if not cell_value:
                    continue
                    
                # 获取格式类型
                fmt = column_formats.get(col_idx, 'normal')
                
                # 根据格式添加到文档
                if fmt == 'title':
                    doc.add_heading(cell_value, level=0)
                elif fmt == 'heading_1':
                    doc.add_heading(cell_value, level=1)
                elif fmt == 'heading_2':
                    doc.add_heading(cell_value, level=2)
                elif fmt == 'heading_3':
                    doc.add_heading(cell_value, level=3)
                elif fmt == 'normal':
                    doc.add_paragraph(cell_value)
                elif fmt == 'bold':
                    p = doc.add_paragraph()
                    run = p.add_run(cell_value)
                    run.bold = True
                elif fmt == 'italic':
                    p = doc.add_paragraph()
                    run = p.add_run(cell_value)
                    run.italic = True
                elif fmt == 'bullet':
                    doc.add_paragraph(cell_value, style='List Bullet')
                else:
                    # 默认作为普通段落
                    doc.add_paragraph(cell_value)
        
        # 保存Word文档
        doc.save(docx_path)
        return True

if __name__ == "__main__":
    # 测试代码
    processor = DocumentProcessor()
    print("文档处理器初始化完成")