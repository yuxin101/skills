# [OC-WM] licensed-to: macmini@MacminideMac-mini | bundle: vendor-suite | ts: 2026-03-09T17:30:16Z
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(output_path):
    doc = Document()

    # Title
    heading = doc.add_heading('🔷 Skill Trigger V2: 智能触发的艺术 🔷', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Powered by
    p_powered = doc.add_paragraph()
    p_powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_powered = p_powered.add_run('Powered by halfmoon82 🔷')
    run_powered.bold = True
    run_powered.font.color.rgb = RGBColor(0, 112, 192)  # Blue-ish

    # Quote
    p_quote = doc.add_paragraph()
    p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_quote.italic = True
    run_quote = p_quote.add_run('\n"真正的智能，不在于听懂每一个字，而在于洞察每一个意图。"\n')
    run_quote.font.size = Pt(12)
    run_quote.font.color.rgb = RGBColor(80, 80, 80)

    # Intro
    doc.add_paragraph('在 AI Agent 的世界里，最微妙的挑战莫过于此：如何在用户那只言片语的自然表达中，精准捕捉到那一闪而过的执行意图？')
    doc.add_paragraph('Skill Trigger V2 便是为此而生。它不仅仅是一个触发器，更是一套精心编织的决策艺术。')

    # Section 1
    doc.add_heading('🎭 核心哲学：统一与仲裁', level=1)
    doc.add_paragraph('如果说上一代的触发器是简单的“关键词匹配”，那么 V2 则是一场优雅的交响乐指挥。')

    doc.add_heading('1. 统一阈值 (Unified Threshold)', level=2)
    p_s1 = doc.add_paragraph('我们摒弃了各自为战的混乱标准。所有的技能，无论出身 L0 还是 L3，都必须在一个统一的置信度标尺下接受检验。')
    p_s1.add_run('0.50 (50%)').bold = True
    p_s1.add_run(' —— 这是我们划定的黄金分割线。只有越过这条线的意图，才会被系统认真对待。')

    doc.add_heading('2. 优先级仲裁 (Priority Arbitration)', level=2)
    doc.add_paragraph('当多个技能同时举手示意时，谁该获得执行权？V2 引入了 L0-L3 权重仲裁机制。')
    
    # List
    doc.add_paragraph('• L0 (系统级)：如基石般稳固，拥有最高话语权 (1.2x 权重)。', style='List Bullet')
    doc.add_paragraph('• L1 (核心能力)：中流砥柱，优先响应 (1.1x 权重)。', style='List Bullet')
    doc.add_paragraph('• L2/L3 (扩展应用)：丰富多彩，但在冲突时懂得礼让。', style='List Bullet')

    # Section 2
    doc.add_heading('🌟 为什么选择 V2？', level=1)
    doc.add_paragraph('• 精准而不失灵动：它能听懂“帮我写个爬虫”，也能理解“看看今天的天气”。', style='List Bullet')
    doc.add_paragraph('• 秩序井然：告别技能冲突的混沌，每一次触发都是深思熟虑的最优解。', style='List Bullet')
    doc.add_paragraph('• 透明可溯：每一次决策都会留下 fit_result 的足迹，让你知道 AI 为什么这样做。', style='List Bullet')

    # Section 3
    doc.add_heading('📜 快速接入', level=1)
    doc.add_paragraph('只需简单的一行命令，即可让你的 Agent 拥有这份智慧：')
    p_code = doc.add_paragraph('clawhub install skill-trigger-v2@latest')
    p_code.style = 'Quote'

    # Footer / Attribution
    doc.add_heading('⚖️ 归属与致谢', level=1)
    doc.add_paragraph('本作品承载着 halfmoon82 对 AI 交互体验的极致追求。')
    
    doc.add_paragraph('• 核心算法: halfmoon82', style='List Bullet')
    doc.add_paragraph('• 版权所有: © 2026 halfmoon82. All rights reserved.', style='List Bullet')
    doc.add_paragraph('• 发布平台: ClawHub (https://clawhub.ai/halfmoon82/skill-trigger-v2)', style='List Bullet')

    # Final touch
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.add_run('\n— 让每一次触发，都成为一次心有灵犀的默契 —').italic = True

    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    output_file = os.path.expanduser("~/.openclaw/workspace/skills/skill-trigger-v2/Skill_Trigger_V2_README_Humanities_CN.docx")
    create_docx(output_file)
