#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
企业红头文件生成器
新增功能：
1. 支持页脚（自动页码）
2. 增加内容自动跨页
3. 优化多页排版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


class RedHeadDocument:
    """企业红头文件生成器"""
    
    def __init__(self, output_dir="/workspace/output"):
        self.doc = Document()
        self.output_dir = output_dir
        self._setup_page()
        self._setup_footer()  # 初始化页脚设置
        
    def _setup_page(self):
        """设置标准 A4 页面格式"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.top_margin = Mm(25.4)
            section.bottom_margin = Mm(25.4)
            section.left_margin = Mm(31.8)
            section.right_margin = Mm(31.8)
            # 设置页脚距离底部的距离
            section.footer_distance = Mm(15)

    def _setup_footer(self):
        """设置页脚页码（居中显示）"""
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加 "第 X 页 共 Y 页" 格式
        def add_page_number_field(paragraph):
            # "第"
            run = paragraph.add_run("第 ")
            self._set_font(run, "仿宋", size=Pt(10))
            
            # 当前页码字段
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), 'PAGE')
            p._element.append(fldSimple)
            
            # " 页 共 "
            run = paragraph.add_run(" 页 共 ")
            self._set_font(run, "仿宋", size=Pt(10))
            
            # 总页数地段
            fldSimple2 = OxmlElement('w:fldSimple')
            fldSimple2.set(qn('w:instr'), 'NUMPAGES')
            p._element.append(fldSimple2)
            
            # " 页"
            run = paragraph.add_run(" 页")
            self._set_font(run, "仿宋", size=Pt(10))

        add_page_number_field(p)

    def _set_font(self, run, font_name="宋体", size=Pt(12), bold=False, color=None):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = size
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_red_header_title(self, text="企业文件"):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        self._set_font(run, "方正小标宋简体", size=Pt(48), bold=True, color=(255, 0, 0))

    def add_doc_number(self, org="XXX", year="20XX", num="6"):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        text = f"{org}〔{year}〕{num} 号"
        run = p.add_run(text)
        self._set_font(run, "仿宋", size=Pt(16))

    def add_red_separator(self):
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'FF0000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_main_title(self, text="工程开工通知"):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        self._set_font(run, "黑体", size=Pt(26), bold=True)

    def add_recipient(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        self._set_font(run, "仿宋", size=Pt(16), bold=True)

    def add_section_title(self, text):
        """添加小节标题"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        self._set_font(run, "黑体", size=Pt(16), bold=True)

    def add_body_paragraph(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(text)
        self._set_font(run, "仿宋", size=Pt(16))

    def add_info_row(self, left_text, right_text):
        table = self.doc.add_table(rows=1, cols=2)
        table.width = Inches(6)
        p_left = table.cell(0, 0).paragraphs[0]
        run_l = p_left.add_run(left_text)
        self._set_font(run_l, "仿宋", size=Pt(16))
        p_right = table.cell(0, 1).paragraphs[0]
        run_r = p_right.add_run(right_text)
        self._set_font(run_r, "仿宋", size=Pt(16))

    def add_signature_area(self, company, date):
        # 增加一些空行使落款更有层次感
        for _ in range(2):
            self.doc.add_paragraph()
        for text in [company, date]:
            p = self.doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.paragraph_format.right_indent = Pt(20)
            run = p.add_run(text)
            self._set_font(run, "仿宋", size=Pt(16))

    def save(self, name="企业红头文件.docx"):
        os.makedirs(self.output_dir, exist_ok=True)
        path = os.path.join(self.output_dir, name)
        self.doc.save(path)
        print(f"文件已保存至: {path}")
        return path


def run_generator():
    gen = RedHeadDocument()
    
    # --- 首页红头部分 ---
    gen.add_red_header_title("企业文件")
    gen.add_doc_number("XXX", "20XX", "6")
    gen.add_red_separator()
    gen.add_main_title("关于 20XX 年度有线电视线路\n专项抢修工程的开工通知")
    
    gen.add_recipient("XX 通信工程有限责任公司：")
    
    # --- 第一页正文内容 ---
    gen.add_section_title("一、工程背景与任务概述")
    gen.add_body_paragraph(
        "经我公司研究决定，由贵公司负责的有线电视抢修任务，工程施工地点位于 XXX 区域。 "
        "本工程为单项抢修工程，由于近期线路信号波动剧烈，部分节点信号电平低于标准值，"
        "严重影响了周边用户的收视体验。经现场勘测，决定将原有的旧外三分配器统一更换为高性能的旧外四分配器，"
        "并需同步增铺相关配套电缆，路由长度共计约 200 米。"
    )
    
    gen.add_section_title("二、施工技术要求")
    gen.add_body_paragraph(
        "1. 电缆铺设：所有新增电缆必须符合国家广电行业标准，铺设过程中需做好防水与防雷接地处理。 "
        "2. 分配器安装：更换后的外四分配器需进行逐级电平测试，确保末端用户端输出信号强度不低于 65dBμV。 "
        "3. 标识标注：施工完成后，必须在关键节点加挂醒目的资产标识牌，注明施工单位及日期。"
    )
    
    # --- 第二页内容（增加内容使其跨页） ---
    gen.add_section_title("三、安全施工与环境保护")
    gen.add_body_paragraph(
        "施工单位在作业过程中必须严格遵守《安全生产法》，高空作业必须佩戴安全绳。 "
        "由于施工点位于居民区周边，务必在施工区域外侧拉起警戒线，防止非作业人员进入引发危险。 "
        "严禁在施工现场乱扔废旧分配器及线缆碎料，完工后必须做到‘工完、料净、场清’。"
    )
    
    gen.add_section_title("四、时间进度要求")
    gen.add_body_paragraph(
        "请贵公司务必根据要求时间 X 月 X 日组织施工队伍准时进场。工程总体时限为 X 天。 "
        "如遇不可抗力（如极端天气）需延期，必须提前 24 小时向我公司工程部提交书面申请，经核准后方可顺延。"
    )
    
    gen.add_section_title("五、联系与协调机制")
    gen.add_body_paragraph(
        "施工期间，贵公司项目经理需保持电话 24 小时畅通。我公司将委派专职监理人员对施工质量进行实时抽检。 "
        "若发现施工质量不达标，我方有权要求立即停工整改，由此产生的费用由贵公司自行承担。"
    )
    
    # --- 结尾部分 ---
    gen.add_info_row("联系人：张经理", "电话：010-8888XXXX")
    gen.add_info_row("工程部负责人签名：", "公司负责人签名：")
    gen.add_signature_area("XX 通信工程有限公司（公章）", "20XX 年 X 月 X 日")
    
    return gen.save()


if __name__ == "__main__":
    run_generator()
