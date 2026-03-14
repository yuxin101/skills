#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TDoc DOCX 差异对比引擎
生成两个 DOCX 文件间的 Unified Diff 报告。

用法:
  python diff_docx.py old.docx new.docx --output diff_report.md
  python diff_docx.py old.docx new.docx
"""

import argparse
import difflib
import os
import sys
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def docx_to_text(docx_path):
    """将 DOCX 转换为纯文本"""
    if not DOCX_AVAILABLE:
        raise ImportError("缺少 python-docx 库")

    doc = Document(str(docx_path))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ''
            if style.startswith('Heading'):
                level = style.replace('Heading ', '').replace('Heading', '')
                prefix = '#' * (int(level) if level.isdigit() else 1)
                lines.append(f"{prefix} {text}")
            else:
                lines.append(text)
        else:
            lines.append('')

    for i, table in enumerate(doc.tables):
        lines.append(f"\n[表格 {i + 1}]")
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            lines.append(row_text)

    return '\n'.join(lines)


def generate_unified_diff(old_text, new_text, old_name, new_name):
    """生成 Unified Diff 格式"""
    old_lines = old_text.splitlines(keepends=True)
    new_lines = new_text.splitlines(keepends=True)

    diff = difflib.unified_diff(
        old_lines, new_lines,
        fromfile=old_name, tofile=new_name,
        lineterm='\n'
    )
    return ''.join(diff)


def generate_summary(old_text, new_text):
    """生成变更摘要"""
    old_lines = old_text.splitlines()
    new_lines = new_text.splitlines()

    matcher = difflib.SequenceMatcher(None, old_lines, new_lines)
    added = 0
    removed = 0
    changed = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'insert':
            added += (j2 - j1)
        elif tag == 'delete':
            removed += (i2 - i1)
        elif tag == 'replace':
            changed += max(i2 - i1, j2 - j1)

    similarity = matcher.ratio()

    return {
        'added_lines': added,
        'removed_lines': removed,
        'changed_lines': changed,
        'similarity': f"{similarity * 100:.1f}%",
        'old_total_lines': len(old_lines),
        'new_total_lines': len(new_lines),
    }


def main():
    parser = argparse.ArgumentParser(description='TDoc DOCX 差异对比引擎')
    parser.add_argument('old_docx', help='旧版文档路径')
    parser.add_argument('new_docx', help='新版文档路径')
    parser.add_argument('--output', '-o', help='输出 diff 报告路径')

    args = parser.parse_args()

    old_path = Path(args.old_docx)
    new_path = Path(args.new_docx)

    if not old_path.exists():
        print(f"❌ 文件不存在: {old_path}", file=sys.stderr)
        sys.exit(1)
    if not new_path.exists():
        print(f"❌ 文件不存在: {new_path}", file=sys.stderr)
        sys.exit(1)

    print(f"📄 读取旧版本: {old_path}")
    old_text = docx_to_text(old_path)

    print(f"📄 读取新版本: {new_path}")
    new_text = docx_to_text(new_path)

    print(f"📊 生成差异报告...")

    diff_text = generate_unified_diff(old_text, new_text, str(old_path), str(new_path))
    summary = generate_summary(old_text, new_text)

    # 构建报告
    report = []
    report.append("# DOCX 差异对比报告\n")
    report.append(f"**旧版本:** {old_path}")
    report.append(f"**新版本:** {new_path}\n")
    report.append("## 变更摘要\n")
    report.append(f"| 指标 | 值 |")
    report.append(f"|------|-----|")
    report.append(f"| 相似度 | {summary['similarity']} |")
    report.append(f"| 新增行 | +{summary['added_lines']} |")
    report.append(f"| 删除行 | -{summary['removed_lines']} |")
    report.append(f"| 修改行 | ~{summary['changed_lines']} |")
    report.append(f"| 旧版总行数 | {summary['old_total_lines']} |")
    report.append(f"| 新版总行数 | {summary['new_total_lines']} |")
    report.append("")

    if diff_text:
        report.append("## 详细差异\n")
        report.append("```diff")
        report.append(diff_text)
        report.append("```")
    else:
        report.append("✅ 两个文档内容完全一致，无差异。")

    report_text = '\n'.join(report)

    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
        print(f"✅ 差异报告已保存: {output_path}")
    else:
        print("\n" + report_text)


if __name__ == "__main__":
    main()
