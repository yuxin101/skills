#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TDoc DOCX 文档读取引擎
支持 .docx/.doc 格式，提取文本、表格、图片、元数据。

用法:
  python read_docx.py document.docx
  python read_docx.py document.docx --format json
  python read_docx.py document.docx --format markdown
  python read_docx.py document.docx --extract tables
  python read_docx.py ./docs_folder --batch --format json --output results.json
"""

import argparse
import json
import os
import subprocess
import sys
import traceback
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordReader:
    """Word 文档读取器"""

    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.document = None
        self.format_type = None

        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {self.file_path.suffix}")

    def read(self):
        """自动检测格式并读取"""
        if self.file_path.suffix.lower() == '.docx':
            return self.read_docx()
        else:
            return self.read_doc()

    def read_docx(self):
        """读取 .docx 格式文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("缺少 python-docx 库。请安装: pip3 install python-docx")
        self.document = Document(str(self.file_path))
        self.format_type = 'docx'
        return True

    def read_doc(self):
        """读取 .doc 格式文档（使用 antiword 或 LibreOffice）"""
        # 先尝试 antiword
        try:
            result = subprocess.run(
                ['antiword', str(self.file_path)],
                capture_output=True, text=True, encoding='utf-8', timeout=30
            )
            if result.returncode == 0:
                self.format_type = 'doc'
                self._doc_text = result.stdout
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 尝试 LibreOffice 转换
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'docx',
                     '--outdir', tmpdir, str(self.file_path)],
                    capture_output=True, timeout=60, check=True
                )
                docx_path = Path(tmpdir) / (self.file_path.stem + '.docx')
                if docx_path.exists():
                    self.document = Document(str(docx_path))
                    self.format_type = 'docx'
                    return True
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            pass

        raise RuntimeError(
            "无法读取 .doc 文件。请安装 antiword (brew install antiword) "
            "或 LibreOffice (brew install --cask libreoffice)"
        )

    def read_metadata(self):
        """读取文档元数据"""
        metadata = {
            'filename': self.file_path.name,
            'size': f"{self.file_path.stat().st_size} bytes",
            'size_kb': f"{self.file_path.stat().st_size / 1024:.1f} KB",
            'created': datetime.fromtimestamp(self.file_path.stat().st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(self.file_path.stat().st_mtime).isoformat(),
        }

        if self.format_type == 'docx' and self.document and hasattr(self.document, 'core_properties'):
            props = self.document.core_properties
            metadata.update({
                'title': getattr(props, 'title', '') or '',
                'author': getattr(props, 'author', '') or '',
                'subject': getattr(props, 'subject', '') or '',
                'keywords': getattr(props, 'keywords', '') or '',
                'comments': getattr(props, 'comments', '') or '',
                'category': getattr(props, 'category', '') or '',
                'paragraph_count': len(self.document.paragraphs),
                'table_count': len(self.document.tables),
            })

        return metadata

    def extract_text(self):
        """提取文档文本"""
        if self.format_type == 'doc' and hasattr(self, '_doc_text'):
            return self._doc_text

        text_parts = []
        if self.format_type == 'docx' and self.document:
            # 遍历文档体的元素，保留段落和表格的顺序
            for para in self.document.paragraphs:
                if para.text.strip():
                    style_name = para.style.name if para.style else ''
                    if style_name.startswith('Heading'):
                        level = style_name.replace('Heading ', '').replace('Heading', '')
                        prefix = '#' * (int(level) if level.isdigit() else 1)
                        text_parts.append(f"{prefix} {para.text}")
                    else:
                        text_parts.append(para.text)

            # 提取表格文本
            for table in self.document.tables:
                table_lines = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_lines.append(' | '.join(row_text))
                if table_lines:
                    text_parts.append('\n'.join(table_lines))

        return '\n\n'.join(text_parts)

    def extract_tables(self):
        """提取表格数据"""
        tables = []
        if self.format_type == 'docx' and self.document:
            for i, table in enumerate(self.document.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'id': i + 1,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'data': table_data
                })
        return tables

    def extract_images(self):
        """提取图片信息"""
        images = []
        if self.format_type == 'docx' and self.document:
            try:
                part = self.document.part
                for rel in part.rels.values():
                    if "image" in rel.reltype:
                        try:
                            image_data = rel.target_part.blob
                            images.append({
                                'id': rel.rId,
                                'filename': f"image_{rel.rId}.{rel.target_ref.split('.')[-1]}",
                                'size': f"{len(image_data)} bytes",
                                'size_kb': f"{len(image_data) / 1024:.1f} KB"
                            })
                        except Exception:
                            images.append({
                                'id': rel.rId,
                                'filename': rel.target_ref if hasattr(rel, 'target_ref') else 'unknown',
                                'size': 'unknown'
                            })
            except Exception:
                pass
        return images

    def extract_all(self):
        """提取所有内容"""
        return {
            'metadata': self.read_metadata(),
            'format': self.format_type,
            'text': self.extract_text(),
            'tables': self.extract_tables(),
            'images': self.extract_images()
        }

    def to_markdown(self, extract_type='all'):
        """转换为 Markdown 格式"""
        if extract_type == 'text':
            return self.extract_text()

        result = self.extract_all()
        md = []

        # 文件名标题
        md.append(f"# {result['metadata']['filename']}")
        md.append("")

        # 元数据
        meta = result['metadata']
        if meta.get('title'):
            md.append(f"**标题**: {meta['title']}")
        if meta.get('author'):
            md.append(f"**作者**: {meta['author']}")
        md.append(f"**文件大小**: {meta.get('size_kb', meta['size'])}")
        md.append(f"**修改时间**: {meta['modified']}")
        md.append("")

        # 正文
        if result['text']:
            md.append("## 正文内容")
            md.append("")
            md.append(result['text'])
            md.append("")

        # 表格
        if result['tables']:
            md.append("## 表格内容")
            md.append("")
            for table in result['tables']:
                md.append(f"### 表格 {table['id']} ({table['rows']}行 × {table['columns']}列)")
                md.append("")
                if table['data']:
                    # Markdown 表格头
                    header = table['data'][0]
                    md.append("| " + " | ".join(str(c) for c in header) + " |")
                    md.append("| " + " | ".join("---" for _ in header) + " |")
                    for row in table['data'][1:]:
                        md.append("| " + " | ".join(str(c) for c in row) + " |")
                md.append("")

        # 图片
        if result['images']:
            md.append("## 图片列表")
            md.append("")
            for img in result['images']:
                md.append(f"- **{img['filename']}** ({img.get('size_kb', img['size'])})")
            md.append("")

        return '\n'.join(md)

    def to_json(self, extract_type='all'):
        """转换为 JSON"""
        if extract_type == 'text':
            return json.dumps({'text': self.extract_text()}, ensure_ascii=False, indent=2)
        elif extract_type == 'tables':
            return json.dumps({'tables': self.extract_tables()}, ensure_ascii=False, indent=2)
        elif extract_type == 'metadata':
            return json.dumps({'metadata': self.read_metadata()}, ensure_ascii=False, indent=2)
        elif extract_type == 'images':
            return json.dumps({'images': self.extract_images()}, ensure_ascii=False, indent=2)
        return json.dumps(self.extract_all(), ensure_ascii=False, indent=2)

    def to_text(self, extract_type='all'):
        """转换为纯文本"""
        if extract_type == 'text':
            return self.extract_text()

        result = self.extract_all()
        lines = []
        lines.append(f"文件: {result['metadata']['filename']}")
        lines.append("=" * 50)

        meta = result['metadata']
        for key in ['title', 'author', 'subject', 'keywords']:
            if meta.get(key):
                lines.append(f"{key}: {meta[key]}")
        lines.append(f"大小: {meta.get('size_kb', meta['size'])}")
        lines.append("")

        if result['text']:
            lines.append("正文内容:")
            lines.append("-" * 20)
            lines.append(result['text'])
            lines.append("")

        if result['tables']:
            lines.append("表格内容:")
            lines.append("-" * 20)
            for table in result['tables']:
                lines.append(f"表格 {table['id']}:")
                for row in table['data']:
                    lines.append("  " + " | ".join(str(c) for c in row))
                lines.append("")

        return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='TDoc DOCX 文档读取引擎')
    parser.add_argument('path', help='文档路径或目录路径（批量模式）')
    parser.add_argument('--format', choices=['json', 'text', 'markdown'],
                        default='text', help='输出格式')
    parser.add_argument('--extract', choices=['text', 'tables', 'images', 'metadata', 'all'],
                        default='all', help='提取内容类型')
    parser.add_argument('--batch', action='store_true', help='批量处理模式')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--encoding', default='utf-8', help='文本编码')

    args = parser.parse_args()

    try:
        if args.batch:
            path = Path(args.path)
            if not path.is_dir():
                print("❌ 批量模式需要指定目录路径", file=sys.stderr)
                sys.exit(1)

            word_files = list(path.glob("**/*.docx")) + list(path.glob("**/*.doc"))
            if not word_files:
                print("未找到 Word 文档")
                sys.exit(0)

            print(f"📄 找到 {len(word_files)} 个 Word 文档")
            results = {}

            for file_path in word_files:
                print(f"  处理: {file_path.name}")
                try:
                    reader = WordReader(file_path)
                    reader.read()
                    if args.format == 'json':
                        results[str(file_path)] = reader.extract_all()
                    elif args.format == 'markdown':
                        results[str(file_path)] = reader.to_markdown(args.extract)
                    else:
                        results[str(file_path)] = reader.to_text(args.extract)
                except Exception as e:
                    results[str(file_path)] = {'error': str(e)}

            output_text = json.dumps(results, ensure_ascii=False, indent=2)
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(output_text)
                print(f"✅ 结果已保存: {args.output}")
            else:
                print(output_text)
        else:
            reader = WordReader(args.path)
            reader.read()

            if args.format == 'json':
                content = reader.to_json(args.extract)
            elif args.format == 'markdown':
                content = reader.to_markdown(args.extract)
            else:
                content = reader.to_text(args.extract)

            if args.output:
                with open(args.output, 'w', encoding=args.encoding) as f:
                    f.write(content)
                print(f"✅ 结果已保存: {args.output}")
            else:
                print(content)

    except Exception as e:
        print(f"❌ 错误: {str(e)}", file=sys.stderr)
        if '--debug' in sys.argv:
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
