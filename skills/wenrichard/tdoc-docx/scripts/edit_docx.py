#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TDoc DOCX 文档编辑引擎
基于 JSON 规则批量编辑 DOCX 文件，支持文本替换、内容添加、样式标记。

用法:
  python edit_docx.py input.docx output.docx edits.json
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 样式应用
# ============================================================
def apply_highlight(run, color='yellow'):
    """应用高亮"""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), color)
    rPr.append(highlight)


def apply_style(run, style_type):
    """对 run 应用样式"""
    if style_type == "highlight":
        apply_highlight(run)
    elif style_type == "delete":
        run.font.strike = True
        apply_highlight(run, 'red')
    elif style_type == "bold":
        run.font.bold = True
    elif style_type == "underline":
        run.font.underline = WD_UNDERLINE.SINGLE
    elif style_type == "italic":
        run.font.italic = True
    elif style_type == "red":
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


# ============================================================
# 替换逻辑
# ============================================================
def replace_in_paragraph(para, search_text, replace_text, style_type="replace"):
    """在段落中替换文字，保留格式"""
    if search_text not in para.text:
        return False

    if style_type == "replace":
        # 简单替换：遍历 runs
        for run in list(para.runs):
            if search_text in run.text:
                run.text = run.text.replace(search_text, replace_text)
                return True

        # 如果文本跨越多个 runs，重建段落
        full_text = para.text
        if search_text in full_text:
            new_text = full_text.replace(search_text, replace_text)
            # 保留第一个 run 的格式
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run.text = ''
                first_run.text = new_text
            return True
    else:
        # 带样式的替换
        for run in list(para.runs):
            if search_text in run.text:
                parts = run.text.split(search_text, 1)
                run.text = parts[0]

                # 添加替换文本（带样式）
                new_run = para.add_run(replace_text)
                # 复制原格式
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                apply_style(new_run, style_type)

                # 添加剩余文本
                if len(parts) > 1 and parts[1]:
                    remaining_run = para.add_run(parts[1])
                    if run.font.name:
                        remaining_run.font.name = run.font.name
                    if run.font.size:
                        remaining_run.font.size = run.font.size

                return True

    return False


def replace_in_table(table, search_text, replace_text, style_type="replace"):
    """在表格中替换文字"""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, search_text, replace_text, style_type):
                    replaced = True
    return replaced


def add_text_in_paragraph(para, after_text, add_text, style_type="highlight"):
    """在段落中指定文本后添加内容"""
    if after_text not in para.text:
        return False
    new_run = para.add_run(add_text)
    apply_style(new_run, style_type)
    return True


# ============================================================
# 主编辑函数
# ============================================================
def edit_docx(input_path, output_path, edits):
    """
    编辑 DOCX 文件

    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        edits: 编辑规则字典
    """
    doc = Document(input_path)
    stats = {"replacements": 0, "additions": 0}

    # 处理替换
    for edit in edits.get("replacements", []):
        search_text = edit["search"]
        replace_text = edit.get("replace", "")
        style = edit.get("style", "replace")

        for para in doc.paragraphs:
            if replace_in_paragraph(para, search_text, replace_text, style):
                stats["replacements"] += 1

        for table in doc.tables:
            if replace_in_table(table, search_text, replace_text, style):
                stats["replacements"] += 1

    # 处理添加
    for edit in edits.get("additions", []):
        after_text = edit["after"]
        add_text = edit["text"]
        style = edit.get("style", "highlight")

        for para in doc.paragraphs:
            if add_text_in_paragraph(para, after_text, add_text, style):
                stats["additions"] += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if add_text_in_paragraph(para, after_text, add_text, style):
                            stats["additions"] += 1

    # 保存
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"✅ 编辑完成: {output_path}")
    print(f"📊 统计: 替换 {stats['replacements']} 处, 添加 {stats['additions']} 处")
    return stats


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    edits_path = Path(sys.argv[3])

    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)
    if not edits_path.exists():
        print(f"❌ 编辑规则文件不存在: {edits_path}")
        sys.exit(1)

    with open(edits_path, 'r', encoding='utf-8') as f:
        edits = json.load(f)

    edit_docx(input_path, output_path, edits)


if __name__ == "__main__":
    main()
