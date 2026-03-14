#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TDoc DOCX 文档创建引擎
支持多种风格（默认/公文/商务/学术），从 Markdown 创建专业 Word 文档。
同时暴露 DocxCreator 类供 Python 脚本直接调用。

用法:
  # CLI 方式（从 Markdown 创建）
  python create_docx.py --title "标题" --output output.docx
  python create_docx.py --from-markdown input.md --output output.docx --style gov

  # Python API 方式（推荐，直接调用 DocxCreator）
  from create_docx import DocxCreator
  creator = DocxCreator(style='default')
  creator.add_title("标题")
  creator.add_heading1("一、概述")
  creator.add_paragraph("正文内容")
  creator.add_table(["列1", "列2"], [["A", "B"]])
  creator.save("output.docx")
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 中文引号处理
# ============================================================
LEFT_QUOTE = '\u201c'   # "
RIGHT_QUOTE = '\u201d'  # "


def convert_quotes(text):
    """将英文引号转换为中文引号（成对替换）"""
    result = []
    in_quote = False
    for char in text:
        if char == '"':
            if in_quote:
                result.append(RIGHT_QUOTE)
                in_quote = False
            else:
                result.append(LEFT_QUOTE)
                in_quote = True
        else:
            result.append(char)
    return ''.join(result)


# ============================================================
# 字体配置
# ============================================================
STYLE_CONFIGS = {
    'default': {
        'title_font': 'Arial',
        'title_font_cn': '微软雅黑',
        'title_size': Pt(22),
        'h1_font': 'Arial',
        'h1_font_cn': '微软雅黑',
        'h1_size': Pt(16),
        'h2_font': 'Arial',
        'h2_font_cn': '微软雅黑',
        'h2_size': Pt(14),
        'body_font': 'Arial',
        'body_font_cn': '微软雅黑',
        'body_size': Pt(12),
        'line_spacing': Pt(22),
        'first_line_indent': Pt(24),
        'page_width': Inches(8.27),
        'page_height': Inches(11.69),
        'margin_top': Inches(1),
        'margin_bottom': Inches(1),
        'margin_left': Inches(1.25),
        'margin_right': Inches(1.25),
    },
    'gov': {
        'title_font': '方正小标宋简体',
        'title_font_cn': '方正小标宋简体',
        'title_size': Pt(22),  # 2号字
        'h1_font': '黑体',
        'h1_font_cn': '黑体',
        'h1_size': Pt(16),    # 3号字
        'h2_font': '楷体_GB2312',
        'h2_font_cn': '楷体_GB2312',
        'h2_size': Pt(16),
        'body_font': '仿宋_GB2312',
        'body_font_cn': '仿宋_GB2312',
        'body_size': Pt(16),
        'line_spacing': Pt(28),
        'first_line_indent': Pt(32),  # 2字符
        'page_width': Inches(8.27),
        'page_height': Inches(11.69),
        'margin_top': Inches(1),
        'margin_bottom': Inches(1),
        'margin_left': Inches(1.25),
        'margin_right': Inches(1.25),
    },
    'business': {
        'title_font': 'Calibri',
        'title_font_cn': '微软雅黑',
        'title_size': Pt(26),
        'h1_font': 'Calibri',
        'h1_font_cn': '微软雅黑',
        'h1_size': Pt(16),
        'h2_font': 'Calibri',
        'h2_font_cn': '微软雅黑',
        'h2_size': Pt(14),
        'body_font': 'Calibri',
        'body_font_cn': '微软雅黑',
        'body_size': Pt(11),
        'line_spacing': Pt(20),
        'first_line_indent': None,
        'page_width': Inches(8.5),
        'page_height': Inches(11),
        'margin_top': Inches(1),
        'margin_bottom': Inches(1),
        'margin_left': Inches(1),
        'margin_right': Inches(1),
    },
    'academic': {
        'title_font': 'Times New Roman',
        'title_font_cn': '宋体',
        'title_size': Pt(22),
        'h1_font': 'Times New Roman',
        'h1_font_cn': '黑体',
        'h1_size': Pt(16),
        'h2_font': 'Times New Roman',
        'h2_font_cn': '楷体',
        'h2_size': Pt(14),
        'body_font': 'Times New Roman',
        'body_font_cn': '宋体',
        'body_size': Pt(12),
        'line_spacing': Pt(22),
        'first_line_indent': Pt(24),
        'page_width': Inches(8.27),
        'page_height': Inches(11.69),
        'margin_top': Inches(1),
        'margin_bottom': Inches(1),
        'margin_left': Inches(1.25),
        'margin_right': Inches(1.25),
    },
}


# ============================================================
# DocxCreator 核心类
# ============================================================
class DocxCreator:
    """Word 文档创建器，支持多种风格"""

    def __init__(self, style='default'):
        self.style_name = style
        self.config = STYLE_CONFIGS.get(style, STYLE_CONFIGS['default'])
        self.doc = Document()
        self._setup_default_style()
        self._setup_page()

    def _set_font(self, run, font_name, font_cn, size, bold=False, color=None):
        """统一设置字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.font.size = size
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color

    def _setup_default_style(self):
        """设置默认样式"""
        style = self.doc.styles['Normal']
        style.font.name = self.config['body_font']
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['body_font_cn'])
        style.font.size = self.config['body_size']

    def _setup_page(self):
        """设置页面"""
        section = self.doc.sections[0]
        section.page_height = self.config['page_height']
        section.page_width = self.config['page_width']
        section.top_margin = self.config['margin_top']
        section.bottom_margin = self.config['margin_bottom']
        section.left_margin = self.config['margin_left']
        section.right_margin = self.config['margin_right']
        self._add_page_number(section)

    def _add_page_number(self, section):
        """添加页码"""
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = self.config['body_font']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['body_font_cn'])
        run.font.size = self.config['body_size']

    def _set_paragraph_format(self, p, first_line_indent=True,
                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """设置段落格式"""
        p.paragraph_format.line_spacing = self.config['line_spacing']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if first_line_indent and self.config['first_line_indent']:
            p.paragraph_format.first_line_indent = self.config['first_line_indent']
        p.paragraph_format.widow_control = False
        p.alignment = alignment

    # ------ 内容添加方法 ------

    def add_title(self, text):
        """添加大标题"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, self.config['title_font'],
                        self.config['title_font_cn'], self.config['title_size'],
                        bold=(self.style_name != 'gov'))
        self._set_paragraph_format(p, first_line_indent=False,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # 标题后空一行
        p.paragraph_format.space_after = Pt(12)
        return p

    def add_author(self, text):
        """添加署名（支持多行，用 \\n 分隔）"""
        text = convert_quotes(text)
        lines = text.split('\\n')
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line.strip())
            if self.style_name == 'gov':
                self._set_font(run, '楷体_GB2312', '楷体_GB2312', Pt(16))
            else:
                self._set_font(run, self.config['body_font'],
                                self.config['body_font_cn'], self.config['body_size'])
            self._set_paragraph_format(p, first_line_indent=False,
                                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return p

    def add_heading1(self, text):
        """添加一级标题"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, self.config['h1_font'],
                        self.config['h1_font_cn'], self.config['h1_size'],
                        bold=(self.style_name != 'gov'))
        self._set_paragraph_format(p, first_line_indent=(self.style_name == 'gov'))
        return p

    def add_heading2(self, text):
        """添加二级标题"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, self.config['h2_font'],
                        self.config['h2_font_cn'], self.config['h2_size'])
        self._set_paragraph_format(p, first_line_indent=True)
        return p

    def add_paragraph(self, text, bold_prefix=None):
        """添加正文段落"""
        text = convert_quotes(text)
        if bold_prefix:
            bold_prefix = convert_quotes(bold_prefix)
        p = self.doc.add_paragraph()
        if bold_prefix:
            run1 = p.add_run(bold_prefix)
            self._set_font(run1, self.config['body_font'],
                            self.config['body_font_cn'], self.config['body_size'], bold=True)
            run2 = p.add_run(text)
            self._set_font(run2, self.config['body_font'],
                            self.config['body_font_cn'], self.config['body_size'])
        else:
            run = p.add_run(text)
            self._set_font(run, self.config['body_font'],
                            self.config['body_font_cn'], self.config['body_size'])
        self._set_paragraph_format(p)
        return p

    def add_table(self, headers, rows, col_widths=None):
        """添加表格"""
        num_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 设置表头
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(convert_quotes(header_text))
            run.font.bold = True
            self._set_font(run, self.config['body_font'],
                            self.config['body_font_cn'], self.config['body_size'], bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 表头背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'D9E2F3')
            cell._tc.get_or_add_tcPr().append(shading)

        # 设置数据行
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < num_cols:
                    cell = row.cells[i]
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(convert_quotes(str(cell_text)))
                    self._set_font(run, self.config['body_font'],
                                    self.config['body_font_cn'], self.config['body_size'])

        # 设置列宽
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    if i < len(row.cells):
                        row.cells[i].width = Cm(width)

        return table

    def add_image(self, image_path, width=None, caption=None):
        """添加图片"""
        if not os.path.exists(image_path):
            print(f"⚠️ 图片文件不存在: {image_path}")
            return None
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if width:
            run.add_picture(image_path, width=Pt(width))
        else:
            run.add_picture(image_path, width=Inches(4))

        if caption:
            cap_p = self.doc.add_paragraph()
            cap_run = cap_p.add_run(convert_quotes(caption))
            self._set_font(cap_run, self.config['body_font'],
                            self.config['body_font_cn'], Pt(10))
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def add_empty_line(self):
        """添加空行"""
        return self.doc.add_paragraph()

    def save(self, filepath):
        """保存文档"""
        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(filepath)
        size_kb = os.path.getsize(filepath) / 1024
        print(f"✅ 文档已保存: {filepath} ({size_kb:.1f} KB)")


# ============================================================
# Markdown 解析器
# ============================================================
def parse_markdown(filepath):
    """解析 Markdown 文件为结构化内容"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    result = {'title': None, 'content': []}

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # 主标题
        if line.startswith('# ') and result['title'] is None:
            result['title'] = line[2:].strip()
            continue

        # 一级标题
        if line.startswith('## '):
            result['content'].append({'type': 'heading1', 'text': line[3:].strip()})
            continue

        # 二级标题
        if line.startswith('### '):
            result['content'].append({'type': 'heading2', 'text': line[4:].strip()})
            continue

        # 分隔线
        if line.strip() in ['---', '***', '___']:
            continue

        # 模式1: 数字. **标题**。内容
        match = re.match(r'^(\d+\.\s+)\*\*([^*]+)\*\*[。\.]\s*(.+)$', line)
        if match:
            prefix = match.group(1)
            bold_text = match.group(2)
            content = match.group(3)
            result['content'].append({
                'type': 'paragraph',
                'text': content,
                'bold_prefix': prefix + bold_text + '。'
            })
            continue

        # 模式2: **前缀** 内容
        if line.startswith('**') and '**' in line[2:]:
            end_pos = line.index('**', 2)
            bold_prefix = line[2:end_pos]
            text = line[end_pos + 2:].strip()
            result['content'].append({
                'type': 'paragraph',
                'text': text,
                'bold_prefix': bold_prefix
            })
            continue

        # 列表项
        if line.startswith('- '):
            result['content'].append({'type': 'paragraph', 'text': line[2:].strip()})
            continue

        # 普通段落
        result['content'].append({'type': 'paragraph', 'text': line})

    return result


# ============================================================
# CLI 入口
# ============================================================
def main():
    parser = argparse.ArgumentParser(description='TDoc DOCX 文档创建引擎')
    parser.add_argument('--title', help='文档标题')
    parser.add_argument('--author', help='署名（可用 \\n 换行）')
    parser.add_argument('--from-markdown', help='从 Markdown 文件创建')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--style', default='default',
                        choices=['default', 'gov', 'business', 'academic'],
                        help='文档风格 (default/gov/business/academic)')

    args = parser.parse_args()

    creator = DocxCreator(style=args.style)

    if args.from_markdown:
        # Markdown 转换模式
        if not os.path.exists(args.from_markdown):
            print(f"❌ 文件不存在: {args.from_markdown}")
            sys.exit(1)

        data = parse_markdown(args.from_markdown)
        if data['title']:
            creator.add_title(data['title'])
        if args.author:
            creator.add_author(args.author)
            creator.add_empty_line()
        for item in data['content']:
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            if item_type == 'heading1':
                creator.add_heading1(text)
            elif item_type == 'heading2':
                creator.add_heading2(text)
            elif item_type == 'paragraph':
                creator.add_paragraph(text, item.get('bold_prefix'))
        creator.save(args.output)
        return

    # 标准创建模式（仅标题 + 可选署名）
    if args.title:
        creator.add_title(args.title)

    if args.author:
        creator.add_author(args.author)

    creator.save(args.output)


if __name__ == '__main__':
    main()
