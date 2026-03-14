#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TDoc DOCX 格式转换引擎
支持 docx→pdf, docx→markdown, doc→docx, docx→images 转换。

用法:
  python convert_docx.py input.docx --to pdf --output output.pdf
  python convert_docx.py input.docx --to markdown --output output.md
  python convert_docx.py input.doc --to docx --output output.docx
"""

import argparse
import os
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================================
# 中文字体发现（跨平台）
# ============================================================
def find_chinese_font():
    """查找并注册中文字体"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return None

    font_candidates = []

    # macOS
    font_candidates.extend([
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
        ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
    ])

    # Windows
    font_candidates.extend([
        (r"C:\Windows\Fonts\msyh.ttc", "MicrosoftYaHei"),
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),
        (r"C:\Windows\Fonts\simsun.ttc", "SimSun"),
        (r"C:\Windows\Fonts\simkai.ttf", "KaiTi"),
    ])

    # Linux
    font_candidates.extend([
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WenQuanYiMicroHei"),
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", "DroidSansFallback"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
    ])

    for font_path, font_name in font_candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                return font_name
            except Exception:
                continue

    return None


# ============================================================
# DOCX → PDF
# ============================================================
def docx_to_pdf(input_file, output_file):
    """将 DOCX 转换为 PDF"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        print("❌ 缺少 reportlab 库。请安装: pip install reportlab")
        return False

    if not DOCX_AVAILABLE:
        print("❌ 缺少 python-docx 库。请安装: pip install python-docx")
        return False

    try:
        doc = Document(input_file)
        font_name = find_chinese_font() or 'Helvetica'

        pdf = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=36
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            fontSize=24, spaceAfter=12,
            fontName=font_name, alignment=TA_CENTER
        )
        h1_style = ParagraphStyle(
            'CustomH1', parent=styles['Heading1'],
            fontSize=18, spaceAfter=10,
            fontName=font_name
        )
        h2_style = ParagraphStyle(
            'CustomH2', parent=styles['Heading2'],
            fontSize=14, spaceAfter=8,
            fontName=font_name
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['Normal'],
            fontSize=11, spaceAfter=8, leading=16,
            fontName=font_name
        )

        story = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            # 转义 XML 特殊字符
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            style_name = para.style.name if para.style else ''
            if style_name == 'Title' or style_name == 'Heading 1':
                story.append(Paragraph(text, title_style if style_name == 'Title' else h1_style))
            elif style_name == 'Heading 2':
                story.append(Paragraph(text, h2_style))
            elif style_name.startswith('List'):
                story.append(Paragraph(f"• {text}", body_style))
            else:
                story.append(Paragraph(text, body_style))

        pdf.build(story)
        size_kb = os.path.getsize(output_file) / 1024
        print(f"✅ PDF 已生成: {output_file} ({size_kb:.1f} KB)")
        return True

    except Exception as e:
        print(f"❌ PDF 转换失败: {e}")
        return False


# ============================================================
# DOCX → Markdown
# ============================================================
def docx_to_markdown(input_file, output_file):
    """将 DOCX 转换为 Markdown"""
    if not DOCX_AVAILABLE:
        print("❌ 缺少 python-docx 库")
        return False

    try:
        doc = Document(input_file)
        md_lines = []

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                md_lines.append('')
                continue

            style_name = para.style.name if para.style else ''

            if style_name == 'Title':
                md_lines.append(f"# {text}")
            elif 'Heading 1' in style_name or style_name == 'Heading1':
                md_lines.append(f"## {text}")
            elif 'Heading 2' in style_name or style_name == 'Heading2':
                md_lines.append(f"### {text}")
            elif 'Heading 3' in style_name or style_name == 'Heading3':
                md_lines.append(f"#### {text}")
            elif 'List Bullet' in style_name:
                md_lines.append(f"- {text}")
            elif 'List Number' in style_name:
                md_lines.append(f"1. {text}")
            else:
                # 检查 run 级别的格式
                formatted_parts = []
                for run in para.runs:
                    t = run.text
                    if not t:
                        continue
                    if run.bold and run.italic:
                        formatted_parts.append(f"***{t}***")
                    elif run.bold:
                        formatted_parts.append(f"**{t}**")
                    elif run.italic:
                        formatted_parts.append(f"*{t}*")
                    else:
                        formatted_parts.append(t)
                md_lines.append(''.join(formatted_parts) if formatted_parts else text)

        # 表格
        for i, table in enumerate(doc.tables):
            md_lines.append('')
            if table.rows:
                header = table.rows[0]
                header_cells = [cell.text.strip() for cell in header.cells]
                md_lines.append('| ' + ' | '.join(header_cells) + ' |')
                md_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    md_lines.append('| ' + ' | '.join(cells) + ' |')
            md_lines.append('')

        md_text = '\n'.join(md_lines)

        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(md_text)

        print(f"✅ Markdown 已生成: {output_file}")
        return True

    except Exception as e:
        print(f"❌ Markdown 转换失败: {e}")
        return False


# ============================================================
# DOC → DOCX (via LibreOffice)
# ============================================================
def doc_to_docx(input_file, output_file):
    """将 .doc 转换为 .docx（需要 LibreOffice）"""
    try:
        output_dir = str(Path(output_file).parent or '.')
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'docx',
             '--outdir', output_dir, str(input_file)],
            capture_output=True, text=True, timeout=60
        )

        expected = Path(output_dir) / (Path(input_file).stem + '.docx')
        if expected.exists():
            if str(expected) != str(output_file):
                expected.rename(output_file)
            print(f"✅ DOCX 已生成: {output_file}")
            return True
        else:
            print(f"❌ 转换失败: LibreOffice 未生成输出文件")
            if result.stderr:
                print(f"   错误: {result.stderr[:200]}")
            return False

    except FileNotFoundError:
        print("❌ 需要 LibreOffice。安装方法:")
        print("   macOS: brew install --cask libreoffice")
        print("   Linux: sudo apt install libreoffice")
        return False
    except subprocess.TimeoutExpired:
        print("❌ LibreOffice 转换超时")
        return False


# ============================================================
# DOCX → Images (via LibreOffice + Poppler)
# ============================================================
def docx_to_images(input_file, output_dir):
    """将 DOCX 转换为图片（逐页）"""
    import tempfile

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # 先转 PDF
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', tmpdir, str(input_file)],
                capture_output=True, text=True, timeout=60
            )

            pdf_path = Path(tmpdir) / (Path(input_file).stem + '.pdf')
            if not pdf_path.exists():
                print("❌ PDF 中间转换失败")
                return False

            # PDF 转图片
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            output_prefix = str(Path(output_dir) / 'page')

            result = subprocess.run(
                ['pdftoppm', '-jpeg', '-r', '150', str(pdf_path), output_prefix],
                capture_output=True, text=True, timeout=120
            )

            images = list(Path(output_dir).glob('page-*.jpg'))
            if images:
                print(f"✅ 已生成 {len(images)} 张图片到: {output_dir}")
                return True
            else:
                print("❌ 图片生成失败。请确保安装了 Poppler: brew install poppler")
                return False

    except FileNotFoundError as e:
        print(f"❌ 缺少依赖: {e}")
        print("   需要: LibreOffice + Poppler (pdftoppm)")
        return False


# ============================================================
# CLI
# ============================================================
def main():
    parser = argparse.ArgumentParser(description='TDoc DOCX 格式转换引擎')
    parser.add_argument('input_file', help='输入文件路径')
    parser.add_argument('--to', required=True,
                        choices=['pdf', 'markdown', 'md', 'docx', 'images'],
                        help='目标格式')
    parser.add_argument('--output', '-o', help='输出路径')

    args = parser.parse_args()

    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)

    target = args.to.lower()
    if target == 'md':
        target = 'markdown'

    # 默认输出路径
    if not args.output:
        ext_map = {'pdf': '.pdf', 'markdown': '.md', 'docx': '.docx', 'images': '_pages'}
        ext = ext_map.get(target, '.out')
        args.output = str(input_path.with_suffix(ext))

    success = False
    if target == 'pdf':
        success = docx_to_pdf(args.input_file, args.output)
    elif target == 'markdown':
        success = docx_to_markdown(args.input_file, args.output)
    elif target == 'docx':
        success = doc_to_docx(args.input_file, args.output)
    elif target == 'images':
        success = docx_to_images(args.input_file, args.output)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
