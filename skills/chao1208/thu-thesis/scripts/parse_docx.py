#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
parse_docx.py — 清华 MBA 毕业论文 .docx → 结构化 JSON 解析器

用法：
    python3 parse_docx.py <input.docx> [output_dir]

输出：
    output_dir/parsed_<论文标题>.json

依赖：
    pip install python-docx
"""

import sys
import re
import json
import os
import shutil
import zipfile
import tempfile
from pathlib import Path
from typing import Any, Optional

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    print("错误：请先安装 python-docx：pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────
# 常量：特殊 Heading 1 章节名（这些不作为正文章节，而是特殊处理）
# ─────────────────────────────────────────────────────────────
SPECIAL_HEADINGS = {
    "摘要", "摘 要",
    "abstract",
    "目录", "目 录",
    "插图清单", "附表清单",
    "参考文献",
    "致谢", "致 谢",
    "声明", "声 明",
    "个人简历", "个人简历、在学期间完成的相关学术成果",
    "指导教师评语",
    "答辩委员会决议书",
}


def normalize_special_heading(text: str) -> str:
    """
    将特殊章节标题规范化（去空格、转小写）用于匹配。
    例如："摘 要" → "摘要"，"Abstract" → "abstract"
    """
    return re.sub(r'\s+', '', text).lower()


def is_special_heading(text: str) -> Optional[str]:
    """
    判断是否是特殊 Heading 1，返回规范化后的标准名称，否则返回 None。
    """
    normalized = normalize_special_heading(text)
    for sh in SPECIAL_HEADINGS:
        if normalize_special_heading(sh) == normalized:
            return sh
    return None


def is_chapter_heading(text: str) -> bool:
    """
    判断是否是正文章节标题（第X章 / 第XX章）。
    """
    return bool(re.match(r'^第[一二三四五六七八九十百\d]+章', text.strip()))


# ─────────────────────────────────────────────────────────────
# 文本清理
# ─────────────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    """
    基础文本清理：
    - 合并连续空白（空格、制表符、NBSP 等）为单个空格
    - 去除首尾空白
    - 不删除换行（段落内换行保留为空格）
    """
    if not text:
        return ""
    # 将各种空白字符（包括全角空格 \u3000、NBSP \xa0、制表符）统一为普通空格
    text = re.sub(r'[\t\u00a0\u3000]+', ' ', text)
    # 合并连续空格（但保留换行）
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def clean_spaced_chinese(text: str) -> str:
    """
    清理中文字符间人为插入的空格，如"吴 刘 驹" → "吴刘驹"。
    这种情况常见于封面的人名、单位名。
    策略：若连续的单汉字之间只有空格，则合并。
    """
    # 匹配「单个汉字 + 空格 + 单个汉字」模式，循环替换
    prev = None
    result = text
    while result != prev:
        prev = result
        result = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', result)
    return result.strip()


# ─────────────────────────────────────────────────────────────
# 封面信息提取
# ─────────────────────────────────────────────────────────────
def _extract_cover_from_tables(doc) -> dict:
    """
    从文档的封面表格中提取元信息（兜底方案）。
    支持两种格式：
      1. 字段名和值在同一单元格（"培养单位：经济管理学院"）
      2. 字段名在一列、值在相邻列（"指导教师" | "肖勇波教授"）
    """
    result = {}
    month_map = {
        "january":"01","february":"02","march":"03","april":"04",
        "may":"05","june":"06","july":"07","august":"08",
        "september":"09","october":"10","november":"11","december":"12"
    }
    FIELD_LABELS = {
        '培养单位':   'department',
        '培 养 单 位': 'department',
        '申请人':     'author',
        '申  请  人': 'author',
        '指导教师':   'supervisor',
        '指 导 教 师': 'supervisor',
    }

    for ti, tbl in enumerate(doc.tables):
        seen = set()
        cells = []
        row_cells = []
        for row in tbl.rows:
            # 每行去重，保留顺序（合并单元格会重复，只取第一次出现）
            row_seen = set()
            row_vals = []
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in row_seen:
                    row_seen.add(t)
                    row_vals.append(t)
                if t and t not in seen:
                    seen.add(t)
                    cells.append(t)
            if row_vals:
                row_cells.append(row_vals)

        if not cells:
            continue

        full_text = "\n".join(cells)

        # ── 按行扫描字段名+值对 ──
        for row_vals in row_cells:
            full_row = ' '.join(row_vals)
            for label, field in FIELD_LABELS.items():
                if result.get(field):
                    continue
                # 情况1：同一单元格含冒号
                m = re.search(re.escape(label) + r'\s*[：:]\s*(.+)', full_row)
                if m:
                    val = clean_spaced_chinese(m.group(1).strip())
                    if val:
                        result[field] = val
                else:
                    # 情况2：字段名在某列，值在后续列（跳过纯标点列）
                    for i, v in enumerate(row_vals):
                        if re.sub(r'\s+', '', v) == re.sub(r'\s+', '', label):
                            for nv in row_vals[i+1:]:
                                nv = nv.strip().lstrip('：:').strip()
                                if nv and nv not in ('：', ':', ''):
                                    result[field] = clean_spaced_chinese(nv)
                                    break
                            break

        # ── 导师：格式化为"姓名, 职称"，thuthesis 要求逗号分隔 ──
        if result.get('supervisor'):
            sv = result['supervisor']
            # 先去掉多余空格/逗号
            sv = re.sub(r'[,，\s]+', '', sv)
            title_m = re.search(r'(教授|副教授|研究员|副研究员|讲师|助理教授)', sv)
            if title_m:
                name = sv[:title_m.start()].strip()
                title = title_m.group(0)
                result['supervisor'] = f"{name}, {title}"
            else:
                result['supervisor'] = sv

        # ── 中文日期 ──
        if not result.get("date"):
            m = re.search(r'([一二三四五六七八九〇零○\d]+年[一二三四五六七八九十\d]+月)', full_text)
            if m:
                result["date"] = convert_chinese_date(m.group(1))

        # ── 中文标题 ──
        if not result.get("title"):
            for c in cells:
                c_clean = clean_text(c)
                cn_chars = len(re.sub(r'[^\u4e00-\u9fff]', '', c_clean))
                if (cn_chars >= 8
                        and not re.search(r'[：:]', c_clean)
                        and not c_clean.startswith(('(', '（'))
                        and c_clean not in ('培养单位', '申请人', '指导教师')):
                    result["title"] = c_clean
                    break

        # ── 英文标题 ──
        if not result.get("title_en"):
            for c in cells:
                if (len(c) > 30 and re.match(r'^[A-Z]', c)
                        and not re.match(r'(Thesis|Master|Doctor|by\b|Tsinghua|in partial)', c, re.IGNORECASE)):
                    result["title_en"] = clean_text(c.replace('\n', ' '))
                    break

        # ── 英文作者："by\nWU Jingjing" ──
        if not result.get("author_en"):
            for c in cells:
                m = re.search(r'by\s*\n\s*([A-Z][A-Za-z\s]+)', c)
                if m:
                    result["author_en"] = m.group(1).strip()
                    break

        # ── 英文导师 ──
        if not result.get("supervisor_en"):
            m = re.search(r'Thesis\s+Supervisor\s*[：:]\s*(.+)', full_text, re.IGNORECASE)
            if m:
                result["supervisor_en"] = m.group(1).strip().split('\n')[0]
            else:
                for row_vals in row_cells:
                    for i, v in enumerate(row_vals):
                        if re.sub(r'\s+', '', v.lower()) in ('thesissupervisor', 'supervisor'):
                            for nv in row_vals[i+1:]:
                                nv = nv.strip().lstrip('：:').strip()
                                if nv and re.match(r'(Professor|Prof\.)', nv, re.IGNORECASE):
                                    result["supervisor_en"] = nv
                                    break

        # ── 英文日期 ──
        if not result.get("date"):
            m = re.search(r'(January|February|March|April|May|June|July|August|'
                          r'September|October|November|December)[,，]\s*(\d{4})',
                          full_text, re.IGNORECASE)
            if m:
                result["date"] = f"{m.group(2)}-{month_map[m.group(1).lower()]}"

    return result


def parse_cover(paragraphs: list, doc=None) -> dict:
    """
    从文档开头（Heading 1 "摘要" 之前）的段落中提取封面元信息。

    中文封面格式示例：
        培 养 单 位：经济管理学院申 请 人： 吴 刘 驹
        指 导 教 师：谢    伟    教 授
        二〇二五年六月

    英文封面格式示例：
        Enterprise Development Strategy Study...
        by
        WU Liuju
        Thesis Supervisor : Professor Xie Wei
        June, 2025
    """
    meta = {
        "title": "",
        "title_en": "",
        "author": "",
        "author_en": "",
        "supervisor": "",
        "supervisor_en": "",
        "department": "",
        "degree_category": "工商管理硕士",
        "degree_category_en": "Master of Business Administration",
        "discipline": "工商管理",
        "discipline_en": "Business Administration",
        "date": "",
    }

    # 收集封面段落（Heading 1 "摘要" 之前）
    cover_paras = []
    for p in paragraphs:
        style = p.style.name
        txt = p.text.strip()
        # 遇到摘要标题就停止
        if style == "Heading 1" and normalize_special_heading(txt) in {"摘要", "摘 要", "abstract"}:
            break
        # 遇到正文章节也停止
        if style == "Heading 1" and is_chapter_heading(txt):
            break
        cover_paras.append(p)

    # ── 提取论文标题（Title 样式）──
    for p in cover_paras:
        if p.style.name == "Title" and p.text.strip():
            meta["title"] = clean_text(p.text)
            break

    # ── 逐段解析封面字段 ──
    # 用于记录英文封面状态机
    en_cover_state = "before"   # before → title_seen → by_seen → author_done
    en_title_lines = []

    for p in cover_paras:
        raw = p.text.strip()
        if not raw:
            continue
        style = p.style.name
        cleaned = clean_text(raw)

        if style == "Title":
            continue  # 已处理

        if style == "Normal":
            # ── 中文封面字段：用冒号（全角/半角）分割 ──
            # 例："培 养 单 位：经济管理学院申 请 人： 吴 刘 驹"
            # 这行可能把两个字段合在一起（因为换行变成空格），所以用正则逐个提取

            # 培养单位
            m = re.search(r'培\s*养\s*单\s*位[：:]\s*(.+?)(?=申\s*请\s*人|指\s*导\s*教\s*师|$)', cleaned)
            if m:
                dept = clean_spaced_chinese(m.group(1))
                # 去掉可能粘连的下一个字段关键字
                dept = re.sub(r'(申请人|指导教师).*', '', dept).strip()
                if dept:
                    meta["department"] = dept

            # 申请人
            m = re.search(r'申\s*请\s*人\s*[：:]\s*(.+?)(?=指\s*导\s*教\s*师|培\s*养|$)', cleaned)
            if m:
                author = clean_spaced_chinese(m.group(1)).strip()
                if author:
                    meta["author"] = author

            # 指导教师（格式：谢    伟    教 授 或 谢伟 教授）
            m = re.search(r'指\s*导\s*教\s*师[：:]\s*(.+?)(?=论文|$)', cleaned)
            if m:
                sv_raw = m.group(1).strip()
                # 先合并汉字间空格，再处理职称
                sv_name_part = re.sub(r'\s+', ' ', sv_raw)
                # 分离姓名和职称
                title_match = re.search(r'(教授|副教授|研究员|副研究员|讲师|助理教授)', sv_name_part)
                if title_match:
                    sv_name = sv_name_part[:title_match.start()].strip()
                    sv_name = clean_spaced_chinese(sv_name)
                    sv_title = title_match.group(0)
                    meta["supervisor"] = f"{sv_name} {sv_title}"
                else:
                    meta["supervisor"] = clean_spaced_chinese(sv_name_part)

            # 中文日期（如"二〇二五年六月"）
            m = re.search(r'([一二三四五六七八九〇零\d]+年[一二三四五六七八九十\d]+月)', cleaned)
            if m and not meta["date"]:
                meta["date"] = convert_chinese_date(m.group(1))

            # 英文日期（如"June, 2025"）
            m = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)[,，]\s*(\d{4})', cleaned, re.IGNORECASE)
            if m and not meta["date"]:
                month_map = {
                    "january": "01", "february": "02", "march": "03", "april": "04",
                    "may": "05", "june": "06", "july": "07", "august": "08",
                    "september": "09", "october": "10", "november": "11", "december": "12"
                }
                month = month_map.get(m.group(1).lower(), "06")
                meta["date"] = f"{m.group(2)}-{month}"

            # 英文论文标题（通常是比较长的英文段落，在英文封面开头）
            if re.match(r'^[A-Z]', cleaned) and len(cleaned) > 30 and not meta["title_en"]:
                # 排除"Thesis submitted to"、"Master of"等套话
                if not re.match(r'(Thesis|Master|Doctor|by\b|Tsinghua|in partial)', cleaned, re.IGNORECASE):
                    en_title_lines.append(cleaned)

            # 英文导师（"Thesis Supervisor : Professor Xie Wei"）
            m = re.search(r'Thesis\s+Supervisor\s*[：:]\s*(.+)', cleaned, re.IGNORECASE)
            if m:
                meta["supervisor_en"] = m.group(1).strip()

            # 英文作者（"by" 后面的下一个 Normal 段落）
            if cleaned.lower() == "by":
                en_cover_state = "by_seen"
                continue
            if en_cover_state == "by_seen" and re.match(r'^[A-Z][A-Za-z\s]+$', cleaned):
                meta["author_en"] = cleaned.strip()
                en_cover_state = "author_done"

    # 合并英文标题行（可能分两行）
    if en_title_lines and not meta["title_en"]:
        meta["title_en"] = " ".join(en_title_lines)

    # ── 兜底：从封面表格提取（应对封面全在表格里的文档）──
    if doc is not None:
        tbl_meta = _extract_cover_from_tables(doc)
        for k, v in tbl_meta.items():
            if v and not meta.get(k):
                meta[k] = v

    return meta


def convert_chinese_date(cn_date: str) -> str:
    """
    将中文日期转换为 YYYY-MM 格式。
    例："二〇二五年六月" → "2025-06"
    """
    cn_num = {"〇": "0", "零": "0", "○": "0", "一": "1", "二": "2", "三": "3",
              "四": "4", "五": "5", "六": "6", "七": "7", "八": "8", "九": "9"}
    month_map = {"一": "01", "二": "02", "三": "03", "四": "04", "五": "05",
                 "六": "06", "七": "07", "八": "08", "九": "09", "十": "10",
                 "十一": "11", "十二": "12"}

    m = re.match(r'([一二三四五六七八九〇零○\d]+)年([一二三四五六七八九十\d]+)月', cn_date)
    if not m:
        return cn_date

    year_cn = m.group(1)
    month_cn = m.group(2)

    # 年份：逐字替换中文数字
    year_str = year_cn
    for k, v in cn_num.items():
        year_str = year_str.replace(k, v)
    # 如果已经是数字则直接用
    try:
        year = int(year_str)
    except ValueError:
        year = 2025

    # 月份
    month = month_map.get(month_cn, f"{int(month_cn):02d}" if month_cn.isdigit() else "06")

    return f"{year}-{month}"


# ─────────────────────────────────────────────────────────────
# 摘要提取
# ─────────────────────────────────────────────────────────────
def parse_abstracts(paragraphs: list) -> dict:
    """
    提取中英文摘要和关键词。

    规则：
    - 中文摘要：Heading 1 "摘要/摘 要" 之后的 Body Text，直到关键词行
    - 英文摘要：Heading 1 "Abstract" 之后的 Body Text，直到关键词行
    - 关键词行：包含 "关键词：" / "关键字：" / "Keywords:"
    """
    result = {
        "abstract_cn": "",
        "keywords_cn": [],
        "abstract_en": "",
        "keywords_en": [],
    }

    # 状态机：None / "cn" / "en"
    state = None
    cn_lines = []
    en_lines = []

    # 预扫描：找到 Normal 样式的 "Abstract" 行和 Heading 1 "Abstract" 行的位置
    # 有些文档中，英文摘要正文出现在 Normal "Abstract" 之后、Heading 1 "Abstract" 之前
    # 需要同时支持这两种情况
    for p in paragraphs:
        style = p.style.name
        txt = p.text.strip()
        cleaned = clean_text(txt)

        # Normal 样式的 "Abstract"（也触发英文摘要状态）
        if style == "Normal" and normalize_special_heading(txt) == "abstract":
            state = "en"
            continue

        if style == "Heading 1":
            norm = normalize_special_heading(txt)
            if norm in {"摘要"}:
                state = "cn"
                continue
            elif norm == "abstract":
                # Heading 1 "Abstract"：如果已经在收集英文摘要，继续；否则开始
                if state != "en":
                    state = "en"
                continue
            else:
                # 遇到其他章节，停止摘要收集
                if state is not None:
                    state = None

        if style == "Heading 4":
            # 英文摘要的关键词可能是 Heading 4 样式
            if re.match(r'keywords\s*[:：]', txt, re.IGNORECASE):
                kw_str = re.sub(r'^keywords\s*[:：]\s*', '', txt, flags=re.IGNORECASE)
                result["keywords_en"] = [k.strip() for k in re.split(r'[;；,，]', kw_str) if k.strip()]
                state = None
            continue

        if not cleaned:
            continue

        if state == "cn":
            # 检测关键词行
            if re.match(r'关键[词字]\s*[：:]', cleaned):
                kw_str = re.sub(r'^关键[词字]\s*[：:]\s*', '', cleaned)
                result["keywords_cn"] = [k.strip() for k in re.split(r'[;；、,，]', kw_str) if k.strip()]
                state = None
                continue
            if style in ("Body Text", "Normal", "段落"):
                cn_lines.append(cleaned)

        elif state == "en":
            # 检测英文关键词行
            if re.match(r'keywords\s*[:：]', cleaned, re.IGNORECASE):
                kw_str = re.sub(r'^keywords\s*[:：]\s*', '', cleaned, flags=re.IGNORECASE)
                result["keywords_en"] = [k.strip() for k in re.split(r'[;；,，]', kw_str) if k.strip()]
                state = None
                continue
            if style in ("Body Text", "Normal", "段落"):
                en_lines.append(cleaned)

    result["abstract_cn"] = " ".join(cn_lines)
    result["abstract_en"] = " ".join(en_lines)
    return result


# ─────────────────────────────────────────────────────────────
# 表格转文本辅助
# ─────────────────────────────────────────────────────────────
def table_to_text(tbl) -> str:
    """
    将 docx 表格转换为简单的文本表示（Markdown 风格）。
    """
    rows = []
    for row in tbl.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        # 去重相邻重复单元格（合并单元格问题）
        deduped = []
        prev = None
        for c in cells:
            if c != prev:
                deduped.append(c)
            prev = c
        rows.append(" | ".join(deduped))
    return "\n".join(rows)


# ─────────────────────────────────────────────────────────────
# 主章节结构解析
# ─────────────────────────────────────────────────────────────
def parse_body(doc) -> dict:
    """
    解析正文章节结构、参考文献、致谢、声明、个人简历等。

    返回：
    {
        "chapters": [...],
        "references": [...],
        "acknowledgements": "...",
        "statement": "...",
        "resume": "...",
        "figures_list": bool,
        "tables_list": bool,
        "supervisor_comment": "...",
        "defense_resolution": "...",
    }
    """
    # 为了正确处理表格在段落中的位置，需要遍历文档 XML body 中的所有元素
    # docx 的段落和表格在 doc.element.body 中交替出现
    body_elements = list(doc.element.body)

    # 建立段落对象索引（段落 XML element → 段落对象）
    para_map = {p._element: p for p in doc.paragraphs}
    # 建立表格对象索引
    tbl_map = {t._element: t for t in doc.tables}

    result = {
        "chapters": [],
        "references": [],
        "acknowledgements": "",
        "statement": "",
        "resume": "",
        "figures_list": False,
        "tables_list": False,
        "supervisor_comment": "",
        "defense_resolution": "",
    }

    # 状态：当前处于哪个特殊区域
    # None / "chapter" / "references" / "acknowledgements" / "statement" / "resume"
    # / "supervisor_comment" / "defense_resolution"
    state = None

    # 当前章节（level 1）
    current_chapter: Optional[dict] = None
    # 当前 section（level 2 or 3），用于将 Body Text 归属到正确的节
    current_section_content: Optional[list] = None  # 指向当前最深层的 content list

    def new_chapter(level1_text: str) -> dict:
        """创建新的一级章节对象"""
        # 分离章节号和标题："第一章 研究背景与意义" → ("第一章", "研究背景与意义")
        m = re.match(r'^(第[一二三四五六七八九十百\d]+章)\s*(.*)$', level1_text.strip())
        if m:
            number = m.group(1)
            title = m.group(2).strip()
        else:
            number = ""
            title = level1_text.strip()
        return {"level": 1, "number": number, "title": title, "content": []}

    def new_section(level: int, heading_text: str) -> dict:
        """创建新的二/三级节对象（加入父章节的 content）"""
        # 分离编号和标题："1.1 研究背景" → ("1.1", "研究背景")
        m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', heading_text.strip())
        if m:
            number = m.group(1)
            title = m.group(2).strip()
        else:
            number = ""
            title = heading_text.strip()
        return {"type": "section", "level": level, "number": number, "title": title}

    special_text_buffers = {
        "acknowledgements": [],
        "statement": [],
        "resume": [],
        "supervisor_comment": [],
        "defense_resolution": [],
    }

    for body_idx, elem in enumerate(body_elements):
        tag = elem.tag.split('}')[-1]  # 去命名空间，得到 'p' 或 'tbl'

        # ── 处理段落 ──
        if tag == 'p' and elem in para_map:
            p = para_map[elem]
            style = p.style.name
            txt = p.text.strip()
            cleaned = clean_text(txt)

            # ── 检测图片段落 ──
            blip_elems = p._element.findall('.//' + qn('a:blip'))
            if blip_elems and state == "chapter" and current_section_content is not None:
                # 提取 rId（embed 属性）
                embed_ids = []
                for blip in blip_elems:
                    embed = blip.get(qn('r:embed'), '')
                    if embed:
                        embed_ids.append(embed)
                if embed_ids:
                    # caption 从 _figure_captions 里查（由 parse_docx 在解析后填充）
                    # 这里先放 embed，caption 在后期填充
                    for embed_id in embed_ids:
                        current_section_content.append({
                            'type': 'figure',
                            'embed': embed_id,
                            'caption': '',  # 后期由 parse_docx 填充
                        })
                continue  # 图片段落不再当文本处理

            # Heading 1：判断是特殊章节还是正文章节
            if style == "Heading 1" and txt:
                norm = normalize_special_heading(txt)

                if norm == "插图清单":
                    result["figures_list"] = True
                    state = "skip"
                    continue

                if norm == "附表清单":
                    result["tables_list"] = True
                    state = "skip"
                    continue

                if norm in {"摘要", "abstract", "目录"}:
                    state = "skip"
                    current_chapter = None
                    current_section_content = None
                    continue

                if norm == "参考文献":
                    state = "references"
                    current_chapter = None
                    current_section_content = None
                    continue

                if norm in {"致谢"}:
                    state = "acknowledgements"
                    continue

                if norm in {"声明"}:
                    state = "statement"
                    continue

                if norm in {"个人简历", "个人简历、在学期间完成的相关学术成果"}:
                    state = "resume"
                    continue

                if norm == "指导教师评语":
                    state = "supervisor_comment"
                    continue

                if norm == "答辩委员会决议书":
                    state = "defense_resolution"
                    continue

                # 正文章节（第X章）
                if is_chapter_heading(txt):
                    state = "chapter"
                    current_chapter = new_chapter(txt)
                    current_section_content = current_chapter["content"]
                    result["chapters"].append(current_chapter)
                    continue

                # 其他 Heading 1（不认识的，作为新章节处理）
                state = "chapter"
                current_chapter = {"level": 1, "number": "", "title": cleaned, "content": []}
                current_section_content = current_chapter["content"]
                result["chapters"].append(current_chapter)
                continue

            # Heading 2：二级节
            if style == "Heading 2" and txt and state == "chapter" and current_chapter is not None:
                sec = new_section(2, txt)
                current_chapter["content"].append(sec)
                current_section_content = current_chapter["content"]
                continue

            # Heading 3：三级节
            if style == "Heading 3" and txt and state == "chapter" and current_chapter is not None:
                sec = new_section(3, txt)
                current_chapter["content"].append(sec)
                current_section_content = current_chapter["content"]
                continue

            # 正文内容（Body Text / Normal 段落按状态归属）
            if not cleaned:
                continue

            if state == "references":
                if style in ("Normal", "Body Text", "段落"):
                    result["references"].append(cleaned)

            elif state in special_text_buffers:
                if style in ("Body Text", "Normal", "List Paragraph", "段落"):
                    special_text_buffers[state].append(cleaned)

            elif state == "chapter" and current_section_content is not None:
                if style in ("Body Text", "Normal", "List Paragraph", "段落"):
                    current_section_content.append({"type": "text", "content": cleaned})

        # ── 处理表格 ──
        elif tag == 'tbl' and elem in tbl_map:
            tbl = tbl_map[elem]
            if state == "chapter" and current_section_content is not None:
                # 提取结构化行列数据（带去重合并单元格）
                tbl_rows = []
                for row in tbl.rows:
                    cells = [clean_text(cell.text) for cell in row.cells]
                    # 去重相邻重复单元格（合并单元格问题）
                    deduped = []
                    prev_cell = None
                    for c in cells:
                        if c != prev_cell:
                            deduped.append(c)
                        prev_cell = c
                    tbl_rows.append(deduped)

                if not tbl_rows or len(tbl_rows) < 1:
                    pass  # 空表格跳过
                else:
                    # 第一行作为表头（如果非空）
                    headers = tbl_rows[0] if tbl_rows else []
                    data_rows = tbl_rows[1:] if len(tbl_rows) > 1 else []

                    # 查找表格 caption：
                    # 先查前面（body element 前3项），再查后面（body element 后3项）
                    tbl_caption = ''
                    tbl_body_idx = body_idx  # 当前表格在 body_elements 中的位置

                    # 前面找
                    for bi in range(max(0, tbl_body_idx - 3), tbl_body_idx):
                        bel = body_elements[bi]
                        if bel in para_map:
                            pt = para_map[bel].text.strip()
                            if pt.startswith('表') or para_map[bel].style.name == 'Caption1':
                                tbl_caption = pt
                                break

                    # 后面找（优先级低于前面，且只有前面没找到时才用）
                    if not tbl_caption:
                        for bi in range(tbl_body_idx + 1, min(len(body_elements), tbl_body_idx + 4)):
                            bel = body_elements[bi]
                            if bel in para_map:
                                pt = para_map[bel].text.strip()
                                if pt.startswith('表') or para_map[bel].style.name == 'Caption1':
                                    tbl_caption = pt
                                    break

                    current_section_content.append({
                        'type': 'table',
                        'caption': tbl_caption,
                        'headers': headers,
                        'rows': data_rows,
                    })
            elif state == "references":
                # 表格里的参考文献（极少见）
                tbl_text = table_to_text(tbl)
                result["references"].append(f"[TABLE]\n{tbl_text}")

    # 整合特殊文本
    result["acknowledgements"] = "\n".join(special_text_buffers["acknowledgements"])
    result["statement"] = "\n".join(special_text_buffers["statement"])
    result["resume"] = "\n".join(special_text_buffers["resume"])
    result["supervisor_comment"] = "\n".join(special_text_buffers["supervisor_comment"])
    result["defense_resolution"] = "\n".join(special_text_buffers["defense_resolution"])

    return result


# ─────────────────────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────────────────────
def parse_docx(input_path: str, output_dir: str = "output") -> dict:
    """
    解析 .docx 文件，返回结构化字典，并保存为 JSON。

    参数：
        input_path: .docx 文件路径
        output_dir: 输出目录（默认 "output"）

    返回：
        解析结果字典
    """
    input_path = Path(input_path).resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"文件不存在：{input_path}")
    if input_path.suffix.lower() != ".docx":
        raise ValueError(f"不支持的文件格式，需要 .docx：{input_path}")

    print(f"\n{'='*60}")
    print(f"解析文件：{input_path.name}")
    print(f"{'='*60}")

    # ── 加载文档 ──
    print("正在加载文档...", end=" ", flush=True)
    doc = docx.Document(str(input_path))
    paragraphs = doc.paragraphs
    print(f"完成（{len(paragraphs)} 段落，{len(doc.tables)} 表格）")

    # ── 统计段落样式分布 ──
    style_counts: dict[str, int] = {}
    for p in paragraphs:
        sn = p.style.name
        style_counts[sn] = style_counts.get(sn, 0) + 1
    print("\n段落样式分布（前10）：")
    for style, count in sorted(style_counts.items(), key=lambda x: -x[1])[:10]:
        print(f"  {style:<25} {count} 段")

    # ── 提取封面元信息 ──
    print("\n正在提取封面信息...", end=" ", flush=True)
    meta = parse_cover(paragraphs, doc=doc)
    print("完成")

    # ── 提取摘要 ──
    print("正在提取摘要...", end=" ", flush=True)
    abstracts = parse_abstracts(paragraphs)
    print("完成")

    # ── 解析正文章节 ──
    print("正在解析正文章节结构...", end=" ", flush=True)
    body = parse_body(doc)
    print("完成")

    # ── 提取图片文件及 caption 映射 ──
    print("正在提取图片...", end=" ", flush=True)
    figures_dir = Path(output_dir) / 'figures'
    figures_dir.mkdir(parents=True, exist_ok=True)

    # 建立 rId → 图片文件名映射（从 document.xml.rels 中读取）
    rid_to_filename: dict[str, str] = {}
    extract_tmp = Path(tempfile.mkdtemp())
    try:
        with zipfile.ZipFile(str(input_path), 'r') as z:
            # 解析关系文件，建立 rId → target 映射
            rels_name = 'word/_rels/document.xml.rels'
            if rels_name in z.namelist():
                import xml.etree.ElementTree as ET
                rels_xml = z.read(rels_name)
                rels_root = ET.fromstring(rels_xml)
                ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                for rel in rels_root.findall('r:Relationship', ns):
                    rid = rel.get('Id', '')
                    target = rel.get('Target', '')
                    rtype = rel.get('Type', '')
                    if 'image' in rtype.lower() and target:
                        fname = Path(target).name
                        rid_to_filename[rid] = fname

            # 提取图片文件
            for name in z.namelist():
                if name.startswith('word/media/'):
                    fname = Path(name).name
                    data_bytes = z.read(name)
                    (figures_dir / fname).write_bytes(data_bytes)
    finally:
        import shutil as _shutil
        _shutil.rmtree(str(extract_tmp), ignore_errors=True)

    img_count = len(list(figures_dir.iterdir()))
    print(f"完成（{img_count} 张图片 → {figures_dir}）")

    # ── 建立图片 rId → {filename, caption} 映射 ──
    # 遍历 doc.paragraphs，找含 a:blip 的段落，按规则提取 caption
    figures_map: dict[str, dict] = {}
    all_paras = doc.paragraphs
    para_count = len(all_paras)
    for idx, p in enumerate(all_paras):
        blip_elems = p._element.findall('.//' + qn('a:blip'))
        if not blip_elems:
            continue
        for blip in blip_elems:
            embed = blip.get(qn('r:embed'), '')
            if not embed:
                continue
            filename = rid_to_filename.get(embed, '')
            caption = ''

            # 规则0：图片段落自身文本以"图"开头（图片和标题在同一段落）
            self_text = clean_text(p.text)
            if self_text.startswith('图'):
                caption = self_text

            # 规则1：图片段落之后紧跟 Caption1 样式的段落
            if not caption and idx + 1 < para_count:
                next_p = all_paras[idx + 1]
                if next_p.style.name == 'Caption1':
                    caption = clean_text(next_p.text)

            # 规则2：图片段落之后2段内有 Normal 样式且文本以"图"开头的段落
            if not caption:
                for fwd in range(1, 3):
                    if idx + fwd < para_count:
                        next_p = all_paras[idx + fwd]
                        if next_p.style.name in ('Normal', 'Caption1') and next_p.text.strip().startswith('图'):
                            caption = clean_text(next_p.text)
                            break

            # 规则3：图片段落之前4段内有 Caption1 或 Normal 样式且文本以"图"开头的段落
            if not caption:
                for back_idx in range(max(0, idx - 4), idx):
                    prev_p = all_paras[back_idx]
                    if prev_p.style.name == 'Caption1':
                        caption = clean_text(prev_p.text)
                        break
                    if prev_p.style.name == 'Normal' and prev_p.text.strip().startswith('图'):
                        caption = clean_text(prev_p.text)
                        break

            figures_map[embed] = {'filename': filename, 'caption': caption}

    # ── 用 caption 映射回填 parse_body 输出的 figure 项 ──
    for chapter in body['chapters']:
        for item in chapter.get('content', []):
            if item.get('type') == 'figure':
                embed = item.get('embed', '')
                if embed in figures_map:
                    if not item.get('caption'):
                        item['caption'] = figures_map[embed].get('caption', '')

    # ── 组装最终结果 ──
    result = {
        "meta": meta,
        "abstract_cn": abstracts["abstract_cn"],
        "keywords_cn": abstracts["keywords_cn"],
        "abstract_en": abstracts["abstract_en"],
        "keywords_en": abstracts["keywords_en"],
        "chapters": body["chapters"],
        "references": body["references"],
        "acknowledgements": body["acknowledgements"],
        "resume": body["resume"],
        "statement": body["statement"],
        "supervisor_comment": body["supervisor_comment"],
        "defense_resolution": body["defense_resolution"],
        "figures_list": body["figures_list"],
        "tables_list": body["tables_list"],
        "figures": figures_map,
    }

    # ── 保存 JSON ──
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    title_slug = meta["title"] or input_path.stem
    # 文件名安全处理：去掉不合法字符
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title_slug)[:50]
    out_file = output_path / f"parsed_{safe_title}.json"

    with open(out_file, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n✓ JSON 已保存：{out_file}")

    # ── 打印统计摘要 ──
    print_summary(result)

    return result


def print_summary(result: dict):
    """打印解析结果的关键摘要信息"""
    meta = result["meta"]
    chapters = result["chapters"]
    refs = result["references"]
    abstract_cn = result["abstract_cn"]
    abstract_en = result["abstract_en"]

    print(f"\n{'─'*60}")
    print("📋 解析结果摘要")
    print(f"{'─'*60}")

    print("\n【元信息 meta】")
    for k, v in meta.items():
        print(f"  {k:<20} {v}")

    print(f"\n【章节结构】共 {len(chapters)} 章")
    for ch in chapters:
        n = ch.get("number", "")
        t = ch.get("title", "")
        items = len(ch.get("content", []))
        print(f"  {n} {t}  （{items} 个内容块）")

    print(f"\n【参考文献】共 {len(refs)} 条")
    if refs:
        print(f"  示例：{refs[0][:80]}")

    print(f"\n【中文摘要】共 {len(abstract_cn)} 字")
    if abstract_cn:
        print(f"  前100字：{abstract_cn[:100]}")

    print(f"\n【英文摘要】共 {len(abstract_en)} 字符")
    if abstract_en:
        print(f"  前100字符：{abstract_en[:100]}")

    print(f"\n【关键词-中文】{result['keywords_cn']}")
    print(f"【关键词-英文】{result['keywords_en']}")
    print(f"\n【插图清单】{'有' if result['figures_list'] else '无'}")
    print(f"【附表清单】{'有' if result['tables_list'] else '无'}")
    print(f"【致谢】{'有 (' + str(len(result['acknowledgements'])) + '字)' if result['acknowledgements'] else '无'}")
    print(f"【声明】{'有' if result['statement'] else '无'}")
    print(f"【个人简历】{'有' if result['resume'] else '无'}")
    print(f"{'─'*60}\n")


# ─────────────────────────────────────────────────────────────
# 命令行入口
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python3 parse_docx.py <input.docx> [output_dir]")
        print("示例：python3 parse_docx.py thesis.docx output/")
        sys.exit(1)

    input_file = sys.argv[1]
    output_directory = sys.argv[2] if len(sys.argv) > 2 else "output"

    try:
        parse_docx(input_file, output_directory)
    except Exception as e:
        print(f"\n❌ 解析失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
