# -*- coding: utf-8 -*-
"""
机关公文生成脚本
基于 GB/T 9704-2012 党政机关公文格式标准

注意：所有字符串使用三引号包裹，避免文本中的引号冲突
"""

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================
# 【行距常量】- 28.5磅
# ============================================
LINE_SPACING = Pt(28.5)

# ============================================
# 【配置区】- 修改以下内容生成不同文档
# ============================================

CONFIG = {
    # 发文字号（设为 None 则不显示）
    "docNumber": "陕XX〔2025〕XX号",

    # 公文标题
    "title": "关于做好全国模范单位奖牌颁发工作的通知",

    # 主送机关
    "mainRecipients": "各设区市、各高校：",

    # 署名
    "author": "陕西省XXXX",

    # 日期
    "date": "2025年X月X日",

    # 附件列表（设为空列表则不显示）
    "attachments": ["全国模范单位奖牌颁发情况统计表"],

    # 联系方式（在落款日期下方，用中文括号）
    "contact": """联系人：XXX 189XXXXXXXX，电子邮箱：XXXXXXXX@qq.com，邮寄地址：陕西省西安市XXXXXXXX""",

    # 是否显示页码
    "showPageNumber": True,

    # 输出文件名
    "outputFile": "公文.docx"
}

# ============================================
# 【正文内容区】- 在这里填入实际内容
# ============================================

CONTENTS = [
    # 格式说明：
    # ("h1", "一、一级标题")
    # ("h2", "（一）二级标题")
    # ("normal", "正文内容")
    # ("blank", None)  # 空行

    ("h1", "一、充分认识奖牌颁发工作的重要意义"),
    ("normal", "全国模范单位奖牌是表彰先进、树立典型的重要载体，做好奖牌颁发工作对于弘扬精神、激发会员积极性具有重要意义。各地要高度重视，精心组织，确保奖牌颁发工作顺利完成。"),
    ("h1", "二、获奖名单"),
    ("normal", "本次表彰共涉及我省XX家全国模范单位，涵盖社区、企业、学校等多种类型。"),
    ("h1", "三、领取颁发要求"),
    ("h2", "（一）领取方式"),
    ("normal", "获奖单位请于2025年X月X日前联系所在设区市领取奖牌。"),
    ("h2", "（二）颁发时间"),
    ("normal", "请于2025年X月X日前完成奖牌颁发工作，并填写奖牌领取确认表。"),
    ("h1", "四、工作要求"),
    ("normal", "各地要加强领导，明确责任，确保奖牌颁发工作规范有序。要以此次颁奖为契机，进一步推动事业发展。"),
    ("h1", "五、联系方式"),
    ("normal", "联系人：XXX    联系电话：XXXX-XXXXXXXX"),
]

# ============================================
# 【辅助函数】
# ============================================
# 【引号标准化】- 将 ASCII 直引号还原为中文弯引号
# 原因：AI 在生成脚本时容易将用户输入的中文弯引号转为 ASCII 直引号，
#       本函数在文本写入 docx 前自动还原，确保公文格式正确。
# ============================================
LEFT_DOUBLE = "\u201C"   # " 中文左双弯引号
RIGHT_DOUBLE = "\u201D"  # " 中文右双弯引号
LEFT_SINGLE = "\u2018"  # ' 中文左单弯引号
RIGHT_SINGLE = "\u2019" # ' 中文右单弯引号

def normalize_chinese_quotes(text):
    """
    将文本中的 ASCII 直引号交替还原为中文弯引号。

    原理：ASCII 直引号 " 和 ' 是对称字符（同一字符既是左引号也是右引号），
    无法直接区分左右。因此采用交替替换策略：第1、3、5...个双引号 → 左双引号；
    第2、4、6...个双引号 → 右双引号。单引号同理，独立计数。
    这是中文排版中配对引号的标准做法。
    """
    if not text:
        return text

    result = []
    double_count = 0  # 双引号计数（独立）
    single_count = 0  # 单引号计数（独立）
    i = 0
    while i < len(text):
        c = text[i]
        if c == '"':
            if double_count % 2 == 0:
                result.append(LEFT_DOUBLE)   # 奇数个双引号 → 左
            else:
                result.append(RIGHT_DOUBLE)  # 偶数个双引号 → 右
            double_count += 1
        elif c == "'":
            if single_count % 2 == 0:
                result.append(LEFT_SINGLE)   # 奇数个单引号 → 左
            else:
                result.append(RIGHT_SINGLE)  # 偶数个单引号 → 右
            single_count += 1
        else:
            result.append(c)
        i += 1
    return ''.join(result)

# ============================================

def set_run_font(run, font_name, font_size, bold=False):
    """设置文本字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def create_paragraph(doc, text, font_name="仿宋_GB2312", font_size=16,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Twips(640), bold=False,
                     space_before=0, space_after=0):
    """
    创建段落

    参数:
        font_name: 字体名称
        font_size: 字号（pt）
        alignment: 对齐方式
        first_line_indent: 首行缩进，Twips(640)=2个汉字，0表示无缩进
        bold: 是否加粗
    """
    p = doc.add_paragraph()
    # 用 paragraph_format.alignment 确保覆盖 Normal 样式
    p.paragraph_format.alignment = alignment

    # 设置段落间距
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # 设置首行缩进
    if first_line_indent and first_line_indent > 0:
        p.paragraph_format.first_line_indent = first_line_indent
    else:
        p.paragraph_format.first_line_indent = Cm(0)

    # 添加文本（自动还原中文弯引号）
    if text:
        text = normalize_chinese_quotes(text)
        run = p.add_run(text)
        set_run_font(run, font_name, font_size, bold)

    return p


def create_blank_line(doc):
    """创建空行"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_content(doc, item):
    """根据类型添加正文内容"""
    content_type, text = item

    if content_type == "blank":
        return create_blank_line(doc)

    elif content_type == "h1":
        # 一级标题：黑体，不加粗
        return create_paragraph(doc, text,
                                font_name="黑体",
                                font_size=16,
                                first_line_indent=Twips(640),
                                bold=False)

    elif content_type == "h2":
        # 二级标题：楷体，不加粗
        return create_paragraph(doc, text,
                                font_name="楷体_GB2312",
                                font_size=16,
                                first_line_indent=Twips(640),
                                bold=False)

    elif content_type == "normal":
        # 正文：仿宋，首行缩进，左对齐
        return create_paragraph(doc, text,
                                font_name="仿宋_GB2312",
                                font_size=16,
                                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                first_line_indent=Twips(640),
                                bold=False)

    else:
        return create_paragraph(doc, text)


def set_page_margins(doc, top=Cm(3.7), bottom=Cm(3.5),
                     left=Cm(2.8), right=Cm(2.6)):
    """设置页边距"""
    sections = doc.sections
    for section in sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right


# ============================================
# 【主函数】
# ============================================

def generate_document():
    print("=== 机关公文生成器 ===")
    print("标题:", CONFIG["title"])
    print("发文字号:", CONFIG["docNumber"])
    print("主送机关:", CONFIG["mainRecipients"])
    print("输出:", CONFIG["outputFile"])
    print("====================\n")

    # 创建文档
    doc = Document()

    # 设置页边距（GB/T 9704-2012）
    set_page_margins(doc)

    # 1. 发文字号 - 右对齐
    if CONFIG["docNumber"]:
        create_paragraph(doc, CONFIG["docNumber"],
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=0)

    # 2. 发文字号和大标题之间空两行
    create_blank_line(doc)
    create_blank_line(doc)

    # 3. 标题 - 居中，方正小标宋，不加粗
    create_paragraph(doc, CONFIG["title"],
                     font_name="方正小标宋简体",
                     font_size=22,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     first_line_indent=0,
                     bold=False)

    # 4. 标题和主送机关之间空一行
    create_blank_line(doc)

    # 5. 主送机关（顶格，不缩进）
    if CONFIG["mainRecipients"]:
        create_paragraph(doc, CONFIG["mainRecipients"],
                        first_line_indent=0)

    # 6. 正文
    for item in CONTENTS:
        add_content(doc, item)

    # 7. 附件说明
    if CONFIG["attachments"] and len(CONFIG["attachments"]) > 0:
        create_blank_line(doc)
        if len(CONFIG["attachments"]) == 1:
            # 单个附件：附件：xxx
            create_paragraph(doc, "附件：" + CONFIG["attachments"][0],
                             first_line_indent=Twips(640))
        else:
            # 多个附件：附件：1.xxx 换行 2.xxx
            create_paragraph(doc, "附件：1." + CONFIG["attachments"][0],
                             first_line_indent=Twips(640))
            for i, att in enumerate(CONFIG["attachments"][1:], start=2):
                create_paragraph(doc, str(i) + "." + att,
                                 first_line_indent=Twips(640))

    # 8. 落款 - 前空3行，直接右对齐
    create_blank_line(doc)
    create_blank_line(doc)
    create_blank_line(doc)

    if CONFIG["author"]:
        create_paragraph(doc, CONFIG["author"],
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=0)

    if CONFIG["date"]:
        create_paragraph(doc, CONFIG["date"],
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=0)

    # 9. 联系方式 - 在落款日期下方空一行，中文括号，缩进2字符
    if CONFIG["contact"]:
        create_blank_line(doc)
        contact_text = "（" + CONFIG["contact"] + "）"
        create_paragraph(doc, contact_text,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Twips(640))

    # 保存文档
    doc.save(CONFIG["outputFile"])
    print("生成成功:", CONFIG["outputFile"])


if __name__ == "__main__":
    generate_document()
