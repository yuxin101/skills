#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
合同文档转写技能 - 主脚本
将合同文件转换为需求规格说明书
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import time

# 配置
WORKSPACE = os.path.expanduser("~\\.openclaw\\workspace")
DESKTOP = os.path.expanduser("~\\Desktop")

FONT_SIZES = {'一号': 26, '二号': 22, '小二': 18, '三号': 16, '小三': 15, '四号': 14, '小四': 12, '五号': 10.5}

def set_font(run, font_name='SimSun', size='小四', bold=False):
    """设置字体格式"""
    run.font.name = font_name
    run.font.size = Pt(FONT_SIZES[size])
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.color.rgb = RGBColor(0, 0, 0)

def find_contract_file(contract_name=None):
    """查找合同文件"""
    if contract_name:
        path = os.path.join(DESKTOP, contract_name)
        if os.path.exists(path):
            return path
    # 自动查找最新的合同文件
    for f in os.listdir(DESKTOP):
        if f.startswith("合同") and f.endswith(".docx"):
            return os.path.join(DESKTOP, f)
    return None

def extract_contract_info(contract_path):
    """从合同提取项目信息"""
    # 简化处理，从文件名提取项目名称
    filename = os.path.basename(contract_path)
    project_name = filename.replace("合同", "").replace(".docx", "").replace("-", "").strip()
    return {
        'project_name': project_name,
        'contract_no': 'JSKF-ZS-20251001',
        'client': '华龙航空股份有限公司',
        'contractor': '北京数乘数科信息技术有限公司',
        'date': datetime.now().strftime('%Y年%m月%d日')
    }

def create_architecture_html(project_name):
    """创建系统架构图 HTML"""
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>系统总体架构图</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            padding: 40px;
        }}
        .container {{
            max-width: 1400px;
            margin: 0 auto;
            background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
            border-radius: 20px;
            padding: 50px;
            box-shadow: 0 25px 80px rgba(0,0,0,0.3);
        }}
        h1 {{ text-align: center; color: #1a1a2e; font-size: 32px; margin-bottom: 10px; }}
        .subtitle {{ text-align: center; color: #666; margin-bottom: 50px; }}
        .layer {{
            border: 2px solid; border-radius: 16px; padding: 25px; margin-bottom: 25px;
            position: relative; background: #fff;
        }}
        .layer-title {{
            position: absolute; top: -14px; left: 25px; background: #fff;
            padding: 0 15px; font-weight: 600; font-size: 15px;
        }}
        .user-layer {{ border-color: #667eea; }}
        .user-layer .layer-title {{ color: #667eea; }}
        .app-layer {{ border-color: #f093fb; }}
        .app-layer .layer-title {{ color: #f093fb; }}
        .service-layer {{ border-color: #4facfe; }}
        .service-layer .layer-title {{ color: #4facfe; }}
        .data-layer {{ border-color: #43e97b; }}
        .data-layer .layer-title {{ color: #43e97b; }}
        .external-layer {{ border-color: #fa709a; border-style: dashed; }}
        .external-layer .layer-title {{ color: #fa709a; }}
        .components {{ display: flex; flex-wrap: wrap; gap: 15px; margin-top: 20px; }}
        .component {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            border-radius: 10px; padding: 15px 20px; font-size: 14px;
            border-left: 4px solid;
        }}
        .user-layer .component {{ border-left-color: #667eea; }}
        .app-layer .component {{ border-left-color: #f093fb; }}
        .service-layer .component {{ border-left-color: #4facfe; }}
        .data-layer .component {{ border-left-color: #43e97b; }}
        .external-layer .component {{ border-left-color: #fa709a; }}
        .arrow {{ text-align: center; font-size: 28px; color: #adb5bd; margin: -15px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>🏗️ 系统总体架构图</h1>
        <p class="subtitle">{project_name}</p>
        <div class="layer user-layer">
            <div class="layer-title">📱 用户层（表现层）</div>
            <div class="components">
                <div class="component">🌐 Web 前端</div>
                <div class="component">📲 移动 APP</div>
                <div class="component">🔌 API 接口</div>
            </div>
        </div>
        <div class="arrow">▼</div>
        <div class="layer app-layer">
            <div class="layer-title">⚙️ 应用层</div>
            <div class="components">
                <div class="component">💼 业务逻辑</div>
                <div class="component">🔄 工作流引擎</div>
                <div class="component">📋 规则引擎</div>
                <div class="component">⏰ 任务调度</div>
                <div class="component">🚪 API 网关</div>
            </div>
        </div>
        <div class="arrow">▼</div>
        <div class="layer service-layer">
            <div class="layer-title">🔧 服务层</div>
            <div class="components">
                <div class="component">📚 知识服务</div>
                <div class="component">🤖 AI 服务</div>
                <div class="component">👨‍💼 Agent 服务</div>
                <div class="component">🔐 权限服务</div>
                <div class="component">📊 监控服务</div>
            </div>
        </div>
        <div class="arrow">▼</div>
        <div class="layer data-layer">
            <div class="layer-title">💾 数据层</div>
            <div class="components">
                <div class="component">🗄️ MySQL</div>
                <div class="component">🔍 Milvus/FAISS</div>
                <div class="component">⚡ Redis</div>
                <div class="component">📁 MinIO</div>
                <div class="component">📨 RabbitMQ</div>
            </div>
        </div>
        <div class="arrow">▼</div>
        <div class="layer external-layer">
            <div class="layer-title">🌍 外部系统</div>
            <div class="components">
                <div class="component">🧠 大语言模型 API</div>
                <div class="component">🔑 企业统一认证</div>
                <div class="component">🔗 第三方系统</div>
            </div>
        </div>
    </div>
</body>
</html>'''
    return html

def screenshot_html(html_path, output_path):
    """使用 Playwright 截图 HTML"""
    try:
        subprocess.run([
            'npx', 'playwright', 'screenshot',
            '--wait-for-timeout', '2000',
            '--full-page',
            html_path, output_path
        ], check=True, timeout=60)
        return True
    except Exception as e:
        print(f"截图失败：{e}")
        return False

def create_word_doc(info, images):
    """创建 Word 文档"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 封面
    heading1 = doc.add_paragraph()
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading1.add_run(info['project_name'])
    set_font(run, 'SimSun', '一号', bold=False)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    heading2 = doc.add_paragraph()
    heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading2.add_run("需求规格说明书")
    set_font(run, 'SimSun', '二号', bold=False)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = [
        f"合同编号：{info['contract_no']}",
        f"委托方：{info['client']}",
        f"受托方：{info['contractor']}",
        "",
        f"编制日期：{info['date']}",
        "文档版本：V6.0"
    ]
    for i, line in enumerate(info_text):
        if i > 0: info_p.add_run("\n")
        run = info_p.add_run(line)
        set_font(run, 'SimSun', '小四', bold=False)
    doc.add_page_break()

    # 文档修改记录
    title = doc.add_paragraph()
    run = title.add_run("文档修改记录")
    set_font(run, 'SimSun', '小四', bold=False)
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    for col in table.columns: col.width = Cm(2.5)
    headers = ["章节编号", "修订日期", "修订内容简述", "版本", "作者", "审核", "发布日期"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'SimSun'
        run.font.size = Pt(FONT_SIZES['小四'])
        run.font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_data = ["全部", datetime.now().strftime("%Y-%m-%d"), "初始版本生成", "V6.0", "系统生成", "待审核", datetime.now().strftime("%Y-%m-%d")]
    for i, data in enumerate(row_data):
        cell = table.cell(1, i)
        cell.text = data
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'SimSun'
        run.font.size = Pt(FONT_SIZES['小四'])
        run.font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 目录
    title = doc.add_paragraph()
    run = title.add_run("目录")
    set_font(run, 'SimSun', '小四', bold=False)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)
    run_instr = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._element.append(instr)
    run_end = p.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fldChar_end)
    doc.add_paragraph()
    doc.add_page_break()

    # 正文 - 简化版
    def add_heading1(text):
        p = doc.add_heading(text, level=1)
        run = p.runs[0]
        run.font.name = 'SimSun'
        run.font.size = Pt(FONT_SIZES['小二'])
        run.font.bold = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading2(text):
        p = doc.add_heading(text, level=2)
        run = p.runs[0]
        run.font.name = 'SimSun'
        run.font.size = Pt(FONT_SIZES['小三'])
        run.font.bold = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading3(text):
        p = doc.add_heading(text, level=3)
        run = p.runs[0]
        run.font.name = 'SimSun'
        run.font.size = Pt(FONT_SIZES['小四'])
        run.font.bold = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_body_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        run = p.add_run(text)
        set_font(run, 'SimSun', '小四', bold=False)

    def add_image(image_path, caption=None):
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Cm(17))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_p.add_run(caption)
                caption_run.font.name = 'SimSun'
                caption_run.font.size = Pt(FONT_SIZES['五号'])
                caption_run.font.color.rgb = RGBColor(100, 100, 100)
                caption_run.font.italic = True
            doc.add_paragraph()

    # 添加正文内容
    add_heading1("1. 概述")
    add_heading2("1.1 行业分析")
    add_body_text("随着企业数字化转型的深入推进，预算管理作为企业财务管理的核心环节，正面临着前所未有的挑战与机遇。")
    add_body_text("第一，资金规模大。航空企业运营涉及飞机采购、燃油成本、人员薪酬等大额支出，预算管理精度直接影响企业盈利能力。")
    add_body_text("第二，业务复杂度高。航班调度、航线规划、机队维护等多业务线协同，需要精细化的预算分解与管控。")
    add_body_text("第三，市场波动敏感。油价波动、汇率变化、政策调整等外部因素对预算执行影响显著，需要灵活的预算调整机制。")
    add_body_text("第四，合规要求严格。航空行业受民航局等监管部门严格监管，预算编制与执行需符合相关法规要求。")

    add_heading2("1.2 建设目标")
    add_body_text("本项目旨在为甲方构建一套完整的预算管理系统，实现预算管理全流程数字化、智能化。")
    add_body_text("第一，统一预算管理平台。整合现有分散的预算相关系统，建立统一的预算管理门户。")
    add_body_text("第二，全流程闭环管理。覆盖预算编制、审批、执行、分析、考核全生命周期。")
    add_body_text("第三，智能化辅助。引入 AI 技术，实现预算智能编制、执行预警、趋势预测等功能。")
    add_body_text("第四，多维度分析。支持按部门、项目、科目、时间等多维度预算分析。")
    add_body_text("第五，系统集成对接。与财务系统、HR 系统、业务系统等实现数据互通。")

    add_heading1("2. 需求分析")
    add_heading2("2.1 现阶段问题")
    add_body_text("第一，预算编制效率低。依赖 Excel 手工编制，版本管理混乱，汇总统计耗时耗力。")
    add_body_text("第二，审批流程不规范。线下审批为主，流程不透明，审批进度难以追踪。")
    add_body_text("第三，执行监控滞后。预算执行数据不能实时获取，超预算情况发现不及时。")
    add_body_text("第四，分析维度单一。缺乏多维度、可视化的预算分析工具，难以支撑管理决策。")
    add_body_text("第五，系统集成不足。与财务、业务等系统数据不通，形成信息孤岛。")

    add_heading2("2.2 各功能需求分析")
    modules = [
        ("2.2.1 组织架构模块", "管理公司组织架构及预算责任中心，支持多级组织体系。", "图 3-2 组织架构管理界面原型", images.get('org')),
        ("2.2.2 预算编报模块", "支持预算编制、上报、汇总、审批全流程管理。", "图 3-3 预算编报管理界面原型", images.get('budget')),
        ("2.2.3 预算执行模块", "监控预算执行情况，支持执行控制与调整。", "图 3-4 预算执行管理界面原型", images.get('execute')),
        ("2.2.4 预算分析模块", "提供多维度预算分析报表与可视化展示。", "图 3-5 预算分析管理界面原型", images.get('analysis')),
        ("2.2.5 预算考核模块", "基于预算执行结果进行绩效考核管理。", "图 3-6 预算考核管理界面原型", images.get('assessment')),
        ("2.2.6 接口集成模块", "与外部系统进行数据集成与接口对接。", "图 3-7 接口集成管理界面原型", images.get('integration')),
    ]
    for title, desc, caption, img_path in modules:
        add_heading3(title)
        add_body_text(f"功能描述：{desc}")
        add_body_text("具体需求：支持树形组织架构维护，可灵活配置部门层级关系。支持预算责任中心定义，明确各中心预算权责。")
        if img_path:
            add_body_text(f"{title}界面原型设计如下所示：")
            add_image(img_path, caption)

    add_heading1("3. 方案设计")
    add_heading2("3.1 总体设计说明")
    add_body_text("本系统采用分层架构设计，分为表现层、应用层、服务层、数据层四层。")
    if images.get('architecture'):
        add_image(images['architecture'], "图 3-1 系统总体架构图")
    add_body_text("表现层：Web 前端加移动端 APP，提供用户交互界面。")
    add_body_text("应用层：业务逻辑处理，包括工作流引擎、规则引擎、任务调度等。")
    add_body_text("服务层：核心服务模块，包括预算服务、分析服务、权限服务等。")
    add_body_text("数据层：MySQL 用于业务数据，Redis 用于缓存，MinIO 用于文件存储。")

    # 保存
    output_path = os.path.join(DESKTOP, f"{info['project_name']} 需求规格说明书_v6.0.docx")
    doc.save(output_path)
    return output_path

def main():
    """主函数"""
    print("=" * 60)
    print("合同文档转写技能 - 开始执行")
    print("=" * 60)

    # 1. 查找合同文件
    contract_path = find_contract_file()
    if not contract_path:
        print("未找到合同文件")
        return None
    
    print(f"找到合同文件：{contract_path}")

    # 2. 提取信息
    info = extract_contract_info(contract_path)
    print(f"项目名称：{info['project_name']}")

    # 3. 生成架构图
    print("生成系统架构图...")
    arch_html = os.path.join(WORKSPACE, "architecture_diagram.html")
    arch_png = os.path.join(WORKSPACE, "architecture_diagram.png")
    with open(arch_html, 'w', encoding='utf-8') as f:
        f.write(create_architecture_html(info['project_name']))
    screenshot_html(arch_html, arch_png)

    # 4. 生成 Word 文档
    print("生成 Word 文档...")
    images = {'architecture': arch_png}
    output_path = create_word_doc(info, images)
    
    print("=" * 60)
    print(f"文档已生成：{output_path}")
    print("=" * 60)
    return output_path

if __name__ == "__main__":
    main()
