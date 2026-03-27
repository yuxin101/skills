#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测绘/GIS 校招信息整理脚本
生成 Word 文档格式的校招日报
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_job_report():
    """创建测绘/GIS 校招日报 Word 文档"""
    
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('测绘/GIS 专业校招日报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = '黑体'
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'生成日期：2026 年 3 月 24 日')
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空行
    
    # 说明
    intro = doc.add_paragraph()
    intro_run = intro.add_run('【说明】本日报汇总近期测绘工程、地理信息系统（GIS）、遥感科学与技术等相关专业校园招聘信息，按单位类型优先级排序：国央企 > 事业单位 > 大厂 > 其他企业。')
    intro_run.font.size = Pt(9)
    intro_run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()  # 空行
    
    # ============ 第一部分：国央企 ============
    h1 = doc.add_heading('一、国央企单位', level=1)
    h1_run = h1.runs[0]
    h1_run.font.name = '黑体'
    h1_run.font.color.rgb = RGBColor(192, 0, 0)
    
    # 1.1 中国土木工程集团
    doc.add_heading('1.1 中国土木工程集团有限公司', level=2)
    doc.add_paragraph('📌 岗位类别：工程技术岗')
    doc.add_paragraph('🎓 专业要求：测绘工程、摄影测量与遥感、土木工程、交通工程、地质工程等相关专业')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：全国多地（根据项目分配）')
    doc.add_paragraph('🔗 信息来源：鼠鼠求职')
    doc.add_paragraph('💡 备注：央企，工作稳定，适合能接受外派的同学')
    
    # 1.2 中铁/中建/中交系统
    doc.add_heading('1.2 中铁、中建、中交等基建央企', level=2)
    doc.add_paragraph('📌 岗位类别：测绘工程师、工程技术岗')
    doc.add_paragraph('🎓 专业要求：测绘工程、地理信息系统、土木工程等')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：全国各地项目')
    doc.add_paragraph('🔗 信息来源：各大高校就业网、国聘网')
    doc.add_paragraph('💡 备注：2026 年春招扩招趋势明显，建议关注各单位官网')
    
    # 1.3 国家电网/南方电网
    doc.add_heading('1.3 国家电网、南方电网', level=2)
    doc.add_paragraph('📌 岗位类别：电力测绘、线路规划')
    doc.add_paragraph('🎓 专业要求：测绘工程、电气工程、线路规划等')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：各省市电力公司')
    doc.add_paragraph('💡 备注：待遇优厚，竞争较激烈，需参加统一笔试')
    
    doc.add_paragraph()  # 空行
    
    # ============ 第二部分：事业单位 ============
    h2 = doc.add_heading('二、事业单位', level=1)
    h2_run = h2.runs[0]
    h2_run.font.name = '黑体'
    h2_run.font.color.rgb = RGBColor(192, 0, 0)
    
    # 2.1 中国地质调查局
    doc.add_heading('2.1 中国地质调查局局属单位', level=2)
    doc.add_paragraph('📌 岗位类别：地质调查、测绘工程')
    doc.add_paragraph('🎓 专业要求：测绘工程、地质工程、遥感科学与技术等')
    doc.add_paragraph('📚 学历：本科及以上（部分岗位要求硕士）')
    doc.add_paragraph('📅 笔试时间：2026 年 3 月已安排统一笔试')
    doc.add_paragraph('📍 考点：北京、武汉、成都、乌鲁木齐')
    doc.add_paragraph('🔗 报名平台：https://www.sydwgkzp.cn/cgs/')
    doc.add_paragraph('💡 备注：自然资源部下属，事业编制，稳定性高')
    
    # 2.2 各省市测绘院
    doc.add_heading('2.2 各省市测绘院/自然资源部门', level=2)
    doc.add_paragraph('📌 岗位类别：测绘工程师、地理信息工程师')
    doc.add_paragraph('🎓 专业要求：测绘工程、地理信息系统、遥感等')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：各省市')
    doc.add_paragraph('💡 备注：关注各省人社厅官网、自然资源厅官网')
    
    # 2.3 浙江省属事业单位
    doc.add_heading('2.3 浙江省属事业单位（2026 年上半年）', level=2)
    doc.add_paragraph('📌 岗位类别：各类专业技术岗')
    doc.add_paragraph('🎓 专业要求：包含测绘、GIS 相关专业')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📅 时间：2026 年上半年集中招聘')
    doc.add_paragraph('🔗 信息来源：浙江省人社厅')
    doc.add_paragraph('💡 备注：含部属驻浙事业单位）')
    
    doc.add_paragraph()  # 空行
    
    # ============ 第三部分：大厂/知名企业 ============
    h3 = doc.add_heading('三、大厂/知名企业', level=1)
    h3_run = h3.runs[0]
    h3_run.font.name = '黑体'
    h3_run.font.color.rgb = RGBColor(192, 0, 0)
    
    # 3.1 超图软件
    doc.add_heading('3.1 超图软件（SuperMap）', level=2)
    doc.add_paragraph('📌 岗位类别：GIS 开发工程师、技术支持、销售经理')
    doc.add_paragraph('🎓 专业要求：地理信息系统、计算机、测绘工程等')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：北京、上海、广州、成都等')
    doc.add_paragraph('💡 备注：国内 GIS 龙头企业，技术实力强')
    
    # 3.2 航天宏图
    doc.add_heading('3.2 航天宏图', level=2)
    doc.add_paragraph('📌 岗位类别：遥感算法工程师、GIS 开发、项目经理')
    doc.add_paragraph('🎓 专业要求：遥感、GIS、计算机、测绘等')
    doc.add_paragraph('📚 学历：本科及以上（研发岗偏好硕士）')
    doc.add_paragraph('📍 工作地点：北京、西安、成都等')
    doc.add_paragraph('💡 备注：科创板上市公司，遥感领域领先')
    
    # 3.3 中科星图
    doc.add_heading('3.3 中科星图（688568）', level=2)
    doc.add_paragraph('📌 岗位类别：数字地球产品研发、GIS 开发')
    doc.add_paragraph('🎓 专业要求：GIS、计算机、测绘、遥感等')
    doc.add_paragraph('📚 学历：本科及以上')
    doc.add_paragraph('📍 工作地点：北京、合肥等')
    doc.add_paragraph('💡 备注：中科院背景，科创板上市，2026 年调入科创 50 指数')
    
    # 3.4 互联网大厂地图部门
    doc.add_heading('3.4 百度/腾讯/高德地图部门', level=2)
    doc.add_paragraph('📌 岗位类别：地图数据工程师、GIS 开发、导航算法')
    doc.add_paragraph('🎓 专业要求：GIS、计算机、测绘、数学等')
    doc.add_paragraph('📚 学历：本科及以上（核心研发岗偏好硕士）')
    doc.add_paragraph('📍 工作地点：北京、上海、深圳等')
    doc.add_paragraph('💡 备注：薪资待遇好，竞争激烈，需较强编程能力')
    
    doc.add_paragraph()  # 空行
    
    # ============ 第四部分：招聘平台汇总 ============
    h4 = doc.add_heading('四、招聘平台汇总', level=1)
    h4_run = h4.runs[0]
    h4_run.font.name = '黑体'
    h4_run.font.color.rgb = RGBColor(192, 0, 0)
    
    doc.add_heading('4.1 核心招聘渠道', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '平台名称'
    hdr_cells[1].text = '特点'
    hdr_cells[2].text = '适用场景'
    
    # 加粗表头
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 数据行
    data = [
        ('国聘网', '国企央企主渠道，岗位真实可靠', '国央企/事业单位招聘'),
        ('智联招聘', '资源丰富，校招体系成熟', '综合求职，国企岗位多'),
        ('BOSS 直聘', '直接沟通，反馈快', '私企/中小企业'),
        ('高校就业网', '针对性强，信息准确', '本校/目标院校招聘'),
        ('小林 coding', '校招汇总表格，支持搜索', '一站式查看多家企业')
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
    
    doc.add_paragraph()  # 空行
    
    # ============ 第五部分：求职建议 ============
    h5 = doc.add_heading('五、求职建议', level=1)
    h5_run = h5.runs[0]
    h5_run.font.name = '黑体'
    h5_run.font.color.rgb = RGBColor(192, 0, 0)
    
    suggestions = [
        '✅ 优先关注国央企和事业单位，稳定性高，福利待遇好',
        '✅ 提前准备笔试面试，国央企多有统一笔试环节',
        '✅ 关注目标单位官网和官方公众号，获取第一手信息',
        '✅ 准备好简历和作品集，GIS 岗位可展示项目经验',
        '✅ 注意春招时间节点，4-5 月为收尾阶段，抓紧投递',
        '✅ 警惕"付费内推""定岗保录"等骗局，通过正规渠道求职'
    ]
    
    for suggestion in suggestions:
        p = doc.add_paragraph(suggestion)
        p.style = 'List Bullet'
    
    doc.add_paragraph()  # 空行
    
    # 页脚
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f'—— 测绘/GIS 校招日报 · 2026 年 3 月 24 日 ——')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    # 保存文件
    output_path = '/root/.openclaw/workspace/references/gis-jobs-daily-20260324.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    return output_path

if __name__ == '__main__':
    path = create_job_report()
    print(f'Word 文档已生成：{path}')
