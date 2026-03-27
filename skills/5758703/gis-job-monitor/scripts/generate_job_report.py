#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测绘/GIS 校招信息日报生成脚本
生成 Word 文档格式的校招信息汇总报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import sys

def create_job_report():
    doc = Document()
    
    # 设置标题样式
    title = doc.add_heading('测绘/GIS 专业校招信息日报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'报告日期：2026 年 3 月 22 日')
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # 空行
    
    # 添加说明
    intro = doc.add_paragraph('本报告汇总了近期发布的测绘工程、地理信息系统（GIS）、遥感科学与技术等相关专业校园招聘信息，按单位类型优先级排序：国央企 > 事业单位 > 行业龙头企业 > 其他企业。')
    intro.runs[0].font.size = Pt(10)
    
    # === 第一部分：国央企 ===
    doc.add_heading('一、国央企单位', level=1)
    
    # 1.1 中国土木工程集团有限公司
    doc.add_heading('1.1 中国土木工程集团有限公司（中国铁建旗下）', level=2)
    doc.add_paragraph('【单位性质】央企 / 世界 500 强 / 海外龙头企业')
    doc.add_paragraph('【招聘岗位】')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '岗位类别'
    hdr_cells[1].text = '专业要求'
    hdr_cells[2].text = '学历'
    hdr_cells[3].text = '工作地点'
    
    jobs = [
        ('工程技术岗', '测绘工程、摄影测量与遥感、土木工程等', '本科及以上', '海外 110+ 国家'),
        ('战新产业岗', '新能源科学与工程、环境科学等', '本科及以上', '海外区域市场'),
        ('测量储备干部', '土木工程、测绘学等', '本科及以上', '广东/海南/河北/海外'),
    ]
    
    for job in jobs:
        row_cells = table.add_row().cells
        for i, text in enumerate(job):
            row_cells[i].text = text
    
    doc.add_paragraph('【招聘条件】2026 届本科及以上学历，英语四级/六级，服从境内外项目安排')
    doc.add_paragraph('【福利待遇】六险二金、海外津贴 6500 元/月、绩效奖金（海外双倍）、提供食宿')
    doc.add_paragraph('【投递方式】公众号"中国土木招贤纳士馆"或 https://ccecc.zhaopin.com/')
    doc.add_paragraph('【截止日期】2026 年 7 月')
    doc.add_paragraph()
    
    # 1.2 中国中铁
    doc.add_heading('1.2 中国中铁 / 中国铁建系统', level=2)
    doc.add_paragraph('【单位性质】央企 / 基建龙头')
    doc.add_paragraph('【招聘岗位】测绘工程、铁道工程、桥梁与隧道工程等相关专业')
    doc.add_paragraph('【学历要求】本科及以上')
    doc.add_paragraph('【投递提示】关注"中国中铁招聘"公众号获取最新岗位信息')
    doc.add_paragraph()
    
    # 1.3 中建一局
    doc.add_heading('1.3 中建一局', level=2)
    doc.add_paragraph('【单位性质】央企 / 建筑行业骨干')
    doc.add_paragraph('【招聘岗位】房屋建筑、基础设施、勘察设计等')
    doc.add_paragraph('【专业要求】土木工程、建筑设计、工程管理、测绘相关')
    doc.add_paragraph('【工作地点】全国各地及"一带一路"沿线国家')
    doc.add_paragraph()
    
    # === 第二部分：事业单位 ===
    doc.add_heading('二、事业单位', level=1)
    
    # 2.1 中国地质调查局
    doc.add_heading('2.1 中国地质调查局局属单位', level=2)
    doc.add_paragraph('【单位性质】自然资源部直属事业单位')
    doc.add_paragraph('【招聘状态】2026 年度第一批公开招聘已启动')
    doc.add_paragraph('【笔试安排】已设置北京、武汉、成都、乌鲁木齐考点')
    doc.add_paragraph('【投递平台】https://www.sydwgkzp.cn/cgs/')
    doc.add_paragraph('【专业方向】地质工程、测绘工程、遥感、GIS 相关')
    doc.add_paragraph()
    
    # 2.2 各地自然资源部门
    doc.add_heading('2.2 地方自然资源系统事业单位', level=2)
    doc.add_paragraph('【无锡市自然资源和规划局】下属单位 2026 年公开选聘')
    doc.add_paragraph('【南京市事业单位】2026 年统一公开招聘，含自然资源相关岗位')
    doc.add_paragraph('【北京市门头沟区】2026 年上半年事业单位公开招聘')
    doc.add_paragraph('【招聘对象】2025-2026 届毕业生，部分岗位限北京生源')
    doc.add_paragraph()
    
    # 2.3 中国测绘科学研究院（信息待确认）
    doc.add_heading('2.3 中国测绘科学研究院', level=2)
    doc.add_paragraph('【单位性质】自然资源部直属科研事业单位')
    doc.add_paragraph('【提示】建议关注官网及"国聘网"获取最新招聘公告')
    doc.add_paragraph('【专业方向】测绘工程、遥感、GIS、摄影测量')
    doc.add_paragraph()
    
    # === 第三部分：行业龙头企业 ===
    doc.add_heading('三、行业龙头企业（GIS/遥感）', level=1)
    
    # 3.1 超图软件
    doc.add_heading('3.1 超图软件（SuperMap）', level=2)
    doc.add_paragraph('【单位性质】上市公司 / GIS 软件龙头')
    doc.add_paragraph('【招聘岗位】GIS 开发工程师、技术支持工程师、产品经理')
    doc.add_paragraph('【专业要求】地理信息系统、计算机、遥感、测绘工程')
    doc.add_paragraph('【学历要求】本科及以上')
    doc.add_paragraph('【投递提示】关注"超图软件招聘"公众号')
    doc.add_paragraph()
    
    # 3.2 航天宏图
    doc.add_heading('3.2 航天宏图', level=2)
    doc.add_paragraph('【单位性质】科创板上市公司 / 遥感龙头')
    doc.add_paragraph('【招聘岗位】遥感算法工程师、GIS 开发、卫星应用')
    doc.add_paragraph('【专业要求】遥感科学与技术、GIS、计算机、测绘')
    doc.add_paragraph('【学历要求】本科及以上，硕士优先')
    doc.add_paragraph()
    
    # 3.3 中科星图
    doc.add_heading('3.3 中科星图', level=2)
    doc.add_paragraph('【单位性质】科创板 / 中科院背景')
    doc.add_paragraph('【招聘岗位】数字地球产品研发、GIS 工程师')
    doc.add_paragraph('【专业要求】GIS、计算机、遥感、测绘')
    doc.add_paragraph()
    
    # === 第四部分：其他优质企业 ===
    doc.add_heading('四、其他优质企业', level=1)
    
    doc.add_heading('4.1 互联网大厂地图相关部门', level=2)
    doc.add_paragraph('【百度地图】地图数据、导航算法、GIS 开发')
    doc.add_paragraph('【高德地图（阿里）】地图引擎、位置服务')
    doc.add_paragraph('【腾讯地图】位置大数据、导航产品')
    doc.add_paragraph('【华为】河图 Cyberspace、地图业务')
    doc.add_paragraph('【提示】2026 春招 AI 岗位激增，建议关注各公司校招官网')
    doc.add_paragraph()
    
    doc.add_heading('4.2 其他测绘地理信息企业', level=2)
    doc.add_paragraph('【南方测绘】测绘仪器、GIS 软件')
    doc.add_paragraph('【华测导航】GNSS 接收机、无人机测绘')
    doc.add_paragraph('【中海达】高精度定位、海洋测绘')
    doc.add_paragraph()
    
    # === 第五部分：投递建议 ===
    doc.add_heading('五、投递建议与时间线', level=1)
    
    doc.add_paragraph('【春招时间线】')
    doc.add_paragraph('• 3 月 -4 月：春招高峰期，大部分岗位集中发布')
    doc.add_paragraph('• 4 月 -5 月：部分企业补招，竞争相对较小')
    doc.add_paragraph('• 5 月 -6 月：收尾阶段，关注未招满岗位')
    doc.add_paragraph()
    
    doc.add_paragraph('【投递渠道优先级】')
    doc.add_paragraph('1. 企业官方校招公众号/官网（信息最准确）')
    doc.add_paragraph('2. 国聘网（www.iguopin.com）- 国企事业单位主渠道')
    doc.add_paragraph('3. 高校就业网 - 武大、同济、矿大等测绘强校')
    doc.add_paragraph('4. 智联招聘/BOSS 直聘（筛选国企标签）')
    doc.add_paragraph()
    
    doc.add_paragraph('【简历准备建议】')
    doc.add_paragraph('• 突出测绘/GIS 相关课程、项目经历、实习经验')
    doc.add_paragraph('• 准备作品集：GIS 项目、地图作品、数据分析报告')
    doc.add_paragraph('• 国央企重视英语等级（四级/六级）、学生干部经历')
    doc.add_paragraph('• 技术岗准备：编程能力（Python/C++）、GIS 软件操作')
    doc.add_paragraph()
    
    # 添加页脚
    doc.add_paragraph()
    doc.add_paragraph('---')
    footer = doc.add_paragraph('报告生成时间：2026 年 3 月 22 日 9:00 | 数据来源：国聘网、企业官网、招聘平台')
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    output_path = '/tmp/openclaw/测绘 GIS 校招日报_20260322.docx'
    doc.save(output_path)
    print(f'报告已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_job_report()
