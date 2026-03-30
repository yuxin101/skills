#!/usr/bin/env python3
"""
md_to_docx - 将 Markdown 转换为 Word 文档，支持自动模板格式适配。

用法:
    python3 md_to_docx.py input.md -o output.docx
    python3 md_to_docx.py input.md -o output.docx -t template.doc

功能:
    - 解析 Markdown 标题、正文、列表、表格、代码块、引用块
    - 将 Mermaid 图表渲染为 PNG 并嵌入文档
    - 自动从 .doc/.docx 模板提取格式（字体、字号、对齐、缩进、间距）
    - 无模板时使用中文排版默认格式（宋体/黑体、五号字、首行缩进两字符）
"""

import argparse
import os
import platform
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 默认格式（无模板时的中文排版标准）
# ============================================================

def default_styles():
    """返回默认的中文排版格式。"""
    return {
        'chapter_title': {
            'font_name': '黑体', 'cn_font': '黑体',
            'font_size': Pt(18), 'alignment': WD_ALIGN_PARAGRAPH.CENTER,
            'space_before': Pt(12), 'space_after': Pt(12), 'bold': True,
        },
        'section_title': {
            'font_name': '黑体', 'cn_font': '黑体',
            'font_size': Pt(16), 'alignment': WD_ALIGN_PARAGRAPH.CENTER,
            'space_before': Pt(8), 'space_after': Pt(8), 'bold': True,
        },
        'subsection_title': {
            'font_name': '黑体', 'cn_font': '黑体',
            'font_size': Pt(14), 'alignment': WD_ALIGN_PARAGRAPH.LEFT,
            'space_before': Pt(6), 'space_after': Pt(6), 'bold': True,
        },
        'heading4': {
            'font_name': '楷体', 'cn_font': '楷体',
            'font_size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.LEFT,
            'space_before': Pt(4), 'space_after': Pt(4), 'bold': True,
        },
        'body': {
            'font_name': '宋体', 'cn_font': '宋体',
            'font_size': Pt(10.5), 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'first_line_indent': Cm(0.74),
        },
        'quote': {
            'font_name': '楷体', 'cn_font': '楷体',
            'font_size': Pt(10.5), 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'left_indent': Cm(1), 'italic': True,
            'color': RGBColor(0x55, 0x55, 0x55),
        },
        'list': {
            'font_name': '宋体', 'cn_font': '宋体',
            'font_size': Pt(10.5), 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'left_indent': Cm(0.74),
        },
        'code': {
            'font_name': 'Consolas', 'cn_font': 'Consolas',
            'font_size': Pt(9), 'alignment': WD_ALIGN_PARAGRAPH.LEFT,
            'left_indent': Cm(0.74),
        },
        'page': {
            'width': Cm(21), 'height': Cm(29.7),
            'margin_top': Cm(2.54), 'margin_bottom': Cm(2.54),
            'margin_left': Cm(3.17), 'margin_right': Cm(3.17),
        },
    }


# ============================================================
# 模板格式提取
# ============================================================

def convert_doc_to_docx(doc_path: str) -> str:
    """将 .doc 转换为 .docx（使用 textutil 或 LibreOffice）。"""
    # 如果已经是 .docx，直接返回
    if doc_path.lower().endswith('.docx'):
        return doc_path

    # macOS textutil
    converted = doc_path.rsplit('.', 1)[0] + '_converted.docx'
    if platform.system() == 'Darwin':
        try:
            result = subprocess.run(
                ['textutil', '-convert', 'docx', doc_path, '-output', converted],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and os.path.exists(converted):
                return converted
        except Exception:
            pass

    # LibreOffice (cross-platform fallback)
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir',
             os.path.dirname(converted) or '.', doc_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and os.path.exists(converted):
            return converted
    except Exception:
        pass

    return None


def extract_styles_from_template(docx_path: str) -> dict:
    """
    从 Word 模板自动提取格式。

    策略：分析模板中各段落的格式特征，按字号和对齐方式分类，
    识别出标题、正文等不同层级的格式。

    Returns:
        dict: 与 default_styles() 结构相同的格式字典
    """
    styles = default_styles()

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  [WARN] 无法读取模板: {e}，使用默认格式")
        return styles

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return styles

    # 收集所有段落的信息
    para_info = []
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        runs = para.runs
        if not runs:
            continue

        # 取第一个 run 的字体信息
        run = runs[0]
        font_name = run.font.name or 'Times New Roman'
        font_size = run.font.size
        if font_size:
            font_size_pt = font_size.pt
        else:
            # 从 XML 获取
            sz = run._element.find(qn('w:rPr'))
            if sz is not None:
                sz_elem = sz.find(qn('w:sz'))
                if sz_elem is not None:
                    font_size_pt = int(sz_elem.get(qn('w:val'), '24')) / 2
                else:
                    font_size_pt = 12
            else:
                font_size_pt = 12

        alignment = para.alignment
        pf = para.paragraph_format

        # 获取缩进信息
        first_line_indent = pf.first_line_indent
        left_indent = pf.left_indent

        # 获取间距
        space_before = pf.space_before
        space_after = pf.space_after

        # 中文字体
        rPr = run._element.find(qn('w:rPr'))
        cn_font = None
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                cn_font = rFonts.get(qn('w:eastAsia'))

        para_info.append({
            'index': i, 'text': text[:60],
            'font_name': font_name, 'font_size_pt': font_size_pt,
            'alignment': alignment, 'bold': run.font.bold,
            'first_line_indent': first_line_indent,
            'left_indent': left_indent,
            'space_before': space_before, 'space_after': space_after,
            'cn_font': cn_font,
        })

    if not para_info:
        return styles

    # 按字号排序，找出不同层级的标题
    # 策略：字号最大的居中段落 = 章标题，次大居中 = 节标题，以此类推

    # 找所有居中段落
    centered = [p for p in para_info if p['alignment'] == WD_ALIGN_PARAGRAPH.CENTER]
    non_centered = [p for p in para_info if p['alignment'] != WD_ALIGN_PARAGRAPH.CENTER]

    # 按字号排序（降序）
    centered.sort(key=lambda p: p['font_size_pt'], reverse=True)
    non_centered.sort(key=lambda p: p['font_size_pt'], reverse=True)

    # 识别标题层级
    title_levels = [
        ('chapter_title', centered, Pt(18)),
        ('section_title', centered, Pt(16)),
        ('subsection_title', non_centered, Pt(14)),
        ('heading4', non_centered, Pt(12)),
    ]

    assigned_fonts = set()

    for level_name, pool, default_size in title_levels:
        if pool:
            # 找池中字号最大的
            candidate = pool[0]
            # 从池中移除已使用的
            if candidate in pool:
                pool.remove(candidate)

            styles[level_name] = {
                'font_name': candidate['font_name'],
                'cn_font': candidate.get('cn_font') or candidate['font_name'],
                'font_size': Pt(candidate['font_size_pt']),
                'alignment': candidate['alignment'] or styles[level_name]['alignment'],
                'bold': candidate['bold'] if candidate['bold'] is not None else True,
            }
            if candidate['space_before']:
                styles[level_name]['space_before'] = candidate['space_before']
            if candidate['space_after']:
                styles[level_name]['space_after'] = candidate['space_after']

            assigned_fonts.add(candidate['font_name'])

    # 识别正文格式：有首行缩进的最大字号段落
    body_candidates = [p for p in non_centered
                       if p['first_line_indent'] and p['font_size_pt'] <= 12]
    if body_candidates:
        body = body_candidates[0]
        styles['body'] = {
            'font_name': body['font_name'],
            'cn_font': body.get('cn_font') or body['font_name'],
            'font_size': Pt(body['font_size_pt']),
            'alignment': body['alignment'] or WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if body['first_line_indent']:
            styles['body']['first_line_indent'] = body['first_line_indent']

        # 列表和引用继承正文字体
        styles['list']['font_name'] = body['font_name']
        styles['list']['cn_font'] = body.get('cn_font') or body['font_name']
        styles['list']['font_size'] = Pt(body['font_size_pt'])

        styles['quote']['font_name'] = body['font_name']
        styles['quote']['cn_font'] = body.get('cn_font') or body['font_name']
        styles['quote']['font_size'] = Pt(body['font_size_pt'])

    # 页面设置
    if doc.sections:
        section = doc.sections[0]
        styles['page'] = {
            'width': section.page_width,
            'height': section.page_height,
            'margin_top': section.top_margin,
            'margin_bottom': section.bottom_margin,
            'margin_left': section.left_margin,
            'margin_right': section.right_margin,
        }

    return styles


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown(text: str) -> list:
    """
    将 Markdown 文本解析为结构化块列表。

    支持的块类型:
        h1-h6, body, bullet_list, ordered_list, table,
        code_block, mermaid, quote, hr
    """
    blocks = []
    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # 标题 (h1-h6)
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            # 映射到语义化的类型名
            type_map = {1: 'h1', 2: 'h2', 3: 'h3', 4: 'h4', 5: 'h5', 6: 'h6'}
            blocks.append({'type': type_map[level], 'content': content})
            i += 1
            continue

        # 分隔线
        if re.match(r'^---+\s*$', line.strip()):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # 引用块
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].strip()))
                i += 1
            blocks.append({'type': 'quote', 'content': '\n'.join(quote_lines)})
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({'type': 'table', 'content': table_lines})
            continue

        # 代码块
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if lang.lower() == 'mermaid':
                blocks.append({'type': 'mermaid', 'content': '\n'.join(code_lines)})
            else:
                blocks.append({'type': 'code_block', 'content': '\n'.join(code_lines), 'lang': lang})
            continue

        # 无序列表
        if re.match(r'^(\s*)[-*]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^(\s*)[-*]\s+', lines[i]):
                item_text = re.sub(r'^(\s*)[-*]\s+', '', lines[i])
                indent = len(lines[i]) - len(lines[i].lstrip())
                list_items.append({'text': item_text, 'indent': indent})
                i += 1
            blocks.append({'type': 'bullet_list', 'content': list_items})
            continue

        # 有序列表
        if re.match(r'^(\s*)\d+[.)]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^(\s*)\d+[.)]\s+', lines[i]):
                item_text = re.sub(r'^(\s*)\d+[.)]\s+', '', lines[i])
                indent = len(lines[i]) - len(lines[i].lstrip())
                list_items.append({'text': item_text, 'indent': indent})
                i += 1
            blocks.append({'type': 'ordered_list', 'content': list_items})
            continue

        # 普通正文
        body_lines = []
        while i < len(lines):
            cur = lines[i]
            if not cur.strip():
                break
            if (cur.strip().startswith('#') or cur.strip().startswith('>')
                    or cur.strip().startswith('```')):
                break
            if re.match(r'^(\s*)[-*]\s+', cur) or re.match(r'^(\s*)\d+[.)]\s+', cur):
                break
            if re.match(r'^---+\s*$', cur.strip()):
                break
            if '|' in cur and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].strip()):
                break
            body_lines.append(cur)
            i += 1

        if body_lines:
            blocks.append({'type': 'body', 'content': '\n'.join(body_lines)})
            continue

        i += 1

    return blocks


# ============================================================
# Word 文档生成
# ============================================================

def get_cn_font(font_name: str, styles: dict, level: str = 'body') -> str:
    """获取中文字体名。"""
    level_style = styles.get(level, {})
    if level_style.get('cn_font'):
        return level_style['cn_font']
    # 常见映射
    cn_map = {
        'Times New Roman': '宋体', 'Arial': '黑体', 'Helvetica': '黑体',
        'Consolas': 'Consolas', 'Courier New': 'Courier New',
        'Georgia': '宋体', 'Verdana': '宋体',
    }
    return cn_map.get(font_name, '宋体')


def set_run_font(run, font_name: str, font_size, styles: dict,
                 level: str = 'body', bold=None, italic=None, color=None):
    """设置 run 的字体（含中文字体）。"""
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color

    # 写入 w:rFonts 中的 eastAsia
    cn_font = get_cn_font(font_name, styles, level)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)


def get_heading_style_key(level: int) -> str:
    """将标题级别映射到样式键名。"""
    return {1: 'chapter_title', 2: 'section_title', 3: 'subsection_title'}.get(
        level, 'heading4')


def add_paragraph(doc, text: str, styles: dict, level: str, runs_data: list = None):
    """添加一个格式化段落。"""
    fmt = styles.get(level, styles['body'])
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = fmt.get('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)

    if 'space_before' in fmt:
        pf.space_before = fmt['space_before']
    if 'space_after' in fmt:
        pf.space_after = fmt['space_after']
    if 'first_line_indent' in fmt:
        pf.first_line_indent = fmt['first_line_indent']
    if 'left_indent' in fmt:
        pf.left_indent = fmt['left_indent']

    if runs_data:
        for rd in runs_data:
            run = para.add_run(rd.get('text', ''))
            set_run_font(
                run, fmt.get('font_name', '宋体'), fmt.get('font_size', Pt(10.5)),
                styles, level,
                bold=rd.get('bold', fmt.get('bold')),
                italic=rd.get('italic', fmt.get('italic')),
                color=rd.get('color', fmt.get('color')),
            )
    else:
        run = para.add_run(text)
        set_run_font(
            run, fmt.get('font_name', '宋体'), fmt.get('font_size', Pt(10.5)),
            styles, level,
            bold=fmt.get('bold'), italic=fmt.get('italic'), color=fmt.get('color'),
        )
    return para


def parse_inline(text: str) -> list:
    """解析 Markdown 行内格式。"""
    runs = []
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|~~(.+?)~~|`(.+?)`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        m = re.match(r'^\*\*\*(.+)\*\*\*$', part)
        if m:
            runs.append({'text': m.group(1), 'bold': True, 'italic': True})
            continue
        m = re.match(r'^\*\*(.+)\*\*$', part)
        if m:
            runs.append({'text': m.group(1), 'bold': True})
            continue
        m = re.match(r'^\*(.+)\*$', part)
        if m:
            runs.append({'text': m.group(1), 'italic': True})
            continue
        m = re.match(r'^~~(.+)~~$', part)
        if m:
            runs.append({'text': m.group(1), 'strikethrough': True})
            continue
        m = re.match(r'^`(.+)`$', part)
        if m:
            runs.append({'text': m.group(1), 'code': True})
            continue
        runs.append({'text': part})
    return runs


def add_body_text(doc, text: str, styles: dict, level: str = 'body'):
    """添加正文段落，保留行内格式。"""
    fmt = styles.get(level, styles['body'])
    runs_data = parse_inline(text)

    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = fmt.get('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)

    if 'first_line_indent' in fmt:
        pf.first_line_indent = fmt['first_line_indent']
    if 'left_indent' in fmt:
        pf.left_indent = fmt['left_indent']
    if 'space_before' in fmt:
        pf.space_before = fmt['space_before']
    if 'space_after' in fmt:
        pf.space_after = fmt['space_after']

    for rd in runs_data:
        run = para.add_run(rd.get('text', ''))
        if rd.get('code'):
            set_run_font(run, 'Consolas', fmt.get('font_size', Pt(10.5)), styles, 'code')
        else:
            set_run_font(
                run, fmt.get('font_name', '宋体'), fmt.get('font_size', Pt(10.5)),
                styles, level,
                bold=rd.get('bold', fmt.get('bold')),
                italic=rd.get('italic', fmt.get('italic')),
                color=rd.get('color', fmt.get('color')),
            )
            if rd.get('strikethrough'):
                run.font.strike = True
    return para


def add_table(doc, table_lines: list, styles: dict):
    """从 Markdown 表格行添加 Word 表格。"""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows_data.append(cells)

    if len(rows_data) < 2:
        return

    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data) - 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    body_fmt = styles.get('body', default_styles()['body'])
    row_idx = 0
    for data_idx, row_data in enumerate(rows_data):
        if data_idx == 1:
            continue
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            font_size = body_fmt.get('font_size', Pt(10.5))
            if data_idx == 0:
                set_run_font(run, body_fmt.get('font_name', '宋体'), font_size, styles, 'body', bold=True)
            else:
                set_run_font(run, body_fmt.get('font_name', '宋体'), font_size, styles, 'body')
        row_idx += 1


def find_chrome_path() -> str:
    """跨平台查找 Chrome/Chromium 可执行文件路径。"""
    system = platform.system()

    candidates = []

    if system == 'Darwin':
        candidates = [
            '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome',
            '/Applications/Chromium.app/Contents/MacOS/Chromium',
            os.path.expanduser('~/Applications/Google Chrome.app/Contents/MacOS/Google Chrome'),
        ]
    elif system == 'Windows':
        candidates = [
            os.path.join(os.environ.get('PROGRAMFILES', ''), 'Google/Chrome/Application/chrome.exe'),
            os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'Google/Chrome/Application/chrome.exe'),
            os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Google/Chrome/Application/chrome.exe'),
        ]
    elif system == 'Linux':
        candidates = [
            '/usr/bin/google-chrome',
            '/usr/bin/chromium-browser',
            '/usr/bin/chromium',
            '/snap/bin/chromium',
        ]

    for path in candidates:
        if os.path.exists(path):
            return path

    return None


def render_mermaid_to_png(mermaid_code: str, output_path: str, img_dir: str) -> bool:
    """使用 mmdc 将 Mermaid 代码渲染为 PNG。"""
    cleaned = mermaid_code.strip()

    # 预处理：修复常见的 Mermaid 兼容性问题
    # subgraph 名称中的冒号会导致解析错误
    cleaned = re.sub(
        r'subgraph\s+([^\n{]+?):\s*',
        lambda m: f'subgraph {m.group(1).strip()}["{m.group(0).strip()}"]\n',
        cleaned
    )
    # subgraph 名称中的括号也需要引号
    cleaned = re.sub(
        r'subgraph\s+([^\n{]*[()（）][^\n{]*)\s*\n',
        lambda m: f'subgraph sg_{hash(m.group(1)) % 10000}["{m.group(1).strip()}"]\n',
        cleaned
    )

    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
        f.write(cleaned)
        mmd_file = f.name

    # Puppeteer 配置
    puppeteer_config = '{"args":["--no-sandbox","--disable-setuid-sandbox","--disable-gpu","--disable-dev-shm-usage"]}'
    config_file = os.path.join(img_dir, '.puppeteer-config.json')
    with open(config_file, 'w') as f:
        f.write(puppeteer_config)

    try:
        env = os.environ.copy()
        chrome_path = find_chrome_path()
        if chrome_path:
            env['PUPPETEER_EXECUTABLE_PATH'] = chrome_path

        result = subprocess.run(
            ['mmdc', '-i', mmd_file, '-o', output_path, '-w', '1200', '-H', '800',
             '-b', 'white', '-p', config_file],
            capture_output=True, text=True, timeout=120, env=env
        )
        if result.returncode != 0:
            print(f"  [WARN] Mermaid render failed: {result.stderr[:300]}")
            return False
        return os.path.exists(output_path)
    except subprocess.TimeoutExpired:
        print("  [WARN] Mermaid render timeout (120s)")
        return False
    except FileNotFoundError:
        print("  [WARN] mmdc not found. Install with: npm install -g @mermaid-js/mermaid-cli")
        return False
    except Exception as e:
        print(f"  [WARN] Mermaid render error: {e}")
        return False
    finally:
        try:
            os.unlink(mmd_file)
        except OSError:
            pass


def add_mermaid_image(doc, mermaid_code: str, img_dir: str, img_index: int, styles: dict):
    """将 Mermaid 图表渲染为图片并添加到文档。"""
    img_path = os.path.join(img_dir, f'mermaid_{img_index:02d}.png')

    print(f"  渲染 Mermaid 图表 #{img_index}...")
    success = render_mermaid_to_png(mermaid_code, img_path, img_dir)

    if success:
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(2)
            run = para.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            return True
        except Exception as e:
            print(f"  [WARN] 插入图片失败: {e}")

    # 渲染失败时回退为代码块
    print(f"  [INFO] Mermaid 渲染失败，以代码形式展示")
    code_fmt = styles.get('code', default_styles()['code'])
    add_paragraph(doc, '[Mermaid 图表]', styles, 'code', runs_data=[{
        'text': '[Mermaid 图表]', 'bold': True,
    }])
    for line in mermaid_code.strip().split('\n'):
        add_paragraph(doc, line, styles, 'code')
    return False


def convert_md_to_docx(md_path: str, output_path: str, template_path: str = None):
    """
    将 Markdown 文件转换为 Word 文档。

    Args:
        md_path: Markdown 源文件路径
        output_path: 输出 .docx 文件路径
        template_path: 可选，模板文件路径 (.doc 或 .docx)
    """
    print(f"读取 Markdown: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 提取或使用默认样式
    styles = default_styles()
    if template_path:
        print(f"分析模板: {template_path}")
        docx_path = convert_doc_to_docx(template_path)
        if docx_path:
            styles = extract_styles_from_template(docx_path)
            print("  模板格式提取完成")
        else:
            print(f"  [WARN] 模板转换失败，使用默认格式")

    print("解析 Markdown 结构...")
    blocks = parse_markdown(md_text)

    type_counts = {}
    for b in blocks:
        type_counts[b['type']] = type_counts.get(b['type'], 0) + 1
    print(f"解析完成: {len(blocks)} 个块")
    for t, c in sorted(type_counts.items()):
        print(f"  {t}: {c}")

    # 创建文档
    doc = Document()

    # 设置页面
    page = styles.get('page', default_styles()['page'])
    for section in doc.sections:
        section.page_width = page.get('width', Cm(21))
        section.page_height = page.get('height', Cm(29.7))
        section.top_margin = page.get('margin_top', Cm(2.54))
        section.bottom_margin = page.get('margin_bottom', Cm(2.54))
        section.left_margin = page.get('margin_left', Cm(3.17))
        section.right_margin = page.get('margin_right', Cm(3.17))

    # 图片目录
    img_dir = os.path.join(os.path.dirname(os.path.abspath(output_path)), '.mermaid_images')
    os.makedirs(img_dir, exist_ok=True)

    code_fmt = styles.get('code', default_styles()['code'])
    list_fmt = styles.get('list', default_styles()['list'])

    # 逐块生成
    mermaid_idx = 0
    for block in blocks:
        btype = block['type']
        content = block.get('content', '')

        # 标题 (h1-h6)
        if btype in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(btype[1])
            style_key = get_heading_style_key(level)
            add_paragraph(doc, content, styles, style_key)

        elif btype == 'body':
            add_body_text(doc, content, styles)

        elif btype == 'bullet_list':
            for item in content:
                text = item['text']
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.alignment = list_fmt.get('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)
                pf.left_indent = list_fmt.get('left_indent', Cm(0.74))
                pf.first_line_indent = Cm(-0.37)

                run = para.add_run('\u2022  ')
                set_run_font(run, list_fmt.get('font_name', '宋体'),
                             list_fmt.get('font_size', Pt(10.5)), styles, 'list')

                for rd in parse_inline(text):
                    r = para.add_run(rd.get('text', ''))
                    set_run_font(r, list_fmt.get('font_name', '宋体'),
                                 list_fmt.get('font_size', Pt(10.5)), styles, 'list',
                                 bold=rd.get('bold'), italic=rd.get('italic'))

        elif btype == 'ordered_list':
            for idx, item in enumerate(content, 1):
                text = item['text']
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.alignment = list_fmt.get('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)
                pf.left_indent = list_fmt.get('left_indent', Cm(0.74))
                pf.first_line_indent = Cm(-0.55)

                run = para.add_run(f'{idx}. ')
                set_run_font(run, list_fmt.get('font_name', '宋体'),
                             list_fmt.get('font_size', Pt(10.5)), styles, 'list')

                for rd in parse_inline(text):
                    r = para.add_run(rd.get('text', ''))
                    set_run_font(r, list_fmt.get('font_name', '宋体'),
                                 list_fmt.get('font_size', Pt(10.5)), styles, 'list',
                                 bold=rd.get('bold'), italic=rd.get('italic'))

        elif btype == 'table':
            add_table(doc, content, styles)

        elif btype == 'code_block':
            lang = block.get('lang', '')
            if lang:
                add_paragraph(doc, f'[{lang}]', styles, 'code')
            for line in content.split('\n'):
                add_paragraph(doc, line if line else ' ', styles, 'code')

        elif btype == 'mermaid':
            mermaid_idx += 1
            add_mermaid_image(doc, content, img_dir, mermaid_idx, styles)

        elif btype == 'quote':
            for line in content.split('\n'):
                if line.strip():
                    add_body_text(doc, line, styles, 'quote')

        elif btype == 'hr':
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run('\u2500' * 30)
            set_run_font(run, '宋体', Pt(8), styles, 'body',
                         color=RGBColor(0xCC, 0xCC, 0xCC))

    doc.save(output_path)
    print(f"\n文档已保存: {output_path}")
    if mermaid_idx > 0:
        print(f"Mermaid 图表目录: {img_dir}")


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='Markdown 转 Word 文档，支持模板格式自动适配'
    )
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 .docx 文件路径（默认与输入同目录同名）')
    parser.add_argument('-t', '--template', help='模板文件路径（.doc 或 .docx）')
    args = parser.parse_args()

    md_path = args.input
    if not os.path.exists(md_path):
        print(f"错误: 文件不存在 - {md_path}")
        sys.exit(1)

    if args.output:
        output_path = args.output
    else:
        output_path = os.path.splitext(md_path)[0] + '.docx'

    template_path = args.template
    if template_path and not os.path.exists(template_path):
        print(f"警告: 模板文件不存在 - {template_path}")

    convert_md_to_docx(md_path, output_path, template_path)


if __name__ == '__main__':
    main()
